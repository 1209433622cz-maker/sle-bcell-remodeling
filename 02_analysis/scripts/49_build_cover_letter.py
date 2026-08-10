#!/usr/bin/env python
"""Build a compact editable cover letter for Genome Medicine."""

from __future__ import annotations

import argparse
import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
SOURCE = (
    ROOT
    / "04_submission"
    / "cover_letter_genome_medicine_v3_AUTHOR_COMPLETION_REQUIRED_2026-07-31.md"
)
OUTDIR = ROOT / "04_submission" / "outputs" / "cover_letter_2026-07-31"
OUTPUT = OUTDIR / "Genome_Medicine_Cover_Letter_v3_AUTHOR_COMPLETION_REQUIRED.docx"

FONT = "Times New Roman"
BLACK = RGBColor(0, 0, 0)
ACTION = RGBColor(156, 0, 6)
MUTED = RGBColor(89, 89, 89)
TOKEN = re.compile(r"(\*[^*]+\*|\[AUTHOR ACTION REQUIRED:.*?\])")


def format_run(run, size: float = 10.5, bold: bool | None = None, italic: bool | None = None):
    run.font.name = FONT
    run.font.size = Pt(size)
    run.font.color.rgb = BLACK
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    rpr = run._element.get_or_add_rPr()
    fonts = rpr.rFonts
    if fonts is None:
        fonts = OxmlElement("w:rFonts")
        rpr.insert(0, fonts)
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        fonts.set(qn(f"w:{attr}"), FONT)
    return run


def add_inline(paragraph, text: str) -> None:
    cursor = 0
    for match in TOKEN.finditer(text):
        if match.start() > cursor:
            format_run(paragraph.add_run(text[cursor : match.start()]))
        token = match.group(0)
        if token.startswith("*"):
            format_run(paragraph.add_run(token[1:-1]), italic=True)
        else:
            run = format_run(paragraph.add_run(token), bold=True)
            run.font.color.rgb = ACTION
            run.font.highlight_color = WD_COLOR_INDEX.YELLOW
        cursor = match.end()
    if cursor < len(text):
        format_run(paragraph.add_run(text[cursor:]))


def blocks(text: str) -> list[str]:
    output: list[str] = []
    paragraph: list[str] = []
    for line in text.splitlines():
        stripped = line.strip()
        if stripped.startswith("# "):
            continue
        if not stripped:
            if paragraph:
                output.append(" ".join(paragraph))
                paragraph.clear()
        else:
            paragraph.append(stripped)
    if paragraph:
        output.append(" ".join(paragraph))
    return output


def build(source: Path, output: Path) -> None:
    text = source.read_text(encoding="utf-8")
    if text.count("AUTHOR ACTION REQUIRED") != 8:
        raise ValueError("Expected exactly eight cover-letter author actions")

    document = Document()
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.78)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)
    section.footer_distance = Inches(0.35)

    style = document.styles["Normal"]
    style.font.name = FONT
    style.font.size = Pt(10.5)
    style.font.color.rgb = BLACK
    style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    style.paragraph_format.line_spacing = 1.02
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.widow_control = True

    content = blocks(text)
    for index, block in enumerate(content):
        paragraph = document.add_paragraph()
        if index == 0:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        elif block.startswith("Dear Editors"):
            paragraph.paragraph_format.space_before = Pt(6)
        elif block == "Sincerely,":
            paragraph.paragraph_format.space_before = Pt(4)
        add_inline(paragraph, block)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.add_run("Pre-submission draft | AUTHOR COMPLETION REQUIRED")
    format_run(footer_run, size=8, italic=True)
    footer_run.font.color.rgb = MUTED

    document.core_properties.title = "Genome Medicine cover letter"
    document.core_properties.subject = "Research manuscript cover letter"
    document.core_properties.author = "Corresponding author"
    document.core_properties.comments = "Complete highlighted author actions before submission."
    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)
    print(f"Wrote: {output}")
    print(f"Paragraphs: {len(document.paragraphs)}")
    print(f"Bytes: {output.stat().st_size}")


if __name__ == "__main__":
    parser = argparse.ArgumentParser()
    parser.add_argument("--source", type=Path, default=SOURCE)
    parser.add_argument("--output", type=Path, default=OUTPUT)
    args = parser.parse_args()
    build(args.source, args.output)
