#!/usr/bin/env python
"""Build a reviewer-friendly Genome Medicine DOCX from the v6 Markdown source."""

from __future__ import annotations

import argparse
import json
import re
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
DEFAULT_INPUT = (
    ROOT / "01_manuscript" / "manuscript_v6_genome_medicine_submission_source.md"
)
DEFAULT_OUTDIR = ROOT / "04_submission" / "outputs" / "manuscript_2026-07-31"
DEFAULT_OUTPUT = (
    DEFAULT_OUTDIR
    / "Genome_Medicine_Manuscript_v6_AUTHOR_COMPLETION_REQUIRED.docx"
)

FONT = "Times New Roman"
CODE_FONT = "Consolas"
BLACK = RGBColor(0, 0, 0)
MUTED = RGBColor(89, 89, 89)
ACTION = RGBColor(156, 0, 6)


def set_run_font(run, name: str = FONT, size: float = 12, bold=None, italic=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.color.rgb = BLACK
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        rfonts.set(qn(f"w:{attr}"), name)
    return run


def set_language(run, language: str = "en-US") -> None:
    rpr = run._element.get_or_add_rPr()
    lang = rpr.find(qn("w:lang"))
    if lang is None:
        lang = OxmlElement("w:lang")
        rpr.append(lang)
    lang.set(qn("w:val"), language)


def add_field(paragraph, instruction: str, display: str) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f" {instruction} "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = display
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])
    set_run_font(run, size=9)


def add_hyperlink(paragraph, label: str, url: str):
    relationship = paragraph.part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship)
    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    fonts = OxmlElement("w:rFonts")
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        fonts.set(qn(f"w:{attr}"), FONT)
    size = OxmlElement("w:sz")
    size.set(qn("w:val"), "24")
    rpr.extend([fonts, color, underline, size])
    run.append(rpr)
    text = OxmlElement("w:t")
    text.text = label
    run.append(text)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def configure_section(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.45)
    section.footer_distance = Inches(0.45)

    sect_pr = section._sectPr
    line_numbers = sect_pr.find(qn("w:lnNumType"))
    if line_numbers is None:
        line_numbers = OxmlElement("w:lnNumType")
        sect_pr.append(line_numbers)
    line_numbers.set(qn("w:countBy"), "1")
    line_numbers.set(qn("w:start"), "1")
    line_numbers.set(qn("w:restart"), "continuous")

    header = section.header
    header.is_linked_to_previous = False
    header_paragraph = header.paragraphs[0]
    header_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header_paragraph.paragraph_format.space_after = Pt(0)
    header_run = header_paragraph.add_run("Donor-aware SLE B-cell remodeling")
    set_run_font(header_run, size=9, italic=True)
    header_run.font.color.rgb = MUTED

    footer = section.footer
    footer.is_linked_to_previous = False
    footer_paragraph = footer.paragraphs[0]
    footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer_paragraph.paragraph_format.space_before = Pt(0)
    footer_run = footer_paragraph.add_run("Page ")
    set_run_font(footer_run, size=9)
    add_field(footer_paragraph, "PAGE", "1")


def configure_styles(document: Document) -> None:
    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = FONT
    normal.font.size = Pt(12)
    normal.font.color.rgb = BLACK
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.paragraph_format.line_spacing = 2
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.widow_control = True

    for style_name, size, before, after in (
        ("Heading 1", 14, 14, 6),
        ("Heading 2", 12, 12, 4),
    ):
        style = styles[style_name]
        style.font.name = FONT
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = BLACK
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style.paragraph_format.line_spacing = 2
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True
        style.paragraph_format.widow_control = True

    if "Manuscript Title" not in styles:
        title_style = styles.add_style("Manuscript Title", WD_STYLE_TYPE.PARAGRAPH)
    else:
        title_style = styles["Manuscript Title"]
    title_style.font.name = FONT
    title_style.font.size = Pt(16)
    title_style.font.bold = True
    title_style.font.color.rgb = BLACK
    title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_style.paragraph_format.line_spacing = 1.15
    title_style.paragraph_format.space_after = Pt(12)
    title_style.paragraph_format.keep_with_next = True

    if "Reference Entry" not in styles:
        reference_style = styles.add_style("Reference Entry", WD_STYLE_TYPE.PARAGRAPH)
    else:
        reference_style = styles["Reference Entry"]
    reference_style.base_style = styles["Normal"]
    reference_style.paragraph_format.left_indent = Inches(0.25)
    reference_style.paragraph_format.first_line_indent = Inches(-0.25)
    reference_style.paragraph_format.line_spacing = 2
    reference_style.paragraph_format.space_after = Pt(0)


INLINE_PATTERN = re.compile(
    r"(\*\*.*?\*\*|`.*?`|https?://[^\s)]+(?:\)[^\s.,;:]*)?|"
    r"\[AUTHOR ACTION REQUIRED:.*?\])"
)


def add_inline(paragraph, text: str) -> None:
    cursor = 0
    for match in INLINE_PATTERN.finditer(text):
        if match.start() > cursor:
            run = paragraph.add_run(text[cursor : match.start()])
            set_run_font(run)
            set_language(run)
        token = match.group(0)
        if token.startswith("**") and token.endswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_run_font(run, bold=True)
            set_language(run)
        elif token.startswith("`") and token.endswith("`"):
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, name=CODE_FONT, size=10)
            set_language(run)
        elif token.startswith("http"):
            url = token.rstrip(".,;:")
            trailing = token[len(url) :]
            add_hyperlink(paragraph, url, url)
            if trailing:
                run = paragraph.add_run(trailing)
                set_run_font(run)
        elif token.startswith("[AUTHOR ACTION REQUIRED:"):
            run = paragraph.add_run(token)
            set_run_font(run, bold=True)
            run.font.color.rgb = ACTION
            run.font.highlight_color = WD_COLOR_INDEX.YELLOW
            set_language(run)
        cursor = match.end()
    if cursor < len(text):
        run = paragraph.add_run(text[cursor:])
        set_run_font(run)
        set_language(run)


def markdown_blocks(text: str) -> list[tuple[str, str]]:
    blocks: list[tuple[str, str]] = []
    paragraph_lines: list[str] = []

    def flush() -> None:
        if paragraph_lines:
            blocks.append(("paragraph", " ".join(line.strip() for line in paragraph_lines)))
            paragraph_lines.clear()

    for line in text.splitlines():
        stripped = line.strip()
        if not stripped:
            flush()
            continue
        if stripped.startswith("# "):
            flush()
            blocks.append(("title", stripped[2:].strip()))
        elif stripped.startswith("## "):
            flush()
            blocks.append(("heading1", stripped[3:].strip()))
        elif stripped.startswith("### "):
            flush()
            blocks.append(("heading2", stripped[4:].strip()))
        else:
            paragraph_lines.append(stripped)
    flush()
    return blocks


def build(input_path: Path, output_path: Path) -> dict[str, int | str]:
    text = input_path.read_text(encoding="utf-8")
    if re.search(r"\[@[^\]]+\]", text):
        raise ValueError("Unrendered citation tokens remain in the manuscript source")
    if "AUTHOR ACTION REQUIRED" not in text:
        raise ValueError("Expected author-action safeguards are missing")

    document = Document()
    configure_section(document)
    configure_styles(document)
    document.core_properties.title = re.sub(r"^#\s+", "", text.splitlines()[0])
    document.core_properties.subject = "Genome Medicine Research manuscript"
    document.core_properties.author = "Manuscript authors"
    document.core_properties.keywords = "SLE; B cells; single-cell RNA sequencing"
    document.core_properties.comments = (
        "Pre-submission review copy; highlighted author actions must be completed."
    )

    current_section = ""
    reference_entries = 0
    action_paragraphs = 0
    for kind, content in markdown_blocks(text):
        if kind == "title":
            paragraph = document.add_paragraph(style="Manuscript Title")
            run = paragraph.add_run(content)
            set_run_font(run, size=16, bold=True)
            set_language(run)
        elif kind == "heading1":
            current_section = content
            paragraph = document.add_paragraph(style="Heading 1")
            add_inline(paragraph, content)
        elif kind == "heading2":
            paragraph = document.add_paragraph(style="Heading 2")
            add_inline(paragraph, content)
        else:
            is_reference = current_section == "References" and bool(
                re.match(r"^\d+\.\s", content)
            )
            style = "Reference Entry" if is_reference else "Normal"
            paragraph = document.add_paragraph(style=style)
            if content.startswith(("**Article type:**", "**Authors:**", "**Affiliations:**", "**Corresponding author:**", "**Running title:**")):
                paragraph.paragraph_format.keep_together = True
            add_inline(paragraph, content)
            if is_reference:
                reference_entries += 1
            if "AUTHOR ACTION REQUIRED" in content:
                action_paragraphs += 1

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)
    metrics = {
        "input": str(input_path),
        "output": str(output_path),
        "paragraphs": len(document.paragraphs),
        "sections": len(document.sections),
        "references": reference_entries,
        "author_action_paragraphs": action_paragraphs,
        "bytes": output_path.stat().st_size,
    }
    qa_path = output_path.with_suffix(".build_qc.json")
    qa_path.write_text(json.dumps(metrics, indent=2), encoding="utf-8")
    return metrics


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--input", type=Path, default=DEFAULT_INPUT)
    parser.add_argument("--output", type=Path, default=DEFAULT_OUTPUT)
    args = parser.parse_args()
    metrics = build(args.input, args.output)
    print(json.dumps(metrics, indent=2))


if __name__ == "__main__":
    main()
