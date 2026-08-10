#!/usr/bin/env python
"""Audit the project workspace file inventory and generated analysis outputs."""

from __future__ import annotations

import argparse
import csv
import hashlib
import json
import os
import re
import subprocess
import sys
import gzip
import tarfile
import traceback
import zipfile
from collections import Counter, defaultdict
from datetime import datetime
from pathlib import Path
from typing import Any


TEXT_EXTS = {
    ".md",
    ".txt",
    ".py",
    ".ps1",
    ".yml",
    ".yaml",
    ".json",
    ".tex",
    ".bib",
    ".sty",
    ".csv",
}
CSV_EXTS = {".csv", ".tsv"}
IMAGE_EXTS = {".png", ".jpg", ".jpeg"}
H5AD_EXTS = {".h5ad"}
DOC_EXTS = {".pdf", ".docx", ".xlsx"}
ARCHIVE_EXTS = {".gz", ".tgz"}
SMALL_HASH_LIMIT = 50 * 1024 * 1024


def relpath(path: Path, root: Path) -> str:
    return str(path.relative_to(root)).replace("\\", "/")


def size_mib(n_bytes: int) -> float:
    return round(n_bytes / (1024**2), 3)


def size_gib(n_bytes: int) -> float:
    return round(n_bytes / (1024**3), 3)


def write_csv(path: Path, rows: list[dict[str, Any]], fieldnames: list[str] | None = None) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    if fieldnames is None:
        keys: list[str] = []
        seen: set[str] = set()
        for row in rows:
            for key in row:
                if key not in seen:
                    seen.add(key)
                    keys.append(key)
        fieldnames = keys
    with path.open("w", encoding="utf-8-sig", newline="") as handle:
        writer = csv.DictWriter(handle, fieldnames=fieldnames, extrasaction="ignore")
        writer.writeheader()
        writer.writerows(rows)


def write_json(path: Path, payload: Any) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")


def sha256_file(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def enumerate_files(root: Path) -> tuple[list[Path], list[Path]]:
    files: list[Path] = []
    dirs: list[Path] = []
    for path in root.rglob("*"):
        if path.is_dir():
            dirs.append(path)
        elif path.is_file():
            files.append(path)
    return sorted(files), sorted(dirs)


def path_is_relative_to(path: Path, parent: Path) -> bool:
    try:
        path.relative_to(parent)
        return True
    except ValueError:
        return False


def make_inventory(root: Path, files: list[Path]) -> list[dict[str, Any]]:
    rows = []
    for path in files:
        stat = path.stat()
        rows.append(
            {
                "relative_path": relpath(path, root),
                "extension": path.suffix.lower() or "[no_ext]",
                "size_bytes": stat.st_size,
                "size_mib": size_mib(stat.st_size),
                "size_gib": size_gib(stat.st_size),
                "mtime": datetime.fromtimestamp(stat.st_mtime).isoformat(timespec="seconds"),
            }
        )
    return rows


def summarize_types(inventory: list[dict[str, Any]]) -> list[dict[str, Any]]:
    by_ext: dict[str, dict[str, Any]] = defaultdict(lambda: {"extension": "", "file_count": 0, "size_bytes": 0})
    for row in inventory:
        ext = str(row["extension"])
        by_ext[ext]["extension"] = ext
        by_ext[ext]["file_count"] += 1
        by_ext[ext]["size_bytes"] += int(row["size_bytes"])
    out = []
    for row in by_ext.values():
        row["size_mib"] = size_mib(row["size_bytes"])
        row["size_gib"] = size_gib(row["size_bytes"])
        out.append(row)
    return sorted(out, key=lambda x: x["size_bytes"], reverse=True)


def summarize_folders(root: Path, files: list[Path]) -> tuple[list[dict[str, Any]], list[dict[str, Any]]]:
    top: dict[str, dict[str, Any]] = defaultdict(lambda: {"folder": "", "file_count": 0, "size_bytes": 0})
    all_dirs: dict[str, dict[str, Any]] = defaultdict(lambda: {"folder": "", "file_count": 0, "size_bytes": 0})
    for path in files:
        r = path.relative_to(root)
        top_name = r.parts[0] if r.parts else "."
        top[top_name]["folder"] = top_name
        top[top_name]["file_count"] += 1
        top[top_name]["size_bytes"] += path.stat().st_size
        parent = path.parent
        while True:
            folder = relpath(parent, root) if parent != root else "."
            all_dirs[folder]["folder"] = folder
            all_dirs[folder]["file_count"] += 1
            all_dirs[folder]["size_bytes"] += path.stat().st_size
            if parent == root:
                break
            parent = parent.parent
    for rows in (top, all_dirs):
        for row in rows.values():
            row["size_mib"] = size_mib(row["size_bytes"])
            row["size_gib"] = size_gib(row["size_bytes"])
    return (
        sorted(top.values(), key=lambda x: x["size_bytes"], reverse=True),
        sorted(all_dirs.values(), key=lambda x: x["size_bytes"], reverse=True),
    )


def audit_text(path: Path, root: Path) -> dict[str, Any]:
    result: dict[str, Any] = {
        "relative_path": relpath(path, root),
        "extension": path.suffix.lower(),
        "status": "ok",
    }
    raw = path.read_bytes()
    if len(raw) <= SMALL_HASH_LIMIT:
        result["sha256"] = sha256_file(path)
    try:
        text = raw.decode("utf-8-sig")
        result["encoding"] = "utf-8-sig"
    except UnicodeDecodeError:
        text = raw.decode("utf-8", errors="replace")
        result["encoding"] = "utf-8-with-replacement"
        result["warnings"] = ["decode replacement was needed"]
    lines = text.splitlines()
    result["line_count"] = len(lines)
    result["nonempty_line_count"] = sum(bool(line.strip()) for line in lines)
    result["char_count"] = len(text)
    if path.suffix.lower() == ".md":
        headings = [line.strip() for line in lines if line.lstrip().startswith("#")]
        result["first_headings"] = headings[:5]
    if path.suffix.lower() == ".py":
        try:
            compile(text, str(path), "exec")
            result["python_compile"] = "ok"
        except SyntaxError as exc:
            result["python_compile"] = "failed"
            result["error"] = str(exc)
    return result


def audit_csv(path: Path, root: Path) -> dict[str, Any]:
    import pandas as pd

    sep = "\t" if path.suffix.lower() == ".tsv" else ","
    result: dict[str, Any] = {"relative_path": relpath(path, root), "status": "ok"}
    try:
        df = pd.read_csv(path, sep=sep, low_memory=False)
        result["n_rows"] = int(df.shape[0])
        result["n_cols"] = int(df.shape[1])
        result["columns"] = list(map(str, df.columns))
        result["missing_total"] = int(df.isna().sum().sum())
        result["empty_columns"] = [str(col) for col in df.columns if int(df[col].isna().sum()) == len(df)]
        result["duplicate_rows"] = int(df.duplicated().sum()) if len(df) else 0
        key_cols = [
            c
            for c in [
                "leiden",
                "draft_state",
                "state_label",
                "cell_type",
                "author_cell_type",
                "ct_cov",
                "disease",
                "disease_state",
                "donor_id",
                "cluster",
            ]
            if c in df.columns
        ]
        result["key_column_missing"] = {str(c): int(df[c].isna().sum()) for c in key_cols}
        result["key_column_nunique"] = {str(c): int(df[c].nunique(dropna=True)) for c in key_cols}
    except Exception as exc:  # noqa: BLE001
        result["status"] = "failed"
        result["error"] = str(exc)
    return result


def audit_image(path: Path, root: Path) -> dict[str, Any]:
    result: dict[str, Any] = {"relative_path": relpath(path, root), "status": "ok"}
    try:
        from PIL import Image

        with Image.open(path) as img:
            result["width"] = img.width
            result["height"] = img.height
            result["mode"] = img.mode
            result["format"] = img.format
    except Exception as exc:  # noqa: BLE001
        result["status"] = "failed"
        result["error"] = str(exc)
    return result


def top_values(series: Any, limit: int = 8) -> dict[str, int]:
    counts = series.astype("object").value_counts(dropna=True).head(limit)
    return {str(k): int(v) for k, v in counts.items()}


def audit_h5ad(path: Path, root: Path) -> dict[str, Any]:
    import anndata as ad

    result: dict[str, Any] = {"relative_path": relpath(path, root), "status": "ok"}
    try:
        adata = ad.read_h5ad(path, backed="r")
        result["shape"] = [int(adata.n_obs), int(adata.n_vars)]
        result["obs_columns_count"] = len(adata.obs.columns)
        result["var_columns_count"] = len(adata.var.columns)
        result["obs_columns"] = list(map(str, adata.obs.columns))
        result["var_columns"] = list(map(str, adata.var.columns))
        result["obsm_keys"] = list(map(str, adata.obsm.keys()))
        result["layers_keys"] = list(map(str, adata.layers.keys()))
        result["uns_keys"] = list(map(str, adata.uns.keys()))[:50]
        result["x_type"] = type(adata.X).__name__
        result["raw_present"] = adata.raw is not None
        selected_obs = [
            c
            for c in [
                "cell_type",
                "author_cell_type",
                "ct_cov",
                "disease",
                "disease_state",
                "donor_id",
                "leiden",
                "draft_state",
                "state_label",
            ]
            if c in adata.obs.columns
        ]
        result["selected_obs_missing"] = {c: int(adata.obs[c].isna().sum()) for c in selected_obs}
        result["selected_obs_nunique"] = {c: int(adata.obs[c].nunique(dropna=True)) for c in selected_obs}
        result["selected_obs_top_values"] = {c: top_values(adata.obs[c]) for c in selected_obs}
        selected_var = [c for c in ["feature_name", "gene_symbol", "ensembl_id"] if c in adata.var.columns]
        result["selected_var_missing"] = {c: int(adata.var[c].isna().sum()) for c in selected_var}
        result["selected_var_examples"] = {
            c: [str(x) for x in adata.var[c].dropna().astype("object").head(10).tolist()] for c in selected_var
        }
        try:
            adata.file.close()
        except Exception:
            pass
    except Exception as exc:  # noqa: BLE001
        result["status"] = "failed"
        result["error"] = str(exc)
        result["traceback"] = traceback.format_exc(limit=3)
    return result


def approximate_pdf_pages(path: Path) -> int | None:
    try:
        data = path.read_bytes()
        matches = re.findall(rb"/Type\s*/Page\b", data)
        return len(matches)
    except Exception:
        return None


def audit_zip_document(path: Path, root: Path) -> dict[str, Any]:
    result: dict[str, Any] = {"relative_path": relpath(path, root), "status": "ok"}
    ext = path.suffix.lower()
    if ext == ".pdf":
        result["page_count_approx"] = approximate_pdf_pages(path)
        return result
    if ext in {".docx", ".xlsx"}:
        try:
            with zipfile.ZipFile(path) as zf:
                names = zf.namelist()
                result["zip_entry_count"] = len(names)
                result["has_content_types"] = "[Content_Types].xml" in names
                result["top_entries"] = names[:20]
                if ext == ".xlsx":
                    result["worksheet_xml_count"] = sum(name.startswith("xl/worksheets/") for name in names)
                if ext == ".docx":
                    result["has_document_xml"] = "word/document.xml" in names
        except Exception as exc:  # noqa: BLE001
            result["status"] = "failed"
            result["error"] = str(exc)
    return result


def audit_archive(path: Path, root: Path) -> dict[str, Any]:
    result: dict[str, Any] = {"relative_path": relpath(path, root), "status": "ok"}
    ext = path.suffix.lower()
    try:
        if ext == ".tgz" or path.name.lower().endswith(".tar.gz"):
            with tarfile.open(path, mode="r:gz") as tf:
                members = tf.getmembers()
                result["member_count"] = len(members)
                result["uncompressed_size_bytes"] = int(sum(m.size for m in members))
                result["top_members"] = [m.name for m in members[:20]]
        elif ext == ".gz":
            total = 0
            with gzip.open(path, "rb") as handle:
                for chunk in iter(lambda: handle.read(1024 * 1024), b""):
                    total += len(chunk)
            result["uncompressed_size_bytes"] = int(total)
    except Exception as exc:  # noqa: BLE001
        result["status"] = "failed"
        result["error"] = str(exc)
    return result


def audit_documents_with_optional_libs(root: Path, files: list[Path]) -> list[dict[str, Any]]:
    rows: list[dict[str, Any]] = []
    for path in files:
        ext = path.suffix.lower()
        row: dict[str, Any] = {"relative_path": relpath(path, root), "status": "ok"}
        try:
            if ext == ".pdf":
                from pypdf import PdfReader

                reader = PdfReader(str(path))
                row["page_count"] = len(reader.pages)
                row["encrypted"] = bool(reader.is_encrypted)
                meta = reader.metadata or {}
                row["title"] = str(meta.get("/Title", "") or "")[:200]
            elif ext == ".docx":
                from docx import Document

                doc = Document(str(path))
                row["paragraph_count"] = len(doc.paragraphs)
                row["table_count"] = len(doc.tables)
                row["nonempty_paragraph_count"] = sum(bool(p.text.strip()) for p in doc.paragraphs)
            elif ext == ".xlsx":
                import openpyxl

                wb = openpyxl.load_workbook(path, read_only=True, data_only=False)
                row["sheet_names"] = wb.sheetnames
                row["sheet_count"] = len(wb.sheetnames)
                row["sheets"] = [
                    {"name": ws.title, "max_row": ws.max_row, "max_column": ws.max_column} for ws in wb.worksheets
                ]
                wb.close()
        except Exception as exc:  # noqa: BLE001
            row["status"] = "failed"
            row["error"] = str(exc)
        rows.append(row)
    return rows


def run_document_sidecar(root: Path, output_dir: Path, bundle_python: Path | None) -> list[dict[str, Any]] | None:
    if not bundle_python or not bundle_python.exists():
        return None
    sidecar = output_dir / "document_audit_exact.json"
    cmd = [
        str(bundle_python),
        str(Path(__file__).resolve()),
        "--root",
        str(root),
        "--outdir",
        str(output_dir),
        "--documents-only",
    ]
    proc = subprocess.run(cmd, check=False, capture_output=True, text=True)
    if proc.returncode != 0:
        return [
            {
                "relative_path": "[document sidecar]",
                "status": "failed",
                "error": proc.stderr[-2000:] or proc.stdout[-2000:],
            }
        ]
    if sidecar.exists():
        return json.loads(sidecar.read_text(encoding="utf-8"))
    return None


def markdown_table(rows: list[dict[str, Any]], columns: list[str], max_rows: int | None = None) -> str:
    visible = rows[:max_rows] if max_rows else rows
    header = "| " + " | ".join(columns) + " |"
    sep = "| " + " | ".join(["---"] * len(columns)) + " |"
    lines = [header, sep]
    for row in visible:
        values = []
        for col in columns:
            value = row.get(col, "")
            if isinstance(value, float):
                value = f"{value:.3f}"
            values.append(str(value).replace("\n", " ")[:180])
        lines.append("| " + " | ".join(values) + " |")
    return "\n".join(lines)


def build_consistency_checks(root: Path, audits: dict[str, Any], inventory: list[dict[str, Any]]) -> list[dict[str, Any]]:
    checks: list[dict[str, Any]] = []
    inv = {row["relative_path"]: row for row in inventory}
    h5 = {row["relative_path"]: row for row in audits.get("h5ad", [])}
    csvs = {row["relative_path"]: row for row in audits.get("csv", [])}

    source = "Data/processed/GSE174188_perez_cellxgene/perez_gse174188_cellxgene.h5ad"
    expected_source_size = 12218105530
    if source in inv:
        checks.append(
            {
                "check": "CELLxGENE source H5AD byte size",
                "status": "pass" if int(inv[source]["size_bytes"]) == expected_source_size else "warning",
                "detail": f'{inv[source]["size_bytes"]} bytes; expected {expected_source_size}',
            }
        )

    expected_full_obs = 152981
    for rel in [
        "Data/processed/GSE174188_perez_cellxgene/bcell_subset_full.h5ad",
        "03_results/first_pass_bcell_full/bcell_first_pass_processed.h5ad",
        "03_results/first_pass_bcell_full/bcell_first_pass_labeled.h5ad",
    ]:
        if rel in h5:
            obs = h5[rel].get("shape", [None])[0]
            checks.append(
                {
                    "check": f"{rel} n_obs",
                    "status": "pass" if obs == expected_full_obs else "warning",
                    "detail": str(h5[rel].get("shape")),
                }
            )

    for rel in [
        "03_results/first_pass_bcell_full/tables/bcell_obs_scores.csv",
        "03_results/first_pass_bcell_full/tables/bcell_obs_scores_labeled.csv",
    ]:
        if rel in csvs:
            rows = csvs[rel].get("n_rows")
            checks.append(
                {
                    "check": f"{rel} rows",
                    "status": "pass" if rows == expected_full_obs else "warning",
                    "detail": f"{rows} rows",
                }
            )

    label_rel = "03_results/first_pass_bcell_full/bcell_first_pass_labeled.h5ad"
    if label_rel in h5:
        obs_missing = h5[label_rel].get("selected_obs_missing", {})
        for col in ["leiden", "draft_state", "state_label", "disease", "disease_state", "donor_id"]:
            if col in obs_missing:
                checks.append(
                    {
                        "check": f"labeled H5AD {col} missingness",
                        "status": "pass" if obs_missing[col] == 0 else "warning",
                        "detail": f"{obs_missing[col]} missing",
                    }
                )

    score_rel = "03_results/first_pass_bcell_full/tables/bcell_obs_scores_labeled.csv"
    if score_rel in csvs:
        missing = csvs[score_rel].get("key_column_missing", {})
        if "ct_cov" in missing:
            checks.append(
                {
                    "check": "ct_cov missingness in labeled score table",
                    "status": "note",
                    "detail": f"{missing['ct_cov']} missing; inherited from source metadata, not a failed run",
                }
            )

    return checks


def generate_report(
    root: Path,
    outdir: Path,
    inventory: list[dict[str, Any]],
    top_summary: list[dict[str, Any]],
    type_summary: list[dict[str, Any]],
    large_files: list[dict[str, Any]],
    empty_dirs: list[str],
    zero_files: list[dict[str, Any]],
    audits: dict[str, Any],
    checks: list[dict[str, Any]],
) -> str:
    total_bytes = sum(int(row["size_bytes"]) for row in inventory)
    lines: list[str] = []
    lines.append("# Workspace Audit Report")
    lines.append("")
    lines.append(f"- Workspace: `{root}`")
    lines.append(f"- Generated: {datetime.now().isoformat(timespec='seconds')}")
    lines.append(f"- Files: {len(inventory)}")
    lines.append(f"- Total size: {total_bytes:,} bytes ({size_gib(total_bytes)} GiB)")
    lines.append(f"- Output folder: `{outdir}`")
    lines.append("")
    lines.append("## Top-Level Folder Summary")
    lines.append(markdown_table(top_summary, ["folder", "file_count", "size_gib", "size_mib"], max_rows=20))
    lines.append("")
    lines.append("## File Type Summary")
    lines.append(markdown_table(type_summary, ["extension", "file_count", "size_gib", "size_mib"], max_rows=30))
    lines.append("")
    lines.append("## Largest Files")
    lines.append(markdown_table(large_files, ["relative_path", "size_gib", "size_mib", "mtime"], max_rows=30))
    lines.append("")
    lines.append("## Consistency Checks")
    lines.append(markdown_table(checks, ["status", "check", "detail"], max_rows=None))
    lines.append("")
    lines.append("## H5AD Files")
    h5_rows = []
    for row in audits.get("h5ad", []):
        h5_rows.append(
            {
                "relative_path": row["relative_path"],
                "status": row.get("status"),
                "shape": row.get("shape"),
                "obs_cols": row.get("obs_columns_count"),
                "var_cols": row.get("var_columns_count"),
                "obsm": ",".join(row.get("obsm_keys", []))[:160],
                "raw": row.get("raw_present"),
            }
        )
    lines.append(markdown_table(h5_rows, ["relative_path", "status", "shape", "obs_cols", "var_cols", "obsm", "raw"]))
    lines.append("")
    lines.append("## CSV Tables")
    csv_rows = [
        {
            "relative_path": row["relative_path"],
            "status": row.get("status"),
            "n_rows": row.get("n_rows"),
            "n_cols": row.get("n_cols"),
            "missing_total": row.get("missing_total"),
        }
        for row in audits.get("csv", [])
    ]
    lines.append(markdown_table(csv_rows, ["relative_path", "status", "n_rows", "n_cols", "missing_total"], max_rows=60))
    lines.append("")
    lines.append("## Figures And Images")
    img_rows = [
        {
            "relative_path": row["relative_path"],
            "status": row.get("status"),
            "width": row.get("width"),
            "height": row.get("height"),
            "format": row.get("format"),
        }
        for row in audits.get("images", [])
    ]
    lines.append(markdown_table(img_rows, ["relative_path", "status", "width", "height", "format"], max_rows=40))
    lines.append("")
    lines.append("## Documents")
    doc_rows = audits.get("documents_exact") or audits.get("documents_basic", [])
    lines.append(markdown_table(doc_rows, ["relative_path", "status", "page_count", "paragraph_count", "sheet_count", "title"], max_rows=40))
    lines.append("")
    lines.append("## Compressed Archives")
    archive_rows = [
        {
            "relative_path": row["relative_path"],
            "status": row.get("status"),
            "member_count": row.get("member_count", ""),
            "uncompressed_size_bytes": row.get("uncompressed_size_bytes", ""),
        }
        for row in audits.get("archives", [])
    ]
    if archive_rows:
        lines.append(markdown_table(archive_rows, ["relative_path", "status", "member_count", "uncompressed_size_bytes"], max_rows=40))
    else:
        lines.append("- None")
    lines.append("")
    lines.append("## Empty Directories")
    if empty_dirs:
        lines.extend([f"- `{d}`" for d in empty_dirs])
    else:
        lines.append("- None")
    lines.append("")
    lines.append("## Zero-Byte Files")
    if zero_files:
        lines.append(markdown_table(zero_files, ["relative_path", "mtime"], max_rows=40))
    else:
        lines.append("- None")
    lines.append("")
    lines.append("## Notes")
    lines.append("- Large H5AD files are expected for this phase; keep them until manuscript figures are frozen.")
    lines.append("- `ct_cov` missingness is source metadata missingness, not evidence that the local run stopped early.")
    lines.append("- Python bytecode under `02_analysis/scripts/__pycache__` is disposable cache and can be removed later.")
    return "\n".join(lines) + "\n"


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--root", type=Path, default=Path.cwd())
    parser.add_argument("--outdir", type=Path, default=None)
    parser.add_argument("--bundle-python", type=Path, default=None)
    parser.add_argument("--documents-only", action="store_true")
    args = parser.parse_args()

    root = args.root.resolve()
    outdir = args.outdir.resolve() if args.outdir else root / "00_project_management" / "audit_2026-06-23"
    outdir.mkdir(parents=True, exist_ok=True)

    files, dirs = enumerate_files(root)

    if args.documents_only:
        doc_files = [p for p in files if p.suffix.lower() in DOC_EXTS]
        doc_rows = audit_documents_with_optional_libs(root, doc_files)
        write_json(outdir / "document_audit_exact.json", doc_rows)
        write_csv(outdir / "document_audit_exact.csv", doc_rows)
        return

    inventory = make_inventory(root, files)
    type_summary = summarize_types(inventory)
    top_summary, folder_summary = summarize_folders(root, files)
    large_files = sorted(inventory, key=lambda x: int(x["size_bytes"]), reverse=True)[:50]
    zero_files = [row for row in inventory if int(row["size_bytes"]) == 0]
    empty_dirs = [relpath(path, root) for path in dirs if not any(path_is_relative_to(f, path) for f in files)]

    audits: dict[str, Any] = {"text": [], "csv": [], "images": [], "h5ad": [], "documents_basic": [], "archives": []}
    for path in files:
        ext = path.suffix.lower()
        if ext in H5AD_EXTS:
            audits["h5ad"].append(audit_h5ad(path, root))
        elif ext in CSV_EXTS:
            audits["csv"].append(audit_csv(path, root))
        elif ext in IMAGE_EXTS:
            audits["images"].append(audit_image(path, root))
        elif ext in DOC_EXTS:
            audits["documents_basic"].append(audit_zip_document(path, root))
        elif ext in ARCHIVE_EXTS:
            audits["archives"].append(audit_archive(path, root))
        elif ext in TEXT_EXTS or (not ext and path.stat().st_size <= SMALL_HASH_LIMIT):
            audits["text"].append(audit_text(path, root))

    documents_exact = run_document_sidecar(root, outdir, args.bundle_python)
    if documents_exact is not None:
        audits["documents_exact"] = documents_exact

    checks = build_consistency_checks(root, audits, inventory)

    write_csv(outdir / "file_inventory.csv", inventory)
    write_csv(outdir / "type_summary.csv", type_summary)
    write_csv(outdir / "folder_summary_top_level.csv", top_summary)
    write_csv(outdir / "folder_summary_all_dirs.csv", folder_summary)
    write_csv(outdir / "large_files_top50.csv", large_files)
    write_csv(outdir / "zero_byte_files.csv", zero_files)
    write_json(outdir / "content_audit.json", audits)
    write_json(outdir / "consistency_checks.json", checks)
    report = generate_report(
        root=root,
        outdir=outdir,
        inventory=inventory,
        top_summary=top_summary,
        type_summary=type_summary,
        large_files=large_files,
        empty_dirs=empty_dirs,
        zero_files=zero_files,
        audits=audits,
        checks=checks,
    )
    (outdir / "workspace_audit_report.md").write_text(report, encoding="utf-8-sig")
    print(f"Audit complete: {outdir}")
    print(f"Files audited: {len(inventory)}")
    print(f"Total size GiB: {size_gib(sum(int(row['size_bytes']) for row in inventory))}")


if __name__ == "__main__":
    main()
