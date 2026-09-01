#!/usr/bin/env python3
"""Build manuscript and pagination candidates for the citation-order refreeze."""

from __future__ import annotations

import hashlib
import json
import re
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn

import phase17_c8s_04_build_documents as documents


ROOT = Path(__file__).resolve().parents[1]
RUN = ROOT / "phase17_v7/npj_sba_supplementary_citation_refreeze/20260901_first_citation_order"
SOURCES = RUN / "sources"
FIGURES = RUN / "figures/figures"
DOCUMENTS = RUN / "documents"
CANDIDATES = RUN / "qa/pagination_candidates"
TITLE = (
    "Disease-blind reconstruction distinguishes reproducible interferon remodeling from less stable "
    "B-cell state assignments in systemic lupus erythematosus"
)


def sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest().upper()


def extract_text(document: Document) -> str:
    values = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            values.extend(cell.text for cell in row.cells)
    return "\n".join(values)


def remove_table_spacers_before_manual_page_breaks(document: Document) -> int:
    removed = 0
    for paragraph in list(document.paragraphs):
        if not paragraph._p.xpath('.//w:br[@w:type="page"]'):
            continue
        previous = paragraph._p.getprevious()
        if previous is None or previous.tag != qn("w:p"):
            continue
        text = "".join(previous.itertext()).strip()
        has_break = bool(previous.xpath('.//w:br[@w:type="page"]'))
        has_drawing = bool(previous.xpath(".//w:drawing"))
        if not text and not has_break and not has_drawing:
            previous.getparent().remove(previous)
            removed += 1
    return removed


def remove_manual_page_break_before_heading(document: Document, heading: str) -> int:
    removed = 0
    for paragraph in document.paragraphs:
        if paragraph.text.strip() != heading:
            continue
        previous = paragraph._p.getprevious()
        if previous is None or previous.tag != qn("w:p"):
            continue
        if not previous.xpath('.//w:br[@w:type="page"]'):
            continue
        previous.getparent().remove(previous)
        removed += 1
    return removed


def patch_properties(path: Path, subject: str, comments: str) -> dict[str, int]:
    document = Document(path)
    removed_spacers = remove_table_spacers_before_manual_page_breaks(document)
    removed_s8_break = 0
    compacted_main_legend_headings = 0
    if path.name.startswith("Supplementary"):
        removed_s8_break = remove_manual_page_break_before_heading(
            document, "Supplementary Table S8 | Full statistical-results archive map"
        )
    else:
        for paragraph in document.paragraphs:
            if re.fullmatch(r"Figure [1-5] \| .+", paragraph.text.strip()):
                paragraph.paragraph_format.space_before = Pt(6)
                paragraph.paragraph_format.space_after = Pt(2)
                compacted_main_legend_headings += 1
    document.core_properties.author = "Zhi Chen; Teng Qi"
    document.core_properties.title = TITLE if path.name.startswith("Manuscript") else "Supplementary information"
    document.core_properties.subject = subject
    document.core_properties.comments = comments
    document.save(path)
    return {
        "redundant_spacers_removed": removed_spacers,
        "renderer_divergent_s8_breaks_removed": removed_s8_break,
        "main_legend_headings_compacted": compacted_main_legend_headings,
    }


def figure_alt_titles(document: Document) -> list[str]:
    values = []
    for shape in document.inline_shapes:
        values.append(shape._inline.docPr.get("title"))
    return values


def main() -> None:
    DOCUMENTS.mkdir(parents=True, exist_ok=True)
    CANDIDATES.mkdir(parents=True, exist_ok=True)
    manuscript_source = SOURCES / "Manuscript_first_citation_order_refreeze.md"
    supplement_source = SOURCES / "Supplementary_Information_first_citation_order_refreeze.md"
    manuscript = DOCUMENTS / "Manuscript_scientific_maintenance_freeze.docx"
    standard = CANDIDATES / "Supplementary_Information_standard_candidate.docx"
    compact = CANDIDATES / "Supplementary_Information_compact_candidate.docx"

    results = {
        "manuscript": documents.markdown_to_docx(
            manuscript_source,
            manuscript,
            body_size=12,
            double_space=True,
            line_numbers=True,
            running_header="npj Systems Biology and Applications | Article",
            title_override=TITLE,
        ),
        "standard_supplement": documents.markdown_to_docx(
            supplement_source,
            standard,
            body_size=10.5,
            double_space=False,
            line_numbers=False,
            running_header="Supplementary information",
            title_override="Supplementary information",
            supplementary_figure_dirs=[FIGURES],
            page_break_before_headings=set(),
        ),
        "compact_supplement": documents.markdown_to_docx(
            supplement_source,
            compact,
            body_size=10.5,
            double_space=False,
            line_numbers=False,
            running_header="Supplementary information",
            title_override="Supplementary information",
            supplementary_figure_dirs=[FIGURES],
            page_break_before_headings=set(),
        ),
    }

    repairs = {
        "manuscript": patch_properties(
            manuscript,
            "Supplementary first-citation-order scientific refreeze",
            "Reader-path and Supplementary Table anchor repair; no scientific-value change.",
        ),
        "standard_supplement": patch_properties(
            standard,
            "Standard pagination candidate",
            "Display-only Supplementary Figure renumbering from byte-identical scientific objects.",
        ),
        "compact_supplement": patch_properties(
            compact,
            "Compact pagination candidate",
            "Display-only Supplementary Figure renumbering; candidate tests removal of the S1 page break.",
        ),
    }

    compact_document = Document(compact)
    removed_s1_break = remove_manual_page_break_before_heading(
        compact_document, "Supplementary Figure S1 | Source integrity and hard-quality-control diagnostics"
    )
    compact_document.save(compact)
    repairs["compact_supplement"]["s1_breaks_removed"] = removed_s1_break
    repairs["standard_supplement"]["s1_breaks_removed"] = 0

    main_document = Document(manuscript)
    standard_document = Document(standard)
    compact_document = Document(compact)
    main_text = extract_text(main_document)
    standard_text = extract_text(standard_document)
    source_text = manuscript_source.read_text(encoding="utf-8")
    body = source_text.split("## Figure legends", 1)[0]
    first_order = []
    for value in re.findall(r"Supplementary Fig(?:ure)?\. S(10|[1-9])", body):
        number = int(value)
        if number not in first_order:
            first_order.append(number)
    reference_numbers = [
        int(value)
        for value in re.findall(
            r"(?m)^(\d+)\. ",
            source_text.split("## References\n", 1)[1].split("## Figure legends\n", 1)[0],
        )
    ]
    abstract_match = re.search(r"Abstract\n(.+?)\nIntroduction", main_text, flags=re.DOTALL)
    if not abstract_match:
        raise RuntimeError("Could not identify manuscript abstract")
    abstract_words = re.findall(r"\b[\w'-]+\b", abstract_match.group(1))
    expected_alt_titles = [f"Supplementary Figure S{number}" for number in range(1, 11)]

    checks = {
        "manuscript_title_exact": TITLE in main_text,
        "abstract_145_words": len(abstract_words) == 145,
        "references_1_to_33": reference_numbers == list(range(1, 34)),
        "main_has_no_inline_figures": len(main_document.inline_shapes) == 0,
        "first_citation_order_s1_to_s10": first_order == list(range(1, 11)),
        "four_functional_table_anchors_present": all(
            phrase in main_text
            for phrase in (
                "Supplementary Tables S1 and S2",
                "Supplementary Table S3",
                "Supplementary Table S4",
                "Supplementary Tables S5-S8",
            )
        ),
        "standard_has_ten_figures": len(standard_document.inline_shapes) == 10,
        "compact_has_ten_figures": len(compact_document.inline_shapes) == 10,
        "standard_has_ten_grid_objects": len(standard_document.tables) == 10,
        "compact_has_ten_grid_objects": len(compact_document.tables) == 10,
        "standard_alt_titles_s1_to_s10": figure_alt_titles(standard_document) == expected_alt_titles,
        "compact_alt_titles_s1_to_s10": figure_alt_titles(compact_document) == expected_alt_titles,
        "supplement_heading_sequence_s1_to_s10": [
            int(value)
            for value in re.findall(r"Supplementary Figure S(10|[1-9]) \|", standard_text)
        ] == list(range(1, 11)),
        "standard_s8_table_break_repaired": repairs["standard_supplement"]["renderer_divergent_s8_breaks_removed"] == 1,
        "compact_s8_table_break_repaired": repairs["compact_supplement"]["renderer_divergent_s8_breaks_removed"] == 1,
        "compact_s1_break_removed_once": removed_s1_break == 1,
        "five_main_legend_headings_compacted": repairs["manuscript"]["main_legend_headings_compacted"] == 5,
    }
    failed = [name for name, passed in checks.items() if not passed]
    status = {
        "created_at": datetime.now().astimezone().isoformat(timespec="seconds"),
        "status": (
            "PASS_SUPPLEMENTARY_CITATION_DOCX_CANDIDATES_BUILT_RENDER_REQUIRED"
            if not failed
            else "FAIL_SUPPLEMENTARY_CITATION_DOCX_BUILD"
        ),
        "build_results": results,
        "layout_repairs": repairs,
        "checks": checks,
        "failed_checks": failed,
        "files": {
            path.name: {
                "path": path.relative_to(ROOT).as_posix(),
                "bytes": path.stat().st_size,
                "sha256": sha256(path),
            }
            for path in (manuscript, standard, compact)
        },
        "scientific_estimates_changed": False,
        "figures_redrawn": False,
        "source_data_values_changed": False,
    }
    (RUN / "05_DOCUMENT_CANDIDATE_BUILD_STATUS.json").write_text(
        json.dumps(status, indent=2) + "\n", encoding="utf-8", newline="\n"
    )
    print(json.dumps(status, indent=2))
    if failed:
        raise RuntimeError(f"Document candidate checks failed: {failed}")


if __name__ == "__main__":
    main()
