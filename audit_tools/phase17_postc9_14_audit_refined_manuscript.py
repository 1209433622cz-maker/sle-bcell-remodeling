"""Extract a received manuscript and compare it without changing frozen sources."""

import argparse
from collections import Counter
from datetime import datetime
import difflib
import hashlib
import json
from pathlib import Path
import re
import unicodedata
import zipfile
from urllib.parse import parse_qs, urlparse

from docx import Document
import pdfplumber

from phase17_postc9_13_audit_target_preparation import word_count


ROOT = Path(__file__).resolve().parents[1]
W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"


def file_record(path):
    data = path.read_bytes()
    return {"name": path.name, "bytes": len(data), "sha256": hashlib.sha256(data).hexdigest().upper()}


def normalized(text):
    return " ".join(unicodedata.normalize("NFKC", text).replace("`", "").replace("*", "").split())


def section(text, name):
    match = re.search(r"(?m)^## " + re.escape(name) + r"\s*\n", text)
    if not match:
        raise ValueError("Missing section: " + name)
    return re.split(r"(?m)^## ", text[match.end():], maxsplit=1)[0].strip()


def reference_identity(text):
    match = re.search(r"\bdoi:\s*(\S+)", text, re.I)
    if match:
        return "doi:" + match[1].rstrip(".").lower()
    for url in re.findall(r"https?://\S+", text):
        parsed = urlparse(url.rstrip("."))
        accession = parse_qs(parsed.query).get("acc", [""])[0]
        if parsed.hostname == "www.ncbi.nlm.nih.gov" and re.fullmatch(r"GSE\d+", accession):
            return "geo:" + accession
    raise ValueError("Reference lacks a recognized stable identity: " + text[:80])


def reference_map(before, after):
    old = [reference_identity(text) for text in before]
    new = [reference_identity(text) for text in after]
    if len(old) != len(set(old)) or len(new) != len(set(new)):
        raise ValueError("Duplicate reference identity")
    if set(old) != set(new):
        raise ValueError("Reference identity set changed")
    return [{"old_number": i + 1, "new_number": new.index(doi) + 1, "identity": doi,
             "bibliography_text_identical": normalized(before[i]) == normalized(after[new.index(doi)])}
            for i, doi in enumerate(old)]


def citation_order(text):
    order = []
    for match in re.finditer(r"\[([\d,\s\-\u2013]+)\]", text):
        for part in match[1].split(","):
            bounds = re.split(r"[-\u2013]", part.strip())
            if len(bounds) == 1:
                numbers = [int(bounds[0])]
            elif len(bounds) == 2 and int(bounds[0]) <= int(bounds[1]):
                numbers = range(int(bounds[0]), int(bounds[1]) + 1)
            else:
                raise ValueError("Malformed citation range")
            for number in numbers:
                if number not in order:
                    order.append(number)
    return order


def numeric_tokens(text):
    text = re.sub(r"\[[\d,\s\-\u2013]+\]", "", text)
    return Counter(re.findall(r"(?<![A-Za-z0-9_])[-+]?(?:\d{1,3}(?:,\d{3})+|\d+)(?:\.\d+)?(?:[eE][-+]?\d+)?", text))


def extract_docx(path):
    document = Document(path)
    if document.tables:
        raise ValueError("Table-aware extraction is required for this manuscript")
    rows, md, refs = [], [], []
    in_references = False
    for index, paragraph in enumerate(document.paragraphs):
        text = paragraph.text
        style = paragraph.style.name
        rows.append({"paragraph": index, "style": style, "text": text})
        if style.startswith("Heading "):
            md.append("#" * int(style.split()[-1]) + " " + text)
            in_references = text == "References"
        elif in_references and text.strip():
            refs.append(text)
            md.append(f"{len(refs)}. {text}")
        elif text.strip():
            # Preserve structured labels in the extracted review source.
            if style in {"First Paragraph", "Body Text"} and re.match(r"^(Background|Methods|Results|Conclusions): ", text):
                label, body = text.split(": ", 1)
                md.append("**" + label + ":** " + body)
            else:
                md.append(text)
    with zipfile.ZipFile(path) as archive:
        from xml.etree import ElementTree as ET
        xml = ET.fromstring(archive.read("word/document.xml"))
        structural = {
            "tracked_insertions": len(xml.findall(".//" + W + "ins")),
            "tracked_deletions": len(xml.findall(".//" + W + "del")),
            "comments_part": "word/comments.xml" in archive.namelist(),
            "embedded_media": [name for name in archive.namelist() if name.startswith("word/media/")],
            "reference_paragraphs_with_numbering": sum(
                p._p.find(W + "pPr/" + W + "numPr") is not None
                for p in document.paragraphs[-len(refs):]
            ),
        }
    return "\n\n".join(md) + "\n", rows, refs, structural


def pdf_body(path):
    pages, bounds = [], []
    with pdfplumber.open(path) as document:
        for number, page in enumerate(document.pages, 1):
            # Geometry observed in the supplied Letter manuscript. Do not remove
            # arbitrary text: margin line numbers and header/footer are excluded.
            pages.append(page.crop((65, 50, page.width, 742)).extract_text() or "")
            outside = [char["text"] for char in page.chars if char["x0"] < 0 or char["x1"] > page.width + .5
                       or char["top"] < 0 or char["bottom"] > page.height + .5]
            margin_numbers = [word for word in page.extract_words()
                              if re.fullmatch(r"\d+", word["text"]) and 25 <= word["x0"] and word["x1"] < 65]
            bounds.append({"page": number, "width": page.width, "height": page.height,
                           "outside_page_characters": outside, "body_characters": len(pages[-1]),
                           "body_margin_line_numbers": [int(w["text"]) for w in margin_numbers if 50 <= w["top"] < 742],
                           "header_footer_margin_numbers": [w["text"] for w in margin_numbers if w["top"] < 50 or w["top"] >= 742]})
    return pages, bounds


def compact_text(text):
    return re.sub(r"\s+", "", unicodedata.normalize("NFKC", text).replace("\u00ad", ""))


def compare_text(before, after):
    a, b = compact_text(before), compact_text(after)
    if a == b:
        return {"normalized_equal": True, "ratio": 1.0, "changes": []}
    matcher = difflib.SequenceMatcher(None, a, b, autojunk=False)
    changes = [{"kind": tag, "expected": a[i:j], "observed": b[k:l],
                "left_context": a[max(0, i - 55):i], "right_context": a[j:j + 55]}
               for tag, i, j, k, l in matcher.get_opcodes() if tag != "equal"]
    return {"normalized_equal": a == b, "ratio": matcher.ratio(), "changes": changes}


def main():
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--docx", required=True, type=Path)
    parser.add_argument("--pdf", required=True, type=Path)
    parser.add_argument("--output-dir", required=True, type=Path)
    parser.add_argument("--no-extracts", action="store_true", help="Write only the JSON audit for a rebuilt review copy")
    args = parser.parse_args()
    output = args.output_dir.resolve()
    if not output.is_relative_to(ROOT / "00_project_management"):
        raise ValueError("Audit outputs must remain in project management")
    output.mkdir(parents=True, exist_ok=True)
    baseline = (ROOT / "01_manuscript/Manuscript.md").read_text(encoding="utf-8-sig")
    refined, paragraphs, refs, structural = extract_docx(args.docx)
    baseline_refs = re.findall(r"(?m)^\d+\. (.+)$", section(baseline, "References"))
    mapping = reference_map(baseline_refs, refs)
    body = refined.split("## Background\n", 1)[1].split("## References\n", 1)[0]
    order = citation_order(body)
    pages, bounds = pdf_body(args.pdf)
    expected_pdf = []
    ref_number = 0
    in_refs = False
    for row in paragraphs:
        if row["text"] == "References" and row["style"].startswith("Heading "):
            in_refs = True
            expected_pdf.append(row["text"])
        elif in_refs and row["text"].strip():
            ref_number += 1
            expected_pdf.append(f"{ref_number}. " + row["text"])
        else:
            expected_pdf.append(row["text"])
    differences = []
    for name in ("Abstract", "Background", "Methods", "Results", "Discussion", "Conclusions", "Figure legends"):
        old, new = section(baseline, name), section(refined, name)
        a, b = numeric_tokens(old), numeric_tokens(new)
        differences.append({"section": name, "old_whitespace_words": word_count(old),
                            "new_whitespace_words": word_count(new),
                            "number_tokens_added": dict(b - a), "number_tokens_removed": dict(a - b),
                            "normalized_section_identical": normalized(old) == normalized(new)})
    record = {
        "created_at": datetime.now().astimezone().isoformat(timespec="seconds"),
        "inputs": [file_record(args.docx), file_record(args.pdf), file_record(ROOT / "01_manuscript/Manuscript.md")],
        "paragraph_count": len(paragraphs), "docx_structure": structural,
        "reference_count": len(refs), "reference_mapping": mapping,
        "first_appearance": order, "first_appearance_contiguous": order == list(range(1, len(refs) + 1)),
        "reference_identity_set_unchanged": True,
        "bibliography_text_unchanged_after_renumbering": all(row["bibliography_text_identical"] for row in mapping),
        "section_comparison": differences,
        "pdf_pages": len(pages), "pdf_geometry_checks": bounds,
        "docx_pdf_body_comparison": compare_text("\n".join(expected_pdf), "\n".join(pages)),
        "pdf_comparison_method": "PDF body crop x=65..page width, y=50..742 pt; NFKC, whitespace and soft-hyphen normalization only. This is not a visual review.",
        "scope": "Extracted review source, structural/identity/citation and body-text checks. Numeric token differences need semantic adjudication. No full source-paper verification or scientific reanalysis.",
        "author_approval": "PENDING_EXACT_REFINED_FILES",
        "canonical_sources_modified": False,
    }
    if not args.no_extracts:
        (output / "Received_Manuscript.md").write_text(refined, encoding="utf-8")
        (output / "Received_Paragraphs.md").write_text("\n\n".join(
            f"## P{row['paragraph']:03d} | {row['style']}\n\n{row['text']}" for row in paragraphs) + "\n", encoding="utf-8")
        (output / "Received_PDF_Body.txt").write_text("\n\n".join(
            f"PAGE {i + 1}\n{text}" for i, text in enumerate(pages)) + "\n", encoding="utf-8")
    (output / "text_integrity_audit.json").write_text(json.dumps(record, indent=2) + "\n", encoding="utf-8")
    print(json.dumps({key: record[key] for key in ("paragraph_count", "reference_count", "first_appearance_contiguous",
          "bibliography_text_unchanged_after_renumbering", "pdf_pages", "section_comparison", "docx_pdf_body_comparison")}, indent=2))


if __name__ == "__main__":
    main()
