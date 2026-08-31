#!/usr/bin/env python3
"""Build paired documents for the reference-terminology and S6 refreeze."""

from __future__ import annotations

import hashlib
import json
import re
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

import phase17_c8s_04_build_documents as documents


ROOT = Path(__file__).resolve().parents[1]
RUN = (
    ROOT
    / "phase17_v7/npj_sba_reference_terminology_lock/"
    "20260901_reference_terminology_s6_refreeze"
)
SOURCES = RUN / "sources"
FIGURES = RUN / "figures/figures"
OUTPUT = RUN / "documents"
TITLE = (
    "Disease-blind reconstruction distinguishes reproducible interferon remodeling from less stable "
    "B-cell state assignments in systemic lupus erythematosus"
)


def sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest().upper()


def extract_text(document: Document) -> str:
    values = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            values.extend(cell.text for cell in row.cells)
    return "\n".join(values)


def remove_table_spacers_before_manual_page_breaks(document: Document) -> int:
    removed = 0
    for paragraph in list(document.paragraphs):
        if not paragraph._p.xpath('.//w:br[@w:type="page"]'):
            continue
        previous = paragraph._p.getprevious()
        if previous is None or previous.tag != qn("w:p"):
            continue
        text = "".join(previous.itertext()).strip()
        has_break = bool(previous.xpath('.//w:br[@w:type="page"]'))
        has_drawing = bool(previous.xpath(".//w:drawing"))
        if not text and not has_break and not has_drawing:
            previous.getparent().remove(previous)
            removed += 1
    return removed


def patch_properties(path: Path, subject: str) -> None:
    document = Document(path)
    removed = remove_table_spacers_before_manual_page_breaks(document)
    document.core_properties.author = "Zhi Chen; Teng Qi"
    document.core_properties.title = TITLE if path.name.startswith("Manuscript") else "Supplementary information"
    document.core_properties.subject = subject
    document.core_properties.comments = (
        "Reference, evidence-terminology and Supplementary Figure S6 scientific refreeze; "
        f"removed {removed} redundant table spacers before manual page breaks."
    )
    document.save(path)


def main() -> None:
    OUTPUT.mkdir(parents=True, exist_ok=True)
    manuscript_source = SOURCES / "Manuscript_reference_terminology_s6_refreeze.md"
    supplement_source = SOURCES / "Supplementary_Information_reference_terminology_s6_refreeze.md"
    manuscript = OUTPUT / "Manuscript_reference_terminology_s6_refreeze.docx"
    supplement = OUTPUT / "Supplementary_Information_reference_terminology_s6_refreeze.docx"

    results = [
        documents.markdown_to_docx(
            manuscript_source,
            manuscript,
            body_size=12,
            double_space=True,
            line_numbers=True,
            running_header="npj Systems Biology and Applications | Article",
            title_override=TITLE,
        ),
        documents.markdown_to_docx(
            supplement_source,
            supplement,
            body_size=10.5,
            double_space=False,
            line_numbers=False,
            running_header="Supplementary information",
            title_override="Supplementary information",
            supplementary_figure_dirs=[FIGURES],
            page_break_before_headings=set(),
        ),
    ]
    patch_properties(manuscript, "Reference and evidence-terminology scientific refreeze")
    patch_properties(supplement, "Supplementary Figure S6 terminology scientific refreeze")

    main_document = Document(manuscript)
    supplementary_document = Document(supplement)
    main_text = extract_text(main_document)
    supplementary_text = extract_text(supplementary_document)
    abstract_match = re.search(r"Abstract\n(.+?)\nIntroduction", main_text, flags=re.DOTALL)
    if not abstract_match:
        raise RuntimeError("Could not identify manuscript abstract")
    abstract_words = re.findall(r"\b[\w'-]+\b", abstract_match.group(1))
    source_text = manuscript_source.read_text(encoding="utf-8")
    reference_numbers = [
        int(value)
        for value in re.findall(
            r"(?m)^(\d+)\. ", source_text.split("## References\n", 1)[1].split("## Figure legends\n", 1)[0]
        )
    ]

    checks = {
        "manuscript": {
            "title_exact": TITLE in main_text,
            "abstract_145_words": len(abstract_words) == 145,
            "no_inline_figures": len(main_document.inline_shapes) == 0,
            "references_1_to_33": reference_numbers == list(range(1, 34)),
            "biological_unit_aware_abstract": "biological-unit-aware inference" in main_text,
            "edgeR_workflow_reference": "From reads to genes to pathways" in main_text,
            "regulatory_context_boundary": "support an IFN-centred regulatory context" in main_text,
        },
        "supplement": {
            "ten_figures_embedded": len(supplementary_document.inline_shapes) == 10,
            "title_exact": TITLE in supplementary_text,
            "s6_replication_title_present": "GSE135779 replication and robustness diagnostics" in supplementary_text,
            "s6_old_title_absent": "Independent-validation diagnostics" not in supplementary_text,
            "internal_replication_present": "Discovery and internal replication" in supplementary_text,
            "independent_replication_present": "Independent SLE replication" in supplementary_text,
        },
    }
    failed = [
        f"{group}.{name}"
        for group, group_checks in checks.items()
        for name, passed in group_checks.items()
        if not passed
    ]
    if failed:
        raise RuntimeError(f"Reference-terminology document checks failed: {failed}")

    status = {
        "created_at": datetime.now().astimezone().isoformat(timespec="seconds"),
        "status": "PASS_REFERENCE_TERMINOLOGY_S6_DOCX_BUILT_DUAL_RENDER_REQUIRED",
        "build_results": results,
        "checks": checks,
        "object_inventory": {
            "manuscript_inline_shapes": len(main_document.inline_shapes),
            "supplement_inline_shapes": len(supplementary_document.inline_shapes),
        },
        "files": {
            path.name: {"bytes": path.stat().st_size, "sha256": sha256(path)}
            for path in (manuscript, supplement)
        },
        "scientific_estimates_changed": False,
        "source_data_changed": False,
        "supplementary_figure_s6_source_redrawn": True,
    }
    (RUN / "01_DOCUMENT_BUILD_STATUS.json").write_text(
        json.dumps(status, indent=2) + "\n", encoding="utf-8", newline="\n"
    )
    print(json.dumps(status, indent=2))


if __name__ == "__main__":
    main()
