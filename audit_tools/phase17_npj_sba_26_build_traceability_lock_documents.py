#!/usr/bin/env python3
"""Build paired documents for the final scientific traceability lock."""

from __future__ import annotations

import hashlib
import json
import re
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

import phase17_c8s_04_build_documents as documents


ROOT = Path(__file__).resolve().parents[1]
RUN = (
    ROOT
    / "phase17_v7/npj_sba_traceability_lock/"
    "20260831_final_scientific_object_lock"
)
SOURCES = RUN / "sources"
FIGURES = (
    ROOT
    / "phase17_v7/npj_sba_scientific_presentation_freeze/"
    "20260831_reader_path_and_legend_economy/figures/figures"
)
OUTPUT = RUN / "documents"
TITLE = (
    "Disease-blind reconstruction distinguishes reproducible interferon remodeling from less stable "
    "B-cell state assignments in systemic lupus erythematosus"
)
UNIT = "Sample-cohort pseudobulk (GSE174188); donor pseudobulk (GSE135779)"
ARCHIVE_SCOPE = (
    "version-specific archive of the released analysis code, Source Data and statistical outputs"
)


def sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest().upper()


def extract_text(document: Document) -> str:
    values = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            values.extend(cell.text for cell in row.cells)
    return "\n".join(values)


def remove_table_spacers_before_manual_page_breaks(document: Document) -> int:
    """Avoid a cross-render blank page when a table spacer precedes a hard break."""
    removed = 0
    for paragraph in list(document.paragraphs):
        has_page_break = bool(paragraph._p.xpath('.//w:br[@w:type="page"]'))
        if not has_page_break:
            continue
        previous = paragraph._p.getprevious()
        if previous is None or previous.tag != qn("w:p"):
            continue
        text = "".join(previous.itertext()).strip()
        has_break = bool(previous.xpath('.//w:br[@w:type="page"]'))
        has_drawing = bool(previous.xpath(".//w:drawing"))
        if not text and not has_break and not has_drawing:
            previous.getparent().remove(previous)
            removed += 1
    return removed


def patch_properties(path: Path, subject: str) -> None:
    document = Document(path)
    removed_spacers = remove_table_spacers_before_manual_page_breaks(document)
    document.core_properties.author = "Zhi Chen; Teng Qi"
    document.core_properties.title = TITLE if path.name.startswith("Manuscript") else "Supplementary information"
    document.core_properties.subject = subject
    document.core_properties.comments = (
        "Final scientific-object and numerical-traceability lock rebuilt from repository sources; "
        f"removed {removed_spacers} redundant table spacers before manual page breaks."
    )
    document.save(path)


def main() -> None:
    OUTPUT.mkdir(parents=True, exist_ok=True)
    manuscript = OUTPUT / "Manuscript_final_scientific_lock.docx"
    supplement = OUTPUT / "Supplementary_Information_final_scientific_lock.docx"
    results = [
        documents.markdown_to_docx(
            SOURCES / "Manuscript_final_scientific_lock.md",
            manuscript,
            body_size=12,
            double_space=True,
            line_numbers=True,
            running_header="npj Systems Biology and Applications | Article",
            title_override=TITLE,
        ),
        documents.markdown_to_docx(
            SOURCES / "Supplementary_Information_final_scientific_lock.md",
            supplement,
            body_size=10.5,
            double_space=False,
            line_numbers=False,
            running_header="Supplementary information",
            title_override="Supplementary information",
            supplementary_figure_dirs=[FIGURES],
            page_break_before_headings=set(),
        ),
    ]
    patch_properties(manuscript, "Final scientific manuscript lock")
    patch_properties(supplement, "Final scientific supplementary-information lock")

    main_document = Document(manuscript)
    supplementary_document = Document(supplement)
    main_text = extract_text(main_document)
    supplementary_text = extract_text(supplementary_document)
    abstract_match = re.search(r"Abstract\n(.+?)\nIntroduction", main_text, flags=re.DOTALL)
    if not abstract_match:
        raise RuntimeError("Could not identify manuscript abstract")
    abstract_words = re.findall(r"\b[\w'-]+\b", abstract_match.group(1))

    checks = {
        "manuscript": {
            "title_exact": TITLE in main_text,
            "abstract_145_words": len(abstract_words) == 145,
            "no_inline_figures": len(main_document.inline_shapes) == 0,
            "figure1_boundary_present": "defines a bounded analysis scaffold" in main_text,
            "figure5_boundary_present": "IFN-centred regulatory context" in main_text,
        },
        "supplement": {
            "ten_figures_embedded": len(supplementary_document.inline_shapes) == 10,
            "title_exact": TITLE in supplementary_text,
            "archive_scope_present": ARCHIVE_SCOPE in supplementary_text,
            "old_archive_claim_absent": "matches the frozen manuscript, figures" not in supplementary_text,
            "dataset_specific_unit_present_twice": supplementary_text.count(UNIT) == 2,
            "umbrella_unit_absent": "Donor/sample pseudobulk" not in supplementary_text,
            "current_doi_present": "10.5281/zenodo.22151739" in supplementary_text,
        },
    }
    failed = [
        f"{group}.{name}"
        for group, group_checks in checks.items()
        for name, passed in group_checks.items()
        if not passed
    ]
    if failed:
        raise RuntimeError(f"Traceability document checks failed: {failed}")

    status = {
        "created_at": datetime.now().astimezone().isoformat(timespec="seconds"),
        "status": "PASS_TRACEABILITY_LOCK_DOCX_BUILT_DUAL_RENDER_REQUIRED",
        "build_results": results,
        "checks": checks,
        "object_inventory": {
            "manuscript_inline_shapes": len(main_document.inline_shapes),
            "supplement_inline_shapes": len(supplementary_document.inline_shapes),
        },
        "files": {
            path.name: {"bytes": path.stat().st_size, "sha256": sha256(path)}
            for path in (manuscript, supplement)
        },
        "scientific_estimates_changed": False,
        "source_data_changed": False,
        "figures_changed": False,
    }
    (RUN / "01_DOCUMENT_BUILD_STATUS.json").write_text(
        json.dumps(status, indent=2) + "\n", encoding="utf-8", newline="\n"
    )
    print(json.dumps(status, indent=2))


if __name__ == "__main__":
    main()
