"""Build npj SBA DOCX artifacts from the target-refrozen Markdown sources."""

from __future__ import annotations

import hashlib
import json
from pathlib import Path

from docx import Document

import phase17_c8s_04_build_documents as documents


ROOT = Path(__file__).resolve().parents[1]
RUN = ROOT / "phase17_v7/npj_sba_target_refreeze/20260830_target_specific_refreeze"
SOURCES = RUN / "sources"
OUTPUT = RUN / "documents"
FIGURE_DIR = RUN / "figures/figures"
TITLE = (
    "Disease-blind reconstruction distinguishes reproducible interferon remodeling from unstable B-cell state assignments "
    "in systemic lupus erythematosus"
)


def sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest().upper()


def docx_text(path: Path) -> str:
    document = Document(path)
    values = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            values.extend(cell.text for cell in row.cells)
    return "\n".join(values)


def patch_properties(path: Path, title: str, subject: str) -> None:
    document = Document(path)
    document.core_properties.author = "Zhi Chen; Teng Qi"
    document.core_properties.title = title
    document.core_properties.subject = subject
    document.core_properties.comments = "Generated reproducibly from the npj SBA target-refrozen Markdown source."
    document.save(path)


def main() -> None:
    OUTPUT.mkdir(parents=True, exist_ok=True)
    manuscript = OUTPUT / "Manuscript.docx"
    supplement = OUTPUT / "Supplementary_Information.docx"
    cover = OUTPUT / "Cover_Letter.docx"
    outputs = [
        documents.markdown_to_docx(
            SOURCES / "Manuscript.md",
            manuscript,
            body_size=12,
            double_space=True,
            line_numbers=True,
            running_header="npj Systems Biology and Applications | Article",
            title_override=TITLE,
        ),
        documents.markdown_to_docx(
            SOURCES / "Supplementary_Information.md",
            supplement,
            body_size=10.5,
            double_space=False,
            line_numbers=False,
            running_header="Supplementary information",
            title_override="Supplementary information",
            supplementary_figure_dirs=[FIGURE_DIR],
            page_break_before_headings={
                "Supplementary Table S7 | Statistical tests and multiplicity families"
            },
        ),
        documents.markdown_to_docx(
            SOURCES / "Cover_Letter.md",
            cover,
            body_size=10.5,
            double_space=False,
            line_numbers=False,
            running_header=None,
            title_override="Cover letter to npj Systems Biology and Applications",
            compact=True,
        ),
    ]
    patch_properties(manuscript, TITLE, "Article for npj Systems Biology and Applications")
    patch_properties(supplement, "Supplementary information", TITLE)
    patch_properties(cover, "Cover letter", "npj Systems Biology and Applications")
    manuscript_text = docx_text(manuscript)
    supplement_doc = Document(supplement)
    supplement_text = docx_text(supplement)
    cover_text = docx_text(cover)
    checks = {
        "manuscript_title_exact": TITLE in manuscript_text,
        "manuscript_article_type": "Article type: Article" in manuscript_text,
        "manuscript_introduction": "Introduction" in manuscript_text,
        "manuscript_no_background_heading": "\nBackground\n" not in "\n" + manuscript_text + "\n",
        "manuscript_no_conclusions_heading": "\nConclusions\n" not in "\n" + manuscript_text + "\n",
        "manuscript_data_availability": "Data availability" in manuscript_text,
        "manuscript_code_availability": "Code availability" in manuscript_text,
        "manuscript_ai_disclosure": "Generative AI assistance" in manuscript_text,
        "manuscript_current_doi": "10.5281/zenodo.22151739" in manuscript_text,
        "manuscript_no_old_doi": "10.5281/zenodo.22086892" not in manuscript_text,
        "supplement_title_exact": TITLE in supplement_text,
        "supplement_no_methods": "Supplementary Methods" not in supplement_text,
        "supplement_figures_embedded": len(supplement_doc.inline_shapes) == 10,
        "cover_names_target": "npj Systems Biology and Applications" in cover_text,
        "cover_does_not_claim_exact_file_approval": all(
            phrase not in cover_text
            for phrase in (
                "approval of the exact npj-formatted files",
                "the exact npj-formatted files have been approved",
                "all authors have approved this submission",
            )
        ),
        "cover_no_submission_authorization_claim": all(
            phrase not in cover_text
            for phrase in (
                "authorization for portal submission remain pending",
                "will be completed before submission",
                "submission has been authorized",
            )
        ),
    }
    if not all(checks.values()):
        raise RuntimeError(f"Document structural checks failed: {[key for key, value in checks.items() if not value]}")
    status = {
        "created_at": "2026-08-30",
        "status": "PASS_NPJ_SBA_DOCX_BUILT_RENDER_QA_REQUIRED",
        "design": "Nature Portfolio submission manuscript; Times New Roman body, double-spaced main text, continuous line numbering",
        "outputs": outputs,
        "checks": checks,
        "files": {
            path.name: {"bytes": path.stat().st_size, "sha256": sha256(path)}
            for path in (manuscript, supplement, cover)
        },
        "scientific_reanalysis": False,
        "exact_file_author_approval": False,
        "submission_authorized": False,
    }
    (RUN / "02_DOCUMENT_BUILD_STATUS.json").write_text(json.dumps(status, indent=2) + "\n", encoding="utf-8")
    print(json.dumps({"status": status["status"], "checks_passed": sum(checks.values()), "outputs": status["files"]}, indent=2))


if __name__ == "__main__":
    main()
