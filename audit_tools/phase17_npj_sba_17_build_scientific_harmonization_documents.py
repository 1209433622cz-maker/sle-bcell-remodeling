#!/usr/bin/env python3
"""Build the scientific-harmonization manuscript and supplement candidates."""

from __future__ import annotations

import hashlib
import json
from datetime import datetime
from pathlib import Path

from docx import Document

import phase17_c8s_04_build_documents as documents


ROOT = Path(__file__).resolve().parents[1]
RUN = (
    ROOT
    / "phase17_v7/npj_sba_selected_supplementary_refinement/"
    "20260831_s4_s10_semantic_harmonization"
)
SOURCES = RUN / "sources"
FIGURES = RUN / "figures/figures"
OUTPUT = RUN / "documents"
TITLE = (
    "Disease-blind reconstruction distinguishes reproducible interferon remodeling from unstable B-cell state assignments "
    "in systemic lupus erythematosus"
)


def sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest().upper()


def docx_text(path: Path) -> str:
    document = Document(path)
    values = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            values.extend(cell.text for cell in row.cells)
    return "\n".join(values)


def patch_properties(path: Path, title: str, subject: str) -> None:
    document = Document(path)
    document.core_properties.author = "Zhi Chen; Teng Qi"
    document.core_properties.title = title
    document.core_properties.subject = subject
    document.core_properties.comments = (
        "Generated reproducibly from the scientific-harmonization candidate; "
        "the exact submission package remains unchanged."
    )
    document.save(path)


def main() -> None:
    OUTPUT.mkdir(parents=True, exist_ok=True)
    manuscript = OUTPUT / "Manuscript_scientific_harmonization_candidate.docx"
    supplement = OUTPUT / "Supplementary_Information_scientific_harmonization_candidate.docx"
    build_results = [
        documents.markdown_to_docx(
            SOURCES / "Manuscript_scientific_harmonization_candidate.md",
            manuscript,
            body_size=12,
            double_space=True,
            line_numbers=True,
            running_header="npj Systems Biology and Applications | Article",
            title_override=TITLE,
        ),
        documents.markdown_to_docx(
            SOURCES / "Supplementary_Information_scientific_harmonization_candidate.md",
            supplement,
            body_size=10.5,
            double_space=False,
            line_numbers=False,
            running_header="Supplementary information",
            title_override="Supplementary information",
            supplementary_figure_dirs=[FIGURES],
            page_break_before_headings=set(),
        ),
    ]
    patch_properties(
        manuscript,
        TITLE,
        "Scientific-harmonization candidate guided by npj Systems Biology and Applications",
    )
    patch_properties(supplement, "Supplementary information", TITLE)

    manuscript_text = docx_text(manuscript)
    supplement_document = Document(supplement)
    supplement_text = docx_text(supplement)
    manuscript_checks = {
        "title_exact": TITLE in manuscript_text,
        "article_type_present": "Article type: Article" in manuscript_text,
        "retained_scaffold_language": "retained disease-blind scaffold" in manuscript_text,
        "fine_state_boundary": "hard fine-state assignments" in manuscript_text,
        "source_matrix_language": "B-lineage source matrix contained 152,981 cells" in manuscript_text,
        "statistical_implementation_language": "used to test the statistical implementation" in manuscript_text,
        "anti_pseudoreplication_boundary": (
            "no inferential test treated genes as biological replicates" in manuscript_text
        ),
        "camera_capitalized": "CAMERA FDR of 1.85" in manuscript_text,
        "british_analogue": "source-label-defined broad B-cell analogue" in manuscript_text,
        "figure1_legend_scope": "permissible B_CONV/B_ASC scaffold" in manuscript_text,
        "figure5_ulm_ownership": "ULM STAT1/STAT2 activity was positive" in manuscript_text,
        "figure5_positive_arm": "all 12 frozen positive-arm genes" in manuscript_text,
        "causal_boundary": (
            "not a causal regulator, direct binding or a uniquely upstream ligand" in manuscript_text
        ),
        "current_doi_present": "10.5281/zenodo.22151739" in manuscript_text,
        "old_doi_absent": "10.5281/zenodo.22086892" not in manuscript_text,
        "no_embedded_figures": len(Document(manuscript).inline_shapes) == 0,
    }
    supplement_checks = {
        "ten_figures_embedded": len(supplement_document.inline_shapes) == 10,
        "s4_log_scale_legend": "Panels c-d use logarithmic ratio axes" in supplement_text,
        "s4_null_value_defined": "the null fixed at one" in supplement_text,
        "s10_mapper_colour": "Mapper colour is held constant across panels b-d" in supplement_text,
        "s10_state_shape": "marker shape distinguishes B_CONV from B_ASC" in supplement_text,
        "s10_balanced_accuracy_diagnostic": "balanced accuracy is diagnostic only" in supplement_text,
        "current_doi_present": "10.5281/zenodo.22151739" in supplement_text,
        "old_doi_absent": "10.5281/zenodo.22086892" not in supplement_text,
    }
    checks = {"manuscript": manuscript_checks, "supplement": supplement_checks}
    failed = [
        f"{group}.{name}"
        for group, group_checks in checks.items()
        for name, passed in group_checks.items()
        if not passed
    ]
    if failed:
        raise RuntimeError(f"Scientific candidate document checks failed: {failed}")

    status = {
        "created_at": datetime.now().astimezone().isoformat(timespec="seconds"),
        "status": "PASS_SCIENTIFIC_HARMONIZATION_DOCX_BUILT_RENDER_QA_REQUIRED",
        "build_results": build_results,
        "files": {
            path.name: {
                "bytes": path.stat().st_size,
                "sha256": sha256(path),
            }
            for path in (manuscript, supplement)
        },
        "checks": checks,
        "scientific_estimates_changed": False,
        "source_data_changed": False,
        "exact_submission_package_modified": False,
    }
    (RUN / "02_DOCUMENT_BUILD_STATUS.json").write_text(
        json.dumps(status, indent=2) + "\n", encoding="utf-8", newline="\n"
    )
    print(
        json.dumps(
            {
                "status": status["status"],
                "manuscript_checks": sum(manuscript_checks.values()),
                "supplement_checks": sum(supplement_checks.values()),
                "files": status["files"],
            },
            indent=2,
        )
    )


if __name__ == "__main__":
    main()
