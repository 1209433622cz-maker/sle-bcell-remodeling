import unittest

from docx import Document

from phase17_postc9_14_audit_refined_manuscript import citation_order, compare_text, numeric_tokens, reference_map
from phase17_postc9_15_build_refined_review import apply_changes, suppress_header_footer_numbers


class RefinedManuscriptTests(unittest.TestCase):
    def test_reference_map_handles_doi_and_geo_without_loss(self):
        a = "A. doi:10.1/example."
        b = "NCBI. https://www.ncbi.nlm.nih.gov/geo/query/acc.cgi?acc=GSE135779."
        result = reference_map([a, b], [b, a])
        self.assertEqual([row["new_number"] for row in result], [2, 1])

    def test_reference_omission_fails(self):
        with self.assertRaises(ValueError):
            reference_map(["A. doi:10.1/a.", "B. doi:10.1/b."], ["A. doi:10.1/a."])

    def test_duplicate_reference_fails(self):
        with self.assertRaises(ValueError):
            reference_map(["A. doi:10.1/a."], ["A. doi:10.1/a."] * 2)

    def test_citation_ranges_and_repeats(self):
        self.assertEqual(citation_order("Text [1-3], more [2,4]."), [1, 2, 3, 4])

    def test_reversed_citation_range_fails(self):
        with self.assertRaises(ValueError):
            citation_order("Text [4-1].")

    def test_numeric_punctuation_and_reference_exclusion(self):
        self.assertEqual(numeric_tokens("-1, 150,402 and 2.98 x 10^-6 [1-3]"),
                         numeric_tokens("-1; 150,402 and 2.98 x 10^-6."))

    def test_pdf_comparison_does_not_hide_an_extra_x(self):
        self.assertFalse(compare_text("Data available.", "Data available.X")["normalized_equal"])
        self.assertTrue(compare_text("B-cell\nresponse.", "B-cell response.")["normalized_equal"])

    def test_local_edit_preserves_unrelated_paragraph(self):
        doc = Document()
        doc.add_paragraph("Effect 0.947 was not higher.")
        doc.add_paragraph("Other paragraph.")
        apply_changes(doc, [{"id": "x", "paragraph": 0, "old": "was not higher", "new": "lacked statistical support"}])
        self.assertEqual(doc.paragraphs[1].text, "Other paragraph.")

    def test_changed_estimate_fails(self):
        doc = Document()
        doc.add_paragraph("Effect 0.947.")
        with self.assertRaises(ValueError):
            apply_changes(doc, [{"id": "x", "paragraph": 0, "old": "0.947", "new": "1.947"}])

    def test_ambiguous_replacement_fails(self):
        doc = Document()
        doc.add_paragraph("repeat repeat")
        with self.assertRaises(ValueError):
            apply_changes(doc, [{"id": "x", "paragraph": 0, "old": "repeat", "new": "changed"}])

    def test_header_footer_suppression_is_idempotent(self):
        doc = Document()
        for _ in range(2):
            self.assertEqual(suppress_header_footer_numbers(doc), 2)
        for container in (doc.sections[0].header, doc.sections[0].footer):
            self.assertEqual(len(container.paragraphs[0]._p.xpath("./w:pPr/w:suppressLineNumbers")), 1)


if __name__ == "__main__":
    unittest.main()
