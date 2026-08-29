"""Build the journal-neutral current cover-letter draft as an editable DOCX."""

from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
WORK = ROOT / "00_project_management/current_submission_package_2026-08-30"
SOURCE = WORK / "Cover_Letter_Draft.md"
OUTPUT = WORK / "Cover_Letter_Draft.docx"
TITLE = "Disease-blind single-cell reconstruction distinguishes unstable B-cell state assignments from reproducible interferon remodeling in systemic lupus erythematosus"
CURRENT_DOI = "10.5281/zenodo.22151739"
FONT = "Times New Roman"
BLACK = RGBColor(0, 0, 0)
MUTED = RGBColor(89, 89, 89)


def require(condition, message):
    if not condition:
        raise ValueError(message)


def blocks(text):
    output = []
    pending = []
    for line in text.splitlines():
        stripped = line.strip()
        if stripped.startswith("# "):
            continue
        if not stripped:
            if pending:
                output.append(" ".join(pending))
                pending.clear()
        else:
            pending.append(stripped)
    if pending:
        output.append(" ".join(pending))
    return output


def format_run(run, size=10.5, bold=None, italic=None, color=BLACK):
    run.font.name = FONT
    run.font.size = Pt(size)
    run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    rpr = run._element.get_or_add_rPr()
    fonts = rpr.rFonts
    if fonts is None:
        fonts = OxmlElement("w:rFonts")
        rpr.insert(0, fonts)
    for attribute in ("ascii", "hAnsi", "eastAsia", "cs"):
        fonts.set(qn("w:" + attribute), FONT)
    return run


def validate_source(text):
    for phrase in (
        TITLE,
        CURRENT_DOI,
        "permanent R1 HOLD",
        "C9R HOLD",
        "no corrected external disease outcome was estimated",
        "five vector main figures",
        "ten supplementary figures",
        "not under consideration by another journal",
    ):
        require(phrase in text, "Cover-letter boundary is missing: " + phrase)
    for stale in (
        "Genome Medicine",
        "10.5281/zenodo.22086892",
        "matching revised archive remains required",
        "Both authors approved the manuscript, supplementary information, figures, source data, cover letter and submission",
    ):
        require(stale not in text, "Stale or unauthorized cover-letter statement remains: " + stale)


def build(source, output):
    text = source.read_text(encoding="utf-8")
    validate_source(text)
    content = blocks(text)
    require(content[0] == "30 August 2026", "Cover-letter date block differs")
    require(content[-1] == "tengqi@link.cuhk.edu.cn", "Cover-letter sign-off differs")

    document = Document()
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.72)
    section.bottom_margin = Inches(0.68)
    section.left_margin = Inches(0.88)
    section.right_margin = Inches(0.88)
    section.footer_distance = Inches(0.32)

    normal = document.styles["Normal"]
    normal.font.name = FONT
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = BLACK
    for attribute in ("ascii", "hAnsi", "eastAsia", "cs"):
        normal._element.rPr.rFonts.set(qn("w:" + attribute), FONT)
    normal.paragraph_format.line_spacing = 1.02
    normal.paragraph_format.space_after = Pt(5.5)
    normal.paragraph_format.widow_control = True

    for index, block in enumerate(content):
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.keep_together = True
        if index == 0:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        if block == "Dear Editors,":
            paragraph.paragraph_format.space_before = Pt(4)
        if block == "Sincerely,":
            paragraph.paragraph_format.space_before = Pt(4)
        run = paragraph.add_run(block)
        format_run(run)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    format_run(
        footer.add_run("Journal-neutral draft | Target selection and exact-file author approval pending"),
        size=8,
        italic=True,
        color=MUTED,
    )

    properties = document.core_properties
    properties.title = "Journal-neutral cover letter draft"
    properties.subject = "SLE B-cell remodeling manuscript"
    properties.author = "Zhi Chen and Teng Qi"
    properties.comments = "Not authorized for journal submission; adapt after target freeze and obtain exact-file approval."
    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)
    print(f"Wrote: {output}")
    print(f"Paragraphs: {len(document.paragraphs)}")
    print(f"Bytes: {output.stat().st_size}")


if __name__ == "__main__":
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--source", type=Path, default=SOURCE)
    parser.add_argument("--output", type=Path, default=OUTPUT)
    args = parser.parse_args()
    build(args.source.resolve(), args.output.resolve())
