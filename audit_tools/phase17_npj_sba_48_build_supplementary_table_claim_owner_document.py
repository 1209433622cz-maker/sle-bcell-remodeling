#!/usr/bin/env python3
"""Build the manuscript after the narrow Supplementary Table claim-owner pass."""

from __future__ import annotations

import hashlib
import json
import re
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Pt

import phase17_c8s_04_build_documents as documents


ROOT = Path(__file__).resolve().parents[1]
RUN = ROOT / "phase17_v7/npj_sba_supplementary_table_claim_owner/20260902_semantic_micropass"
SOURCE = RUN / "sources/Manuscript_claim_owner_semantic_micropass.md"
DOCUMENTS = RUN / "documents"
OUTPUT = DOCUMENTS / "Manuscript_claim_owner_semantic_micropass.docx"
TITLE = (
    "Disease-blind reconstruction distinguishes reproducible interferon remodeling from less stable "
    "B-cell state assignments in systemic lupus erythematosus"
)


def sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest().upper()


def extract_text(document: Document) -> str:
    values = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            values.extend(cell.text for cell in row.cells)
    return "\n".join(values)


def patch_properties(path: Path) -> int:
    document = Document(path)
    compacted = 0
    for paragraph in document.paragraphs:
        if re.fullmatch(r"Figure [1-5] \| .+", paragraph.text.strip()):
            paragraph.paragraph_format.space_before = Pt(6)
            paragraph.paragraph_format.space_after = Pt(2)
            compacted += 1
    document.core_properties.author = "Zhi Chen; Teng Qi"
    document.core_properties.title = TITLE
    document.core_properties.subject = "Supplementary Table claim-owner semantic micropass"
    document.core_properties.comments = "Six exact citation-anchor operations; no scientific-value or artwork change."
    document.save(path)
    return compacted


def main() -> None:
    integration = json.loads((RUN / "00_CLAIM_OWNER_MICROPASS_INTEGRATION_STATUS.json").read_text(encoding="utf-8"))
    if integration["failed_checks"]:
        raise RuntimeError("Integration status contains failed checks")
    DOCUMENTS.mkdir(parents=True, exist_ok=True)
    build_result = documents.markdown_to_docx(
        SOURCE,
        OUTPUT,
        body_size=12,
        double_space=True,
        line_numbers=True,
        running_header="npj Systems Biology and Applications | Article",
        title_override=TITLE,
    )
    compacted = patch_properties(OUTPUT)
    document = Document(OUTPUT)
    text = extract_text(document)
    source_text = SOURCE.read_text(encoding="utf-8")
    abstract_match = re.search(r"Abstract\n(.+?)\nIntroduction", text, flags=re.DOTALL)
    if not abstract_match:
        raise RuntimeError("Could not identify manuscript abstract")
    abstract_words = re.findall(r"\b[\w'-]+\b", abstract_match.group(1))
    reference_numbers = [
        int(value)
        for value in re.findall(
            r"(?m)^(\d+)\. ",
            source_text.split("## References\n", 1)[1].split("## Figure legends\n", 1)[0],
        )
    ]
    checks = {
        "title_exact": TITLE in text,
        "abstract_145_words": len(abstract_words) == 145,
        "references_1_to_33": reference_numbers == list(range(1, 34)),
        "main_has_no_inline_figures": len(document.inline_shapes) == 0,
        "five_main_legend_headings_compacted": compacted == 5,
        "s1_s2_owner_present": "Supplementary Tables S1 and S2" in text,
        "s3_quantitative_owner_present": "principal quantitative anchors summarized in Supplementary Table S3" in text,
        "s3_causal_misowner_absent": "causal regulation in SLE (Supplementary Table S3)" not in text,
        "s4a_owner_present": "Supplementary Fig. S9; Supplementary Table S4a" in text,
        "s4b_owner_present": "Supplementary Fig. S10; Supplementary Table S4b" in text,
        "generic_s4_owner_absent": "Supplementary Table S4)" not in text,
        "s5_s8_reproducibility_owner_present": "SHA-256 manifests (Supplementary Tables S5-S8)" in text,
        "s5_s8_superseded_misowner_absent": "present version (Supplementary Tables S5-S8)" not in text,
        "s9_owner_present": "Supplementary Table S9 and Supplementary Fig. S8" in text,
    }
    failed = [name for name, passed in checks.items() if not passed]
    status = {
        "created_at": datetime.now().astimezone().isoformat(timespec="seconds"),
        "status": "PASS_CLAIM_OWNER_DOCX_BUILT_RENDER_REQUIRED" if not failed else "FAIL_CLAIM_OWNER_DOCX_BUILD",
        "build_result": build_result,
        "checks": checks,
        "failed_checks": failed,
        "document": {
            "path": OUTPUT.relative_to(ROOT).as_posix(),
            "bytes": OUTPUT.stat().st_size,
            "sha256": sha256(OUTPUT),
        },
        "scientific_estimates_changed": False,
        "figures_redrawn": False,
        "source_data_values_changed": False,
    }
    (RUN / "06_DOCUMENT_BUILD_STATUS.json").write_text(
        json.dumps(status, indent=2) + "\n", encoding="utf-8", newline="\n"
    )
    print(json.dumps(status, indent=2))
    if failed:
        raise RuntimeError(f"Document build checks failed: {failed}")


if __name__ == "__main__":
    main()
