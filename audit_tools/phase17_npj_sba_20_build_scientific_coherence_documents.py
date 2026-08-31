#!/usr/bin/env python3
"""Build DOCX files for the scientific-coherence refreeze candidate."""

from __future__ import annotations

import hashlib
import json
import re
from datetime import datetime
from pathlib import Path

from docx import Document

import phase17_c8s_04_build_documents as documents


ROOT = Path(__file__).resolve().parents[1]
RUN = ROOT / "phase17_v7/npj_sba_scientific_coherence_refreeze/20260831_claim_order_reader_boundaries"
SOURCES = RUN / "sources"
FIGURES = RUN / "figures/figures"
OUTPUT = RUN / "documents"
TITLE = (
    "Disease-blind reconstruction distinguishes reproducible interferon remodeling from less stable "
    "B-cell state assignments in systemic lupus erythematosus"
)


def sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest().upper()


def extract_text(document: Document) -> str:
    values = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            values.extend(cell.text for cell in row.cells)
    return "\n".join(values)


def patch_properties(path: Path, subject: str) -> None:
    document = Document(path)
    document.core_properties.author = "Zhi Chen; Teng Qi"
    document.core_properties.title = TITLE if "Manuscript" in path.name else "Supplementary information"
    document.core_properties.subject = subject
    document.core_properties.comments = (
        "Scientific-coherence candidate regenerated from the frozen repository evidence; "
        "the author-confirmed exact submission package remains unchanged."
    )
    document.save(path)


def main() -> None:
    OUTPUT.mkdir(parents=True, exist_ok=True)
    manuscript = OUTPUT / "Manuscript_scientific_coherence_refreeze_candidate.docx"
    supplement = OUTPUT / "Supplementary_Information_scientific_coherence_refreeze_candidate.docx"
    build_results = [
        documents.markdown_to_docx(
            SOURCES / "Manuscript_scientific_coherence_refreeze_candidate.md",
            manuscript,
            body_size=12,
            double_space=True,
            line_numbers=True,
            running_header="npj Systems Biology and Applications | Article",
            title_override=TITLE,
        ),
        documents.markdown_to_docx(
            SOURCES / "Supplementary_Information_scientific_coherence_refreeze_candidate.md",
            supplement,
            body_size=10.5,
            double_space=False,
            line_numbers=False,
            running_header="Supplementary information",
            title_override="Supplementary information",
            supplementary_figure_dirs=[FIGURES],
            page_break_before_headings=set(),
        ),
    ]
    patch_properties(
        manuscript,
        "Scientific-coherence manuscript candidate guided by npj Systems Biology and Applications",
    )
    patch_properties(supplement, "Scientific-coherence supplementary information candidate")

    manuscript_document = Document(manuscript)
    supplement_document = Document(supplement)
    manuscript_text = extract_text(manuscript_document)
    supplement_text = extract_text(supplement_document)
    abstract_match = re.search(r"Abstract\n(.+?)\nIntroduction", manuscript_text, flags=re.DOTALL)
    if not abstract_match:
        raise RuntimeError("Could not identify the abstract in the generated manuscript")
    abstract_words = re.findall(r"\b[\w'-]+\b", abstract_match.group(1))

    manuscript_checks = {
        "title_exact": TITLE in manuscript_text,
        "article_type_present": "Article type: Article" in manuscript_text,
        "abstract_145_words": len(abstract_words) == 145,
        "comparative_landing": (
            "reproducibility was stronger for a process-level interferon program" in manuscript_text
        ),
        "composition_null_language_absent": "primary composition null" not in manuscript_text,
        "figure1_retained_analysis_scope": (
            "Disease-blind reconstruction defines the retained analysis scope" in manuscript_text
        ),
        "first_evidence_owners_present": all(
            token in manuscript_text
            for token in (
                "Supplementary Fig. S1",
                "Supplementary Fig. S2",
                "Supplementary Fig. S3",
                "Fig. 2a-d; Supplementary Fig. S4",
                "Fig. 3a-c; Supplementary Fig. S5",
                "Fig. 4a,b; Supplementary Fig. S6",
                "Supplementary Fig. S7",
                "Supplementary Fig. S8",
                "Supplementary Fig. S10",
            )
        ),
        "source_label_independent_heading": (
            "Corrected source-label-independent remapping does not satisfy" in manuscript_text
        ),
        "causal_boundary": (
            "not a causal regulator, direct binding or a uniquely upstream ligand" in manuscript_text
        ),
        "current_doi_present": "10.5281/zenodo.22151739" in manuscript_text,
        "old_doi_absent": "10.5281/zenodo.22086892" not in manuscript_text,
        "no_embedded_figures": len(manuscript_document.inline_shapes) == 0,
    }
    supplement_checks = {
        "ten_figures_embedded": len(supplement_document.inline_shapes) == 10,
        "reader_facing_s3_heading": (
            "Quantitative anchors and prespecified boundaries" in supplement_text
        ),
        "criterion_not_hold": (
            "prespecified criterion not met" in supplement_text
            and "formal HOLD" not in supplement_text
        ),
        "composition_not_null": (
            "Primary B_ASC contrast lacks statistical support" in supplement_text
            and "Null primary B_ASC" not in supplement_text
        ),
        "ulm_ownership": "ULM STAT1/STAT2 activity concordant" in supplement_text,
        "s9_legend_synchronized": (
            "four met their criteria and minimum state-median Jaccard did not" in supplement_text
        ),
        "s10_source_label_independent": (
            "Reference calibration limits source-label-independent external transfer" in supplement_text
        ),
        "internal_c9_pass_absent": "C9 PASS" not in supplement_text,
        "current_doi_present": "10.5281/zenodo.22151739" in supplement_text,
        "old_doi_absent": "10.5281/zenodo.22086892" not in supplement_text,
    }
    checks = {"manuscript": manuscript_checks, "supplement": supplement_checks}
    failed = [
        f"{group}.{name}"
        for group, group_checks in checks.items()
        for name, passed in group_checks.items()
        if not passed
    ]
    if failed:
        raise RuntimeError(f"Scientific-coherence document checks failed: {failed}")

    status = {
        "created_at": datetime.now().astimezone().isoformat(timespec="seconds"),
        "status": "PASS_SCIENTIFIC_COHERENCE_DOCX_BUILT_DUAL_RENDER_REQUIRED",
        "build_results": build_results,
        "object_inventory": {
            "manuscript_inline_shapes": len(manuscript_document.inline_shapes),
            "supplement_inline_shapes": len(supplement_document.inline_shapes),
            "supplement_expected_figures": 10,
            "supplement_object_count_matches": len(supplement_document.inline_shapes) == 10,
        },
        "files": {
            path.name: {"bytes": path.stat().st_size, "sha256": sha256(path)}
            for path in (manuscript, supplement)
        },
        "checks": checks,
        "scientific_estimates_changed": False,
        "source_data_changed": False,
        "exact_submission_package_modified": False,
    }
    (RUN / "02_DOCUMENT_BUILD_STATUS.json").write_text(
        json.dumps(status, indent=2) + "\n", encoding="utf-8", newline="\n"
    )
    print(json.dumps({
        "status": status["status"],
        "abstract_words": len(abstract_words),
        "object_inventory": status["object_inventory"],
        "files": status["files"],
    }, indent=2))


if __name__ == "__main__":
    main()
