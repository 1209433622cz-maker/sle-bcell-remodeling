#!/usr/bin/env python3
"""Build the main manuscript after Figure 1 boundary promotion."""

from __future__ import annotations

import hashlib
import json
import re
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Pt

import phase17_c8s_04_build_documents as documents


ROOT = Path(__file__).resolve().parents[1]
RUN = ROOT / "phase17_v7/npj_sba_figure1_boundary_promotion/20260902_source_rerender_gate"
SOURCE = RUN / "sources/Manuscript_figure1_boundary_promotion.md"
DOCUMENTS = RUN / "documents"
OUTPUT = DOCUMENTS / "Manuscript_Figure1_boundary_promotion.docx"
TITLE = (
    "Disease-blind reconstruction distinguishes reproducible interferon remodeling from less stable "
    "B-cell state assignments in systemic lupus erythematosus"
)


def sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest().upper()


def extract_text(document: Document) -> str:
    values = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            values.extend(cell.text for cell in row.cells)
    return "\n".join(values)


def patch_properties(path: Path) -> int:
    document = Document(path)
    compacted = 0
    for paragraph in document.paragraphs:
        if re.fullmatch(r"Figure [1-5] \| .+", paragraph.text.strip()):
            paragraph.paragraph_format.space_before = Pt(6)
            paragraph.paragraph_format.space_after = Pt(2)
            compacted += 1
    document.core_properties.author = "Zhi Chen; Teng Qi"
    document.core_properties.title = TITLE
    document.core_properties.subject = "Figure 1 end-to-end identity-boundary promotion"
    document.core_properties.comments = (
        "Source-rerendered Figure 1; three exact text operations; no statistical-value change."
    )
    document.save(path)
    return compacted


def main() -> None:
    integration = json.loads(
        (RUN / "00_FIGURE1_BOUNDARY_PROMOTION_INTEGRATION_STATUS.json").read_text(encoding="utf-8")
    )
    if integration["failed_checks"]:
        raise RuntimeError("Figure 1 integration status contains failed checks")
    DOCUMENTS.mkdir(parents=True, exist_ok=True)
    build_result = documents.markdown_to_docx(
        SOURCE,
        OUTPUT,
        body_size=12,
        double_space=True,
        line_numbers=True,
        running_header="npj Systems Biology and Applications | Article",
        title_override=TITLE,
    )
    compacted = patch_properties(OUTPUT)
    document = Document(OUTPUT)
    text = extract_text(document)
    source_text = SOURCE.read_text(encoding="utf-8")
    abstract_match = re.search(r"Abstract\n(.+?)\nIntroduction", text, flags=re.DOTALL)
    if not abstract_match:
        raise RuntimeError("Could not identify manuscript abstract")
    abstract_words = re.findall(r"\b[\w'-]+\b", abstract_match.group(1))
    reference_numbers = [
        int(value)
        for value in re.findall(
            r"(?m)^(\d+)\. ",
            source_text.split("## References\n", 1)[1].split("## Figure legends\n", 1)[0],
        )
    ]
    checks = {
        "title_exact": TITLE in text,
        "abstract_145_words": len(abstract_words) == 145,
        "references_1_to_33": reference_numbers == list(range(1, 34)),
        "main_has_no_inline_figures": len(document.inline_shapes) == 0,
        "five_main_legend_headings_compacted": compacted == 5,
        "fixed_result_points_to_figure1_a_c": "itself (Fig. 1a-c)." in text,
        "end_to_end_result_points_to_figure1d_and_s4": "Fig. 1d; Supplementary Fig. S4" in text,
        "figure1c_legend_owns_fixed_jaccard": "c, Fixed-representation minimum-to-median state Jaccard" in text,
        "figure1d_legend_owns_end_to_end_boundary": "d, End-to-end minimum-to-median state Jaccard across 20 rebuilds" in text,
        "s4_retains_detailed_owner": "Supplementary Fig. S4 provides replicate diagnostics/downstream propagation" in text,
        "old_figure1c_legend_removed": "c, Mapped adjusted Rand index and mapping agreement for each two-compartment resample" not in text,
    }
    failed = [name for name, passed in checks.items() if not passed]
    status = {
        "created_at": datetime.now().astimezone().isoformat(timespec="seconds"),
        "status": "PASS_FIGURE1_BOUNDARY_DOCUMENT_BUILT_RENDER_REQUIRED" if not failed else "FAIL_FIGURE1_BOUNDARY_DOCUMENT_BUILD",
        "build_result": build_result,
        "checks": checks,
        "failed_checks": failed,
        "document": {
            "path": OUTPUT.relative_to(ROOT).as_posix(),
            "bytes": OUTPUT.stat().st_size,
            "sha256": sha256(OUTPUT),
        },
        "scientific_estimates_changed": False,
        "figure1_redrawn": True,
        "source_data_values_changed": False,
    }
    (RUN / "05_DOCUMENT_BUILD_STATUS.json").write_text(
        json.dumps(status, indent=2) + "\n", encoding="utf-8", newline="\n"
    )
    print(json.dumps(status, indent=2))
    if failed:
        raise RuntimeError(f"Document build checks failed: {failed}")


if __name__ == "__main__":
    main()
