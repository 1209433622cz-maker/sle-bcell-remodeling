"""Build the DOI-integrated manuscript without changing frozen science."""

import argparse
from copy import deepcopy
from datetime import datetime
import json
from pathlib import Path

from docx import Document

from phase17_postc9_14_audit_refined_manuscript import extract_docx, file_record
from phase17_postc9_16_verify_scientific_freeze import ROOT, scientific_hashes


BASE = ROOT / "00_project_management/qiteng_text_audit_2026-08-29/review_candidate"
EXPECTED_DOI = "10.5281/zenodo.22151739"
OLD_DOI = "10.5281/zenodo.22086892"

OLD_DATA = (
    "The datasets analysed are publicly available through NCBI GEO under GSE174188, "
    "GSE135779 and GSE23307 [17,18,20]. The project repository is "
    "https://github.com/1209433622cz-maker/sle-bcell-remodeling; an initial immutable "
    "snapshot is archived at doi:10.5281/zenodo.22086892 [32]. That snapshot predates "
    "the end-to-end reconstruction and corrected external-mapping audits. A matching "
    "version-specific archive of the revised code, decisions, source data and SHA-256 "
    "records is required before submission. Original project code is licensed under the "
    "MIT License; original manuscript text, composite figures, project documentation and "
    "project-generated derived source-data tables are licensed under CC BY 4.0. These "
    "licences do not relicense GEO, CELLxGENE or other third-party source material. Large "
    "recomputable matrices are not duplicated from their source repositories."
)
NEW_DATA = (
    "The datasets analysed are publicly available through NCBI GEO under GSE174188, "
    "GSE135779 and GSE23307 [17,18,20]. The project repository is "
    "https://github.com/1209433622cz-maker/sle-bcell-remodeling; the version-specific "
    "reproducibility archive accompanying this manuscript is available at "
    "doi:10.5281/zenodo.22151739 [32] and supersedes the initial archive at "
    "doi:10.5281/zenodo.22086892. Original project code is licensed under the MIT License; "
    "original manuscript text, composite figures, project documentation and "
    "project-generated derived source-data tables are licensed under CC BY 4.0. These "
    "licences do not relicense GEO, CELLxGENE or other third-party source material. Large "
    "recomputable matrices are not duplicated from their source repositories."
)
OLD_CONTRIBUTIONS = (
    "ZC: Conceptualization, Data curation, Formal analysis, Investigation, Methodology, "
    "Software, Visualization, Writing - original draft. TQ: Conceptualization, Methodology, "
    "Project administration, Validation, Writing - review and editing. Both authors "
    "approved the preceding reviewed snapshot. Final approval of these exact refined "
    "manuscript files is pending."
)
NEW_CONTRIBUTIONS = (
    "ZC: Conceptualization, Data curation, Formal analysis, Investigation, Methodology, "
    "Software, Visualization, Writing - original draft. TQ: Conceptualization, Methodology, "
    "Project administration, Validation, Writing - review and editing. Both authors "
    "approved the scientific manuscript baseline, declarations, figures, source data and "
    "archive content. This administrative DOI integration does not alter the approved "
    "scientific text."
)
OLD_REFERENCE = (
    "32. Chen Z, Qi T. SLE B-cell remodeling analysis: code, source data and reproducible "
    "release. Zenodo. 2026. doi:10.5281/zenodo.22086892."
)
NEW_REFERENCE = OLD_REFERENCE.replace(OLD_DOI, EXPECTED_DOI)
MARKDOWN_CHANGES = {
    OLD_DATA: NEW_DATA,
    OLD_CONTRIBUTIONS: NEW_CONTRIBUTIONS,
    OLD_REFERENCE: NEW_REFERENCE,
}
DOCX_CHANGES = {
    OLD_DATA: NEW_DATA,
    OLD_CONTRIBUTIONS: NEW_CONTRIBUTIONS,
    OLD_REFERENCE.removeprefix("32. "): NEW_REFERENCE.removeprefix("32. "),
}


def replace_markdown(text):
    result = text
    for old, new in MARKDOWN_CHANGES.items():
        if result.count(old) != 1:
            raise ValueError("Expected exactly one Markdown match for an administrative paragraph")
        result = result.replace(old, new, 1)
    return result


def replace_docx(source, destination):
    document = Document(source)
    changed = []
    before = [paragraph.text for paragraph in document.paragraphs]
    for index, paragraph in enumerate(document.paragraphs):
        if paragraph.text not in DOCX_CHANGES:
            continue
        replacement = DOCX_CHANGES[paragraph.text]
        source_run = max(paragraph.runs, key=lambda run: len(run.text))
        properties = deepcopy(source_run._r.rPr)
        paragraph.clear()
        run = paragraph.add_run(replacement)
        if properties is not None:
            run._r.insert(0, properties)
        changed.append(index)
    if changed != [109, 115, 168]:
        raise ValueError(f"Unexpected DOCX paragraph changes: {changed}")
    after = [paragraph.text for paragraph in document.paragraphs]
    if len(before) != len(after):
        raise ValueError("DOCX paragraph count changed")
    for index, (old, new) in enumerate(zip(before, after)):
        if index not in changed and old != new:
            raise ValueError(f"Unrequested DOCX paragraph change: {index}")
    document.save(destination)
    return changed


def main():
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--doi", required=True)
    parser.add_argument("--output-dir", required=True, type=Path)
    parser.add_argument("--receipt", required=True, type=Path)
    args = parser.parse_args()
    if args.doi != EXPECTED_DOI:
        raise ValueError("DOI differs from the reserved linked-draft DOI")
    output = args.output_dir.resolve()
    release_root = (ROOT / "04_submission/zenodo_release").resolve()
    if not output.is_relative_to(release_root):
        raise ValueError("Manuscript output must stay in the Zenodo release directory")
    receipt_path = args.receipt.resolve()
    management = (ROOT / "00_project_management/qiteng_r2_release_2026-08-29").resolve()
    if not receipt_path.is_relative_to(management) or receipt_path.suffix != ".json":
        raise ValueError("Receipt must stay in the release project-management directory")

    baseline_md = (BASE / "Manuscript.md").read_text(encoding="utf-8-sig")
    baseline_science = scientific_hashes(baseline_md)
    candidate_md = replace_markdown(baseline_md)
    if scientific_hashes(candidate_md) != baseline_science:
        raise ValueError("Administrative Markdown update changed frozen science")

    output.mkdir(parents=True, exist_ok=True)
    markdown_path = output / "Manuscript.md"
    docx_path = output / "Manuscript.docx"
    markdown_path.write_text(candidate_md, encoding="utf-8")
    changed = replace_docx(BASE / "Manuscript.docx", docx_path)
    extracted, _, references, _ = extract_docx(docx_path)
    if scientific_hashes(extracted) != baseline_science:
        raise ValueError("Administrative DOCX update changed frozen science")
    if extracted != candidate_md:
        raise ValueError("DOCX and Markdown text differ after administrative update")
    if len(references) != 32 or EXPECTED_DOI not in references[-1]:
        raise ValueError("Archive reference was not updated correctly")
    if OLD_DOI not in extracted or EXPECTED_DOI not in extracted:
        raise ValueError("New and superseded DOI provenance is incomplete")

    receipt = {
        "created_at": datetime.now().astimezone().isoformat(timespec="seconds"),
        "status": "PASS_ADMINISTRATIVE_DOI_INTEGRATION_NOT_PUBLICATION",
        "reserved_doi": EXPECTED_DOI,
        "zenodo_record_id": "22151739",
        "source_files": [file_record(BASE / "Manuscript.docx"), file_record(BASE / "Manuscript.md")],
        "output_files": [file_record(docx_path), file_record(markdown_path)],
        "changed_docx_paragraphs": changed,
        "changes": [
            "Data Availability: replace the required-future archive statement with the reserved version DOI and retain the superseded DOI",
            "Authors' contributions: replace the historical pending sentence with the reported author-confirmed scope",
            "Reference 32: cite the reserved version DOI",
        ],
        "scientific_section_sha256": baseline_science,
        "scientific_sections_unchanged": True,
        "reference_count": len(references),
        "publication_claimed": False,
        "submission_authorized": False,
    }
    receipt_path.parent.mkdir(parents=True, exist_ok=True)
    receipt_path.write_text(json.dumps(receipt, indent=2) + "\n", encoding="utf-8")
    print(json.dumps(receipt, indent=2))


if __name__ == "__main__":
    main()
