#!/usr/bin/env python3
"""Build the Supplementary Table S4 reader-path maintenance candidate."""

from __future__ import annotations

import hashlib
import json
import re
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

import phase17_c8s_04_build_documents as documents


ROOT = Path(__file__).resolve().parents[1]
RUN = ROOT / "phase17_v7/npj_sba_supplementary_table_reader_path/20260901_s4_reader_path_refreeze"
SOURCE = RUN / "sources/Supplementary_Information_s4_reader_path_micropass.md"
FIGURES = RUN / "figures/figures"
OUTPUT = RUN / "documents"
TITLE = (
    "Disease-blind reconstruction distinguishes reproducible interferon remodeling from less stable "
    "B-cell state assignments in systemic lupus erythematosus"
)


def sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest().upper()


def extract_text(document: Document) -> str:
    values = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            values.extend(cell.text for cell in row.cells)
    return "\n".join(values)


def remove_table_spacers_before_manual_page_breaks(document: Document) -> int:
    removed = 0
    for paragraph in list(document.paragraphs):
        if not paragraph._p.xpath('.//w:br[@w:type="page"]'):
            continue
        previous = paragraph._p.getprevious()
        if previous is None or previous.tag != qn("w:p"):
            continue
        text = "".join(previous.itertext()).strip()
        has_break = bool(previous.xpath('.//w:br[@w:type="page"]'))
        has_drawing = bool(previous.xpath(".//w:drawing"))
        if not text and not has_break and not has_drawing:
            previous.getparent().remove(previous)
            removed += 1
    return removed


def remove_manual_page_break_before_heading(document: Document, heading: str) -> int:
    removed = 0
    for paragraph in document.paragraphs:
        if paragraph.text.strip() != heading:
            continue
        previous = paragraph._p.getprevious()
        if previous is None or previous.tag != qn("w:p"):
            continue
        if not previous.xpath('.//w:br[@w:type="page"]'):
            continue
        previous.getparent().remove(previous)
        removed += 1
    return removed


def main() -> None:
    OUTPUT.mkdir(parents=True, exist_ok=True)
    output = OUTPUT / "Supplementary_Information_scientific_maintenance_freeze.docx"
    result = documents.markdown_to_docx(
        SOURCE,
        output,
        body_size=10.5,
        double_space=False,
        line_numbers=False,
        running_header="Supplementary information",
        title_override="Supplementary information",
        supplementary_figure_dirs=[FIGURES],
        page_break_before_headings=set(),
    )

    document = Document(output)
    removed = remove_table_spacers_before_manual_page_breaks(document)
    removed_s8_break = remove_manual_page_break_before_heading(
        document, "Supplementary Table S8 | Full statistical-results archive map"
    )
    document.core_properties.author = "Zhi Chen; Teng Qi"
    document.core_properties.title = "Supplementary information"
    document.core_properties.subject = "Supplementary Table S4 reader-path micropass"
    document.core_properties.comments = (
        "Scientific-presentation maintenance freeze; S4/S4B reader numbering repaired without "
        f"scientific-value changes; removed {removed} redundant table spacers and "
        f"{removed_s8_break} renderer-divergent S8 page break."
    )
    document.save(output)

    document = Document(output)
    text = extract_text(document)
    source_text = SOURCE.read_text(encoding="utf-8")
    table_numbers = [
        int(number)
        for number in re.findall(r"(?m)^## Supplementary Table S(\d+) \|", source_text)
    ]
    figure_numbers = [
        int(number)
        for number in re.findall(r"(?m)^## Supplementary Figure S(\d+) \|", source_text)
    ]
    checks = {
        "title_exact": TITLE in text,
        "ten_figures_embedded": len(document.inline_shapes) == 10,
        "ten_grid_objects_retained": len(document.tables) == 10,
        "numbered_tables_are_exactly_s1_to_s9": table_numbers == list(range(1, 10)),
        "figures_are_exactly_s1_to_s10": figure_numbers == list(range(1, 11)),
        "s4_parent_heading_present": "Supplementary Table S4 | Regulator-sensitivity summaries"
        in text,
        "s4a_subheading_present": "a, Correlation-aware core-regulator sensitivity" in text,
        "s4b_subheading_present": "b, IFN-overlap-depletion summary" in text,
        "orphan_s4b_absent": "Supplementary Table S4B" not in text,
        "reader_wording_present": (
            "Sample-level composition in the 43-control/47-managed-SLE primary comparison"
            in text
        ),
        "implementation_wording_absent": "asserted 43/47 primary groups" not in text,
        "renderer_divergent_s8_page_break_removed": removed_s8_break == 1,
    }
    failed = [name for name, passed in checks.items() if not passed]
    status = {
        "created_at": datetime.now().astimezone().isoformat(timespec="seconds"),
        "status": (
            "PASS_SUPPLEMENTARY_TABLE_S4_DOCX_BUILT_DUAL_RENDER_REQUIRED"
            if not failed
            else "FAIL_SUPPLEMENTARY_TABLE_S4_DOCX_BUILD"
        ),
        "build_result": result,
        "checks": checks,
        "failed_checks": failed,
        "object_inventory": {
            "inline_shapes": len(document.inline_shapes),
            "grid_objects": len(document.tables),
            "numbered_supplementary_tables": len(table_numbers),
        },
        "file": {
            "path": output.relative_to(ROOT).as_posix(),
            "bytes": output.stat().st_size,
            "sha256": sha256(output),
        },
        "main_manuscript_rebuilt": False,
        "figures_rebuilt": False,
        "scientific_estimates_changed": False,
    }
    (RUN / "04_DOCUMENT_BUILD_STATUS.json").write_text(
        json.dumps(status, indent=2) + "\n", encoding="utf-8", newline="\n"
    )
    print(json.dumps(status, indent=2))
    if failed:
        raise RuntimeError(f"Document checks failed: {failed}")


if __name__ == "__main__":
    main()
