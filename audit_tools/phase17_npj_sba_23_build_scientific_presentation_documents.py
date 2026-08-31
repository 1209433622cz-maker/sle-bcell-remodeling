#!/usr/bin/env python3
"""Build DOCX files for the scientific-presentation freeze candidate."""

from __future__ import annotations

import hashlib
import json
import re
from datetime import datetime
from pathlib import Path

from docx import Document

import phase17_c8s_04_build_documents as documents


ROOT = Path(__file__).resolve().parents[1]
RUN = (
    ROOT
    / "phase17_v7/npj_sba_scientific_presentation_freeze/"
    "20260831_reader_path_and_legend_economy"
)
SOURCES = RUN / "sources"
FIGURES = RUN / "figures/figures"
OUTPUT = RUN / "documents"
TITLE = (
    "Disease-blind reconstruction distinguishes reproducible interferon remodeling from less stable "
    "B-cell state assignments in systemic lupus erythematosus"
)


def sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest().upper()


def extract_text(document: Document) -> str:
    values = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            values.extend(cell.text for cell in row.cells)
    return "\n".join(values)


def patch_properties(path: Path, subject: str) -> None:
    document = Document(path)
    document.core_properties.author = "Zhi Chen; Teng Qi"
    document.core_properties.title = TITLE if "Manuscript" in path.name else "Supplementary information"
    document.core_properties.subject = subject
    document.core_properties.comments = (
        "Scientific-presentation candidate regenerated from repository sources; "
        "the author-confirmed exact submission package remains unchanged."
    )
    document.save(path)


def main() -> None:
    OUTPUT.mkdir(parents=True, exist_ok=True)
    manuscript = OUTPUT / "Manuscript_scientific_presentation_freeze_candidate.docx"
    supplement = OUTPUT / "Supplementary_Information_scientific_presentation_freeze_candidate.docx"
    build_results = [
        documents.markdown_to_docx(
            SOURCES / "Manuscript_scientific_presentation_freeze_candidate.md",
            manuscript,
            body_size=12,
            double_space=True,
            line_numbers=True,
            running_header="npj Systems Biology and Applications | Article",
            title_override=TITLE,
        ),
        documents.markdown_to_docx(
            SOURCES / "Supplementary_Information_scientific_presentation_freeze_candidate.md",
            supplement,
            body_size=10.5,
            double_space=False,
            line_numbers=False,
            running_header="Supplementary information",
            title_override="Supplementary information",
            supplementary_figure_dirs=[FIGURES],
            page_break_before_headings=set(),
        ),
    ]
    patch_properties(manuscript, "Reader-path and legend-economy scientific manuscript candidate")
    patch_properties(supplement, "Scientific-presentation supplementary information candidate")

    manuscript_document = Document(manuscript)
    supplement_document = Document(supplement)
    manuscript_text = extract_text(manuscript_document)
    supplement_text = extract_text(supplement_document)
    abstract_match = re.search(r"Abstract\n(.+?)\nIntroduction", manuscript_text, flags=re.DOTALL)
    if not abstract_match:
        raise RuntimeError("Could not identify the abstract in the generated manuscript")
    abstract_words = re.findall(r"\b[\w'-]+\b", abstract_match.group(1))

    manuscript_checks = {
        "title_exact": TITLE in manuscript_text,
        "article_type_present": "Article type: Article" in manuscript_text,
        "abstract_145_words": len(abstract_words) == 145,
        "figure1_bounded_scaffold_title": (
            "Disease-blind reconstruction defines a bounded analysis scaffold" in manuscript_text
        ),
        "figure1_sample_cohort_units": all(
            token in manuscript_text
            for token in ("B_ASC composition", "B_CONV transcription", "sample-cohort stratum")
        ),
        "figure5_regulatory_context_title": (
            "Convergent observational evidence supports an IFN-centred regulatory context"
            in manuscript_text
        ),
        "figure5_parallel_roles": all(
            token in manuscript_text
            for token in ("confirmatory observational", "response-set concordance", "descriptive IFN-beta")
        ),
        "figure5_causal_boundary": all(
            token in manuscript_text
            for token in ("causal regulator", "direct binding", "unique upstream stimulus")
        ),
        "current_doi_present": "10.5281/zenodo.22151739" in manuscript_text,
        "old_doi_absent": "10.5281/zenodo.22086892" not in manuscript_text,
        "no_embedded_figures": len(manuscript_document.inline_shapes) == 0,
    }
    supplement_checks = {
        "ten_figures_embedded": len(supplement_document.inline_shapes) == 10,
        "title_synchronized": TITLE in supplement_text,
        "corrected_calibration_failure_present": (
            "No new multiplicity family was evaluated after corrected calibration failed."
            in supplement_text
        ),
        "calibration_hold_absent": "calibration HOLD" not in supplement_text,
        "formal_hold_absent": "formal HOLD" not in supplement_text,
        "internal_c9_pass_absent": "C9 PASS" not in supplement_text,
        "current_doi_present": "10.5281/zenodo.22151739" in supplement_text,
        "old_doi_absent": "10.5281/zenodo.22086892" not in supplement_text,
    }
    checks = {"manuscript": manuscript_checks, "supplement": supplement_checks}
    failed = [
        f"{group}.{name}"
        for group, group_checks in checks.items()
        for name, passed in group_checks.items()
        if not passed
    ]
    if failed:
        raise RuntimeError(f"Scientific-presentation document checks failed: {failed}")

    status = {
        "created_at": datetime.now().astimezone().isoformat(timespec="seconds"),
        "status": "PASS_SCIENTIFIC_PRESENTATION_DOCX_BUILT_DUAL_RENDER_REQUIRED",
        "build_results": build_results,
        "object_inventory": {
            "manuscript_inline_shapes": len(manuscript_document.inline_shapes),
            "supplement_inline_shapes": len(supplement_document.inline_shapes),
            "supplement_expected_figures": 10,
            "supplement_object_count_matches": len(supplement_document.inline_shapes) == 10,
        },
        "files": {
            path.name: {"bytes": path.stat().st_size, "sha256": sha256(path)}
            for path in (manuscript, supplement)
        },
        "checks": checks,
        "scientific_estimates_changed": False,
        "source_data_changed": False,
        "exact_submission_package_modified": False,
    }
    (RUN / "02_DOCUMENT_BUILD_STATUS.json").write_text(
        json.dumps(status, indent=2) + "\n", encoding="utf-8", newline="\n"
    )
    print(
        json.dumps(
            {
                "status": status["status"],
                "abstract_words": len(abstract_words),
                "object_inventory": status["object_inventory"],
                "files": status["files"],
            },
            indent=2,
        )
    )


if __name__ == "__main__":
    main()
