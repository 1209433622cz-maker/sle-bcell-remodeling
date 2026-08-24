#!/usr/bin/env python3
"""Qualify the portable Gate C8BR release runtime before artifact creation."""

from __future__ import annotations

import argparse
import json
import platform
import shutil
import sys
from importlib.metadata import version
from pathlib import Path

import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
from PIL import Image
from docx import Document
from pypdf import PdfReader


EXPECTED = {
    "matplotlib": "3.10.7",
    "numpy": "2.3.3",
    "pandas": "2.3.3",
    "pillow": "12.3.0",
    "python-docx": "1.2.0",
    "pypdf": "6.10.0",
}


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser()
    parser.add_argument("--output-dir", type=Path, required=True)
    return parser.parse_args()


def main() -> None:
    args = parse_args()
    output_dir = args.output_dir.resolve()
    if output_dir.exists():
        shutil.rmtree(output_dir)
    output_dir.mkdir(parents=True)

    actual = {name: version(name) for name in EXPECTED}
    versions_match = actual == EXPECTED

    x = np.linspace(0, 2 * np.pi, 200)
    frame = pd.DataFrame({"x": x, "signal": np.sin(x)})
    figure, axis = plt.subplots(figsize=(4.0, 2.5), constrained_layout=True)
    axis.plot(frame["x"], frame["signal"], color="#2C6EAD", linewidth=1.2)
    axis.set(title="Gate C8BR savefig qualification", xlabel="x", ylabel="sin(x)")
    png = output_dir / "savefig_smoke.png"
    pdf = output_dir / "savefig_smoke.pdf"
    figure.savefig(png, dpi=150)
    figure.savefig(pdf)
    plt.close(figure)

    with Image.open(png) as image:
        png_size = image.size
        png_ok = image.width >= 500 and image.height >= 300
    pdf_ok = len(PdfReader(pdf).pages) == 1 and pdf.stat().st_size > 1_000

    docx = output_dir / "docx_smoke.docx"
    document = Document()
    document.add_heading("Gate C8BR document qualification", level=1)
    document.add_paragraph("Portable release environment smoke test.")
    document.save(docx)
    docx_ok = len(Document(docx).paragraphs) == 2 and docx.stat().st_size > 10_000

    passed = versions_match and png_ok and pdf_ok and docx_ok
    result = {
        "created_at": "2026-08-25",
        "status": "PASS_GATE_C8BR_RELEASE_RUNTIME" if passed else "HOLD_GATE_C8BR_RELEASE_RUNTIME",
        "python_executable": sys.executable,
        "python_version": platform.python_version(),
        "platform": platform.platform(),
        "expected_versions": EXPECTED,
        "actual_versions": actual,
        "versions_match": versions_match,
        "png": {"bytes": png.stat().st_size, "pixels": png_size, "pass": png_ok},
        "pdf": {"bytes": pdf.stat().st_size, "pages": 1, "pass": pdf_ok},
        "docx": {"bytes": docx.stat().st_size, "paragraphs": 2, "pass": docx_ok},
    }
    (output_dir.parent / "00_GATE_C8BR_RELEASE_RUNTIME_STATUS.json").write_text(
        json.dumps(result, indent=2) + "\n", encoding="utf-8", newline="\n"
    )
    print(json.dumps(result, indent=2))
    if not passed:
        raise SystemExit(1)


if __name__ == "__main__":
    main()
