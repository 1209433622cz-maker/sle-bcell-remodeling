#!/usr/bin/env python3
"""Build the full-main-figure candidate manuscript without touching the package."""

from __future__ import annotations

import hashlib
import json
from datetime import datetime
from pathlib import Path

from docx import Document

import phase17_c8s_04_build_documents as documents


ROOT = Path(__file__).resolve().parents[1]
RUN_DIR = (
    ROOT
    / "phase17_v7/npj_sba_full_main_figure_refinement/20260831_figure5e_and_figures2to4_adjudication"
)
RECOMMENDED = RUN_DIR / "recommended_full_main_figure_set"
SOURCE = RECOMMENDED / "sources/Manuscript_full_main_figure_candidate.md"
OUTPUT_DIR = RECOMMENDED / "documents"
OUTPUT = OUTPUT_DIR / "Manuscript_scientific_candidate.docx"
TITLE = (
    "Disease-blind reconstruction distinguishes reproducible interferon remodeling from unstable B-cell state assignments "
    "in systemic lupus erythematosus"
)


def sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest().upper()


def docx_text(path: Path) -> str:
    document = Document(path)
    values = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            values.extend(cell.text for cell in row.cells)
    return "\n".join(values)


def main() -> None:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    result = documents.markdown_to_docx(
        SOURCE,
        OUTPUT,
        body_size=12,
        double_space=True,
        line_numbers=True,
        running_header="npj Systems Biology and Applications | Article",
        title_override=TITLE,
    )
    document = Document(OUTPUT)
    document.core_properties.author = "Zhi Chen; Teng Qi"
    document.core_properties.title = TITLE
    document.core_properties.subject = "Scientific presentation candidate guided by npj Systems Biology and Applications"
    document.core_properties.comments = "Generated reproducibly from the full-main-figure scientific candidate Markdown source."
    document.save(OUTPUT)

    text = docx_text(OUTPUT)
    checks = {
        "title_exact": TITLE in text,
        "article_type_present": "Article type: Article" in text,
        "new_figure1a_legend": "Disease-blind workflow and identity scope" in text,
        "new_figure5a_legend": "Quantitative summary of three evidence classes" in text,
        "new_figure5e_legend": "all 24 donor-gene effects were positive" in text,
        "fine_state_boundary": "hard fine-state assignments" in text,
        "causal_boundary": "not a causal regulator, direct binding or a uniquely upstream ligand" in text,
        "ai_disclosure_present": "Generative AI assistance" in text,
        "current_doi_present": "10.5281/zenodo.22151739" in text,
        "no_old_doi": "10.5281/zenodo.22086892" not in text,
        "no_embedded_figures": len(Document(OUTPUT).inline_shapes) == 0,
    }
    if not all(checks.values()):
        raise RuntimeError(f"Candidate manuscript checks failed: {[name for name, passed in checks.items() if not passed]}")
    status = {
        "created_at": datetime.now().astimezone().isoformat(timespec="seconds"),
        "status": "PASS_CANDIDATE_MANUSCRIPT_DOCX_BUILT_RENDER_QA_REQUIRED",
        "source": SOURCE.relative_to(ROOT).as_posix(),
        "source_sha256": sha256(SOURCE),
        "output": OUTPUT.relative_to(ROOT).as_posix(),
        "output_sha256": sha256(OUTPUT),
        "output_bytes": OUTPUT.stat().st_size,
        "builder_result": result,
        "checks": checks,
        "scientific_estimates_changed": False,
        "exact_submission_package_modified": False,
    }
    (RECOMMENDED / "02_CANDIDATE_MANUSCRIPT_BUILD_STATUS.json").write_text(
        json.dumps(status, indent=2) + "\n", encoding="utf-8", newline="\n"
    )
    print(json.dumps(status, indent=2))


if __name__ == "__main__":
    main()
