"""Independently audit the built journal-neutral submission package."""

import argparse
from datetime import datetime
import json
from pathlib import Path
import zipfile

from docx import Document
from pypdf import PdfReader

from phase17_postc9_23_build_current_submission_package import (
    C9R_HOLD,
    PACKAGE_DIR,
    PACKAGE_ZIP,
    RECEIPT as BUILD_RECEIPT,
    R1_HOLD,
    ROOT,
    checksum,
    require,
    verify_manifest,
)


WORK = ROOT / "00_project_management/current_submission_package_2026-08-30"
OUTPUT = WORK / "package_validation.json"


def read_json(path):
    return json.loads(path.read_text(encoding="utf-8-sig"))


def main():
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--expected-zip-sha256", required=True)
    parser.add_argument("--output", type=Path, default=OUTPUT)
    args = parser.parse_args()
    output = args.output.resolve()
    require(output.is_relative_to((ROOT / "00_project_management").resolve()), "Receipt must remain in project management")

    manifest_files = verify_manifest(PACKAGE_DIR)
    build = read_json(BUILD_RECEIPT)
    metadata = read_json(PACKAGE_DIR / "07_Integrity/PACKAGE_METADATA.json")
    freeze = read_json(WORK / "scientific_freeze_reverification.json")
    a11y = read_json(WORK / "cover_letter_a11y.json")
    zip_sha = checksum(PACKAGE_ZIP)
    require(zip_sha == args.expected_zip_sha256.upper(), "Package ZIP differs from the repeated-build SHA-256")
    require(build["package_zip_sha256"] == zip_sha, "Build receipt ZIP SHA-256 differs")
    require(build["status"] == "PASS_JOURNAL_NEUTRAL_PACKAGE_VERIFIED_NOT_SUBMISSION_AUTHORIZED", "Build status differs")

    with zipfile.ZipFile(PACKAGE_ZIP) as archive:
        require(archive.testzip() is None, "Outer package ZIP CRC verification failed")
        zip_entries = len(archive.infolist())
    nested = {}
    for path in sorted((PACKAGE_DIR / "05_Source_Data").glob("*.zip")):
        with zipfile.ZipFile(path) as archive:
            require(archive.testzip() is None, "Nested ZIP CRC verification failed: " + path.name)
            nested[path.name] = len(archive.infolist())
    require(nested == {
        "Figure_Source_Data.zip": 16,
        "Full_Statistical_Results.zip": 185,
        "Regulator_Sensitivity.zip": 11,
    }, "Nested archive inventory differs")

    pdf_pages = {
        "Manuscript.pdf": len(PdfReader(PACKAGE_DIR / "01_Manuscript/Manuscript.pdf").pages),
        "Supplementary_Information.pdf": len(PdfReader(PACKAGE_DIR / "02_Supplementary_Information/Supplementary_Information.pdf").pages),
        "Cover_Letter_Draft.pdf": len(PdfReader(PACKAGE_DIR / "06_Administrative/Cover_Letter_Draft.pdf").pages),
    }
    require(pdf_pages == {
        "Manuscript.pdf": 18,
        "Supplementary_Information.pdf": 19,
        "Cover_Letter_Draft.pdf": 1,
    }, "Primary PDF page counts differ")
    figure_pdfs = sorted((PACKAGE_DIR / "03_Main_Figures").glob("*.pdf"))
    supplementary_pdfs = sorted((PACKAGE_DIR / "04_Supplementary_Figures").glob("*.pdf"))
    require(len(figure_pdfs) == 5 and all(len(PdfReader(path).pages) == 1 for path in figure_pdfs), "Main-figure PDF inventory differs")
    require(len(supplementary_pdfs) == 10 and all(len(PdfReader(path).pages) == 1 for path in supplementary_pdfs), "Supplementary-figure PDF inventory differs")

    documents = {
        "Manuscript.docx": Document(PACKAGE_DIR / "01_Manuscript/Manuscript.docx"),
        "Supplementary_Information.docx": Document(PACKAGE_DIR / "02_Supplementary_Information/Supplementary_Information.docx"),
        "Cover_Letter_Draft.docx": Document(PACKAGE_DIR / "06_Administrative/Cover_Letter_Draft.docx"),
    }
    require(all(len(document.paragraphs) > 0 for document in documents.values()), "A DOCX has no readable paragraphs")
    require(a11y["counts"] == {"high": 0, "medium": 0, "low": 0}, "Cover-letter accessibility findings remain")
    require(freeze["status"] == "PASS_FREEZE_INTEGRITY_NOT_SCIENTIFIC_GATE_PASS", "Scientific freeze reverification differs")
    require(freeze["R1_decision"] == R1_HOLD and freeze["C9R_decision"] == C9R_HOLD, "Scientific HOLD boundary differs")

    for value in (build, metadata):
        require(value["selected_target"] is None, "Target journal is selected prematurely")
        for key in (
            "jcr_q1_verified",
            "institutional_apc_coverage_verified",
            "exact_package_author_approved",
            "submission_authorized",
            "apc_commitment_authorized",
            "corrected_external_outcome_unlock_authorized",
        ):
            require(value[key] is False, "Governance boundary is overclaimed: " + key)
        require(value["R1_decision"] == R1_HOLD and value["C9R_decision"] == C9R_HOLD, "Package HOLD boundary differs")

    result = {
        "audited_at": datetime.now().astimezone().isoformat(timespec="seconds"),
        "status": "PASS_CURRENT_SUBMISSION_PACKAGE_INTEGRITY_NOT_SUBMISSION_AUTHORIZATION",
        "package_zip_bytes": PACKAGE_ZIP.stat().st_size,
        "package_zip_sha256": zip_sha,
        "deterministic_rebuilds_observed": 2,
        "deterministic_rebuild_sha256_match": True,
        "zip_entries": zip_entries,
        "manifest_files_verified": manifest_files,
        "scientific_files_copied_byte_for_byte": build["scientific_files_copied_byte_for_byte"],
        "pdf_pages": pdf_pages,
        "main_figure_pdfs": 5,
        "supplementary_figure_pdfs": 10,
        "nested_zip_entries": nested,
        "docx_files_opened": len(documents),
        "cover_letter_renderer": "WPS Office",
        "cover_letter_pages_visually_checked": 1,
        "cover_letter_accessibility": a11y["counts"],
        "scientific_freeze_reverified": True,
        "R1_decision": R1_HOLD,
        "C9R_decision": C9R_HOLD,
        "selected_target": None,
        "jcr_q1_verified": False,
        "institutional_apc_coverage_verified": False,
        "exact_package_author_approved": False,
        "submission_authorized": False,
        "apc_commitment_authorized": False,
        "next_gate": "Archive official JCR/APC evidence, freeze the target, adapt once, rebuild and obtain exact-file author approval",
    }
    output.parent.mkdir(parents=True, exist_ok=True)
    output.write_text(json.dumps(result, indent=2) + "\n", encoding="utf-8")
    print(json.dumps(result, indent=2))


if __name__ == "__main__":
    main()
