#!/usr/bin/env python3
"""Create the Gate C8S manuscript, supplement and cover-letter DOCX files."""

from __future__ import annotations

import hashlib
import json
import re
import shutil
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
FIXED_ZIP_TIME = (2026, 8, 21, 0, 0, 0)
RUN_DIR = ROOT / "phase17_v7" / "gateC8S" / "20260821_supplementary_traceability_freeze"
REFERENCE_SOURCE = ROOT / "phase17_v7" / "gateC8R" / "20260820_pre_submission_repair" / "references"
SENSITIVITY_SOURCE = ROOT / "phase17_v7" / "gateC8R" / "20260820_pre_submission_repair"
PACKAGE = ROOT / "04_submission" / "package_genome_medicine_gateC8S_2026-08-21"
MANUSCRIPT_MD = ROOT / "01_manuscript" / "manuscript_v12_genome_medicine_gateC8S_2026-08-21.md"
SUPPLEMENT_MD = ROOT / "01_manuscript" / "supplementary_information_v3_gateC8S_2026-08-21.md"
COVER_MD = ROOT / "04_submission" / "cover_letter_genome_medicine_gateC8S_AUTHOR_COMPLETION_REQUIRED_2026-08-21.md"
FIGURE_SOURCE = RUN_DIR
INK = RGBColor(0x1F, 0x1F, 0x1F)
MUTED = RGBColor(0x59, 0x59, 0x59)
PLACEHOLDER = RGBColor(0x9B, 0x1C, 0x1C)
TABLE_FILL = "E8EEF5"
CONTENT_DXA = 9360
TABLE_INDENT_DXA = 120


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: list[int]) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.insert(0, tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    margins = tbl_pr.find(qn("w:tblCellMar"))
    if margins is None:
        margins = OxmlElement("w:tblCellMar")
        tbl_pr.append(margins)
    for side, value in (("top", 80), ("bottom", 80), ("start", 120), ("end", 120)):
        node = margins.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths[idx])


def repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    run.font.name = "Times New Roman"
    run.font.size = Pt(9)
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr)
    run._r.append(fld_char2)


def set_line_numbers(section) -> None:
    sect_pr = section._sectPr
    node = sect_pr.find(qn("w:lnNumType"))
    if node is None:
        node = OxmlElement("w:lnNumType")
        sect_pr.append(node)
    node.set(qn("w:countBy"), "1")
    node.set(qn("w:distance"), "360")
    node.set(qn("w:restart"), "continuous")


def add_numbering(document: Document, fmt: str, text: str, left: int, hanging: int) -> int:
    numbering = document.part.numbering_part.element
    abstract_ids = [int(x.get(qn("w:abstractNumId"))) for x in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(x.get(qn("w:numId"))) for x in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    level.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), fmt)
    level.append(num_fmt)
    level_text = OxmlElement("w:lvlText")
    level_text.set(qn("w:val"), text)
    level.append(level_text)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), str(left))
    tabs.append(tab)
    p_pr.append(tabs)
    indent = OxmlElement("w:ind")
    indent.set(qn("w:left"), str(left))
    indent.set(qn("w:hanging"), str(hanging))
    p_pr.append(indent)
    level.append(p_pr)
    abstract.append(level)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def apply_numbering(paragraph, num_id: int) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num)


def set_run_font(run, size: float | None = None, bold: bool | None = None, italic: bool | None = None, color: RGBColor | None = None, font: str = "Times New Roman") -> None:
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = color


INLINE_PATTERN = re.compile(r"(\[\[(?:.|\n)*?\]\]|\*\*.*?\*\*|`.*?`)")


def add_inline(paragraph, text: str, size: float | None = None) -> None:
    for piece in INLINE_PATTERN.split(text):
        if not piece:
            continue
        if piece.startswith("[[") and piece.endswith("]]" ):
            run = paragraph.add_run(piece)
            set_run_font(run, size=size, bold=True, color=PLACEHOLDER)
        elif piece.startswith("**") and piece.endswith("**"):
            run = paragraph.add_run(piece[2:-2])
            set_run_font(run, size=size, bold=True, color=INK)
        elif piece.startswith("`") and piece.endswith("`"):
            run = paragraph.add_run(piece[1:-1])
            set_run_font(run, size=(size or 12) - 0.5, font="Courier New", color=INK)
        else:
            run = paragraph.add_run(piece)
            set_run_font(run, size=size, color=INK)


def configure_styles(document: Document, body_size: float, double_space: bool) -> None:
    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(body_size)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0 if double_space else 6)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE if double_space else WD_LINE_SPACING.MULTIPLE
    if not double_space:
        normal.paragraph_format.line_spacing = 1.15

    style_specs = {
        "Title": (16, True, 0, 12, WD_ALIGN_PARAGRAPH.CENTER),
        "Heading 1": (14, True, 12, 6, WD_ALIGN_PARAGRAPH.LEFT),
        "Heading 2": (12, True, 10, 4, WD_ALIGN_PARAGRAPH.LEFT),
        "Heading 3": (11, True, 8, 3, WD_ALIGN_PARAGRAPH.LEFT),
    }
    for name, (size, bold, before, after, align) in style_specs.items():
        style = styles[name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.size = Pt(size)
        style.font.bold = bold
        style.font.color.rgb = INK
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.0
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.alignment = align
        p_pr = style._element.get_or_add_pPr()
        p_bdr = p_pr.find(qn("w:pBdr"))
        if p_bdr is not None:
            p_pr.remove(p_bdr)

    if "Reference" not in [s.name for s in styles]:
        ref = styles.add_style("Reference", WD_STYLE_TYPE.PARAGRAPH)
    else:
        ref = styles["Reference"]
    ref.font.name = "Times New Roman"
    ref._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    ref.font.size = Pt(body_size)
    ref.paragraph_format.space_after = Pt(3 if not double_space else 0)
    ref.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE if double_space else WD_LINE_SPACING.MULTIPLE
    if not double_space:
        ref.paragraph_format.line_spacing = 1.15


def configure_document(document: Document, body_size: float, double_space: bool, line_numbers: bool, running_header: str | None) -> None:
    configure_styles(document, body_size, double_space)
    # WPS treats an absent even-page header as empty even when odd/even headers
    # are nominally disabled, so define matching odd and even parts explicitly.
    document.settings.odd_and_even_pages_header_footer = True
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = False
    if line_numbers:
        set_line_numbers(section)
    if running_header:
        for header in (section.header, section.even_page_header):
            paragraph = header.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = paragraph.add_run(running_header)
            set_run_font(run, size=9, color=MUTED)
            p_pr = paragraph._p.get_or_add_pPr()
            p_bdr = OxmlElement("w:pBdr")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"), "single")
            bottom.set(qn("w:sz"), "4")
            bottom.set(qn("w:space"), "2")
            bottom.set(qn("w:color"), "BFBFBF")
            p_bdr.append(bottom)
            p_pr.append(p_bdr)
    for footer in (section.footer, section.even_page_footer):
        add_page_number(footer.paragraphs[0])


def paragraph_blocks(markdown: str) -> list[str]:
    lines = markdown.replace("\r\n", "\n").split("\n")
    blocks: list[str] = []
    idx = 0
    while idx < len(lines):
        line = lines[idx].rstrip()
        if not line.strip():
            idx += 1
            continue
        if line.startswith("|"):
            table_lines = []
            while idx < len(lines) and lines[idx].lstrip().startswith("|"):
                table_lines.append(lines[idx].strip())
                idx += 1
            blocks.append("\n".join(table_lines))
            continue
        if line.startswith("#") or re.match(r"^(?:- |\d+\. )", line):
            blocks.append(line.strip())
            idx += 1
            continue
        paragraph = [line.strip()]
        idx += 1
        while idx < len(lines):
            candidate = lines[idx].rstrip()
            if not candidate.strip() or candidate.startswith("#") or candidate.startswith("|") or re.match(r"^(?:- |\d+\. )", candidate):
                break
            paragraph.append(candidate.strip())
            idx += 1
        blocks.append(" ".join(paragraph))
    return blocks


def table_widths(rows: list[list[str]]) -> list[int]:
    cols = len(rows[0])
    if cols == 6 and rows[0][:3] == ["Contrast", "Regulator", "Matched targets"]:
        return [2400, 1100, 1350, 2100, 1250, 1160]
    maxima = []
    for col in range(cols):
        maxima.append(max(10, min(60, max(len(row[col]) for row in rows))))
    total = sum(maxima)
    raw = [max(1100, round(CONTENT_DXA * value / total)) for value in maxima]
    scale = CONTENT_DXA / sum(raw)
    widths = [round(value * scale) for value in raw]
    widths[-1] += CONTENT_DXA - sum(widths)
    return widths


def add_markdown_table(document: Document, block: str, body_size: float) -> None:
    raw_rows = [[cell.strip() for cell in line.strip().strip("|").split("|")] for line in block.splitlines()]
    if len(raw_rows) > 1 and all(re.fullmatch(r":?-{3,}:?", cell) for cell in raw_rows[1]):
        raw_rows.pop(1)
    table = document.add_table(rows=len(raw_rows), cols=len(raw_rows[0]))
    table.style = "Table Grid"
    for r_idx, row in enumerate(raw_rows):
        for c_idx, value in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.0
            add_inline(paragraph, value, size=8.5 if len(raw_rows[0]) >= 6 else body_size - 1)
            if r_idx == 0:
                set_cell_shading(cell, TABLE_FILL)
                for run in paragraph.runs:
                    run.bold = True
    repeat_table_header(table.rows[0])
    set_table_geometry(table, table_widths(raw_rows))
    spacer = document.add_paragraph()
    spacer.paragraph_format.space_after = Pt(3)


def markdown_to_docx(source: Path, output: Path, *, body_size: float, double_space: bool, line_numbers: bool, running_header: str | None, title_override: str | None = None, compact: bool = False) -> dict:
    document = Document()
    configure_document(document, body_size, double_space, line_numbers, running_header)
    if compact:
        normal = document.styles["Normal"]
        normal.paragraph_format.space_after = Pt(3)
        normal.paragraph_format.line_spacing = 1.05
        title_style = document.styles["Title"]
        title_style.paragraph_format.space_after = Pt(8)
        section = document.sections[0]
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
    document.core_properties.author = "Zhi Chen; Teng Qi"
    document.core_properties.title = title_override or source.stem
    document.core_properties.subject = "Genome Medicine Research submission"
    document.core_properties.keywords = "SLE; B cells; single-cell RNA sequencing; interferon"
    document.core_properties.comments = "Generated reproducibly from Gate C8S Markdown sources."

    ref_num_id = add_numbering(document, "decimal", "%1.", 720, 360)
    bullet_num_id = add_numbering(document, "bullet", "•", 540, 270)
    blocks = paragraph_blocks(source.read_text(encoding="utf-8-sig"))
    reference_mode = False
    table_count = 0
    placeholders = 0
    for block in blocks:
        placeholders += block.count("[[")
        figure_marker = re.fullmatch(r"\[\[SUPPLEMENTARY_FIGURE:(S[1-7])\]\]", block)
        if figure_marker:
            placeholders -= 1
            figure_id = figure_marker.group(1)
            matches = sorted((RUN_DIR / "supplementary_figures").glob(f"Supplementary_Figure_{figure_id}_*.png"))
            if len(matches) != 1:
                raise RuntimeError(f"Expected one PNG for {figure_id}; found {len(matches)}")
            paragraph = document.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            shape = paragraph.add_run().add_picture(str(matches[0]), width=Inches(6.35))
            shape._inline.docPr.set("title", f"Supplementary Figure {figure_id}")
            shape._inline.docPr.set("descr", f"Supplementary Figure {figure_id}; full panel description appears immediately above the figure.")
            continue
        if block.startswith("|"):
            add_markdown_table(document, block, body_size)
            table_count += 1
            continue
        if block.startswith("# "):
            paragraph = document.add_paragraph(style="Title")
            add_inline(paragraph, block[2:].strip(), size=16)
            continue
        if block.startswith("## "):
            heading = block[3:].strip()
            if block.startswith("## Supplementary Figure S") or block.startswith("## Supplementary Table S8"):
                document.add_page_break()
            paragraph = document.add_paragraph(style="Heading 1")
            add_inline(paragraph, heading, size=14)
            reference_mode = heading == "References"
            continue
        if block.startswith("### "):
            paragraph = document.add_paragraph(style="Heading 2")
            add_inline(paragraph, block[4:].strip(), size=12)
            continue
        numbered = re.match(r"^(\d+)\.\s+(.*)$", block)
        if numbered and reference_mode:
            paragraph = document.add_paragraph(style="Reference")
            apply_numbering(paragraph, ref_num_id)
            add_inline(paragraph, numbered.group(2), size=body_size)
            continue
        bullet = re.match(r"^-\s+(.*)$", block)
        if bullet:
            paragraph = document.add_paragraph()
            apply_numbering(paragraph, bullet_num_id)
            add_inline(paragraph, bullet.group(1), size=body_size)
            continue
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.widow_control = True
        add_inline(paragraph, block, size=body_size)

    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)
    return {
        "source": str(source.relative_to(ROOT)).replace("\\", "/"),
        "output": str(output.relative_to(ROOT)).replace("\\", "/"),
        "paragraphs": len(document.paragraphs),
        "tables": table_count,
        "inline_shapes": len(document.inline_shapes),
        "placeholders": placeholders,
        "bytes": output.stat().st_size,
    }


def sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest().upper()


def write_deterministic_zip(output: Path, files: list[Path]) -> None:
    """Write stable attachment bytes for an unchanged set of source files."""
    with zipfile.ZipFile(output, "w", compression=zipfile.ZIP_DEFLATED, compresslevel=9) as archive:
        for path in sorted(files, key=lambda item: item.name):
            info = zipfile.ZipInfo(path.name, date_time=FIXED_ZIP_TIME)
            info.compress_type = zipfile.ZIP_DEFLATED
            info.external_attr = 0o100644 << 16
            info.create_system = 3
            archive.writestr(info, path.read_bytes(), compress_type=zipfile.ZIP_DEFLATED, compresslevel=9)


def prepare_package_assets() -> dict:
    figures_dir = PACKAGE / "figures"
    supplementary_figures_dir = PACKAGE / "figures_supplementary"
    source_dir = PACKAGE / "additional_files" / "source_data"
    refs_dir = PACKAGE / "references"
    docs_dir = PACKAGE / "submission_docs"
    sensitivity_dir = PACKAGE / "additional_files" / "regulator_sensitivity"
    for directory in (figures_dir, supplementary_figures_dir, source_dir, refs_dir, docs_dir, sensitivity_dir):
        directory.mkdir(parents=True, exist_ok=True)

    figure_files = []
    for path in sorted((FIGURE_SOURCE / "figures").glob("Figure*.*")):
        target = figures_dir / path.name
        shutil.copy2(path, target)
        figure_files.append(target)
    supplementary_figure_files = []
    for path in sorted((FIGURE_SOURCE / "supplementary_figures").glob("Supplementary_Figure_S*.*")):
        target = supplementary_figures_dir / path.name
        shutil.copy2(path, target)
        supplementary_figure_files.append(target)
    source_files = []
    for path in sorted((FIGURE_SOURCE / "source_data").glob("Figure*_source_data.csv")):
        target = source_dir / path.name
        shutil.copy2(path, target)
        source_files.append(target)
    for path in sorted((FIGURE_SOURCE / "supplementary_source_data").glob("Supplementary_Figure_S*_source_data.csv")):
        target = source_dir / path.name
        shutil.copy2(path, target)
        source_files.append(target)

    manifest = source_dir / "SHA256SUMS.csv"
    manifest.write_text(
        "file,bytes,sha256\n" + "\n".join(f"{p.name},{p.stat().st_size},{sha256(p)}" for p in source_files) + "\n",
        encoding="utf-8",
    )
    zip_path = PACKAGE / "additional_files" / "Additional_file_2_Figure_Source_Data_GateC8S.zip"
    write_deterministic_zip(zip_path, source_files + [manifest])

    sensitivity_files = []
    for name in (
        "03_CORRELATION_AWARE_STAT1_STAT2_SENSITIVITY.csv",
        "04_CORRELATION_AWARE_STAT1_STAT2_DECISION.json",
    ):
        source = SENSITIVITY_SOURCE / name
        target = sensitivity_dir / name
        shutil.copy2(source, target)
        sensitivity_files.append(target)
    sensitivity_manifest = sensitivity_dir / "SHA256SUMS.csv"
    sensitivity_manifest.write_text(
        "file,bytes,sha256\n"
        + "\n".join(f"{p.name},{p.stat().st_size},{sha256(p)}" for p in sensitivity_files)
        + "\n",
        encoding="utf-8",
    )
    sensitivity_zip = PACKAGE / "additional_files" / "Additional_file_3_Regulator_Sensitivity_GateC8S.zip"
    write_deterministic_zip(sensitivity_zip, sensitivity_files + [sensitivity_manifest])

    full_stats_source = RUN_DIR / "Additional_file_4_Full_Statistical_Results_GateC8S.zip"
    full_stats_target = PACKAGE / "additional_files" / full_stats_source.name
    shutil.copy2(full_stats_source, full_stats_target)

    for path in sorted(REFERENCE_SOURCE.glob("*")):
        if path.is_file():
            shutil.copy2(path, refs_dir / path.name)
    for source in (
        ROOT / "04_submission" / "author_completion_form_gateC8S_2026-08-21.md",
        ROOT / "04_submission" / "journal_target_decision_gateC8S_2026-08-21.md",
        ROOT / "04_submission" / "reporting_checklist_gateC8S_2026-08-21.md",
        ROOT / "04_submission" / "cover_letter_genome_medicine_gateC8S_AUTHOR_COMPLETION_REQUIRED_2026-08-21.md",
    ):
        shutil.copy2(source, docs_dir / source.name)

    return {
        "figure_files": len(figure_files),
        "supplementary_figure_files": len(supplementary_figure_files),
        "source_files": len(source_files),
        "source_zip": str(zip_path.relative_to(ROOT)).replace("\\", "/"),
        "source_zip_bytes": zip_path.stat().st_size,
        "sensitivity_files": len(sensitivity_files),
        "sensitivity_zip": str(sensitivity_zip.relative_to(ROOT)).replace("\\", "/"),
        "sensitivity_zip_bytes": sensitivity_zip.stat().st_size,
        "full_statistical_zip": str(full_stats_target.relative_to(ROOT)).replace("\\", "/"),
        "full_statistical_zip_bytes": full_stats_target.stat().st_size,
        "full_statistical_zip_sha256": sha256(full_stats_target),
    }


def main() -> None:
    if PACKAGE.exists():
        # Rebuild only the active Gate C8S package; no legacy directory is touched.
        shutil.rmtree(PACKAGE)
    PACKAGE.mkdir(parents=True, exist_ok=True)

    outputs = []
    outputs.append(
        markdown_to_docx(
            MANUSCRIPT_MD,
            PACKAGE / "main_text" / "Genome_Medicine_Manuscript_GateC8S_AUTHOR_COMPLETION_REQUIRED.docx",
            body_size=12,
            double_space=True,
            line_numbers=True,
            running_header="Genome Medicine | Research | Gate C8S",
            title_override="Genome Medicine manuscript Gate C8S",
        )
    )
    outputs.append(
        markdown_to_docx(
            SUPPLEMENT_MD,
            PACKAGE / "additional_files" / "Additional_file_1_Supplementary_Information_GateC8S.docx",
            body_size=11,
            double_space=False,
            line_numbers=False,
            running_header="Supplementary information | Gate C8S",
            title_override="Supplementary information Gate C8S",
        )
    )
    outputs.append(
        markdown_to_docx(
            COVER_MD,
            PACKAGE / "submission_docs" / "Genome_Medicine_Cover_Letter_GateC8S_AUTHOR_CONFIRMATION_REQUIRED.docx",
            body_size=10.5,
            double_space=False,
            line_numbers=False,
            running_header=None,
            title_override="Genome Medicine cover letter Gate C8S",
            compact=True,
        )
    )
    assets = prepare_package_assets()
    status = {
        "created_at": "2026-08-21",
        "design_preset": "standard_business_brief",
        "named_override": "Genome Medicine scientific submission: Times New Roman, black hierarchy, manuscript double spacing and continuous line numbering",
        "outputs": outputs,
        "assets": assets,
    }
    RUN_DIR.mkdir(parents=True, exist_ok=True)
    (RUN_DIR / "07_GATE_C8S_DOCUMENT_BUILD_STATUS.json").write_text(json.dumps(status, indent=2) + "\n", encoding="utf-8", newline="\n")
    print(json.dumps(status, indent=2))


if __name__ == "__main__":
    main()
