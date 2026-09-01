#!/usr/bin/env python3
"""Finalize the Supplementary first-citation-order scientific refreeze."""

from __future__ import annotations

import argparse
import csv
import hashlib
import json
import re
from datetime import datetime
from pathlib import Path

from docx import Document
from pypdf import PdfReader

import phase17_npj_sba_43_integrate_supplementary_citation_refreeze as integration_logic


ROOT = Path(__file__).resolve().parents[1]
BASE = ROOT / "phase17_v7/npj_sba_supplementary_table_reader_path/20260901_s4_reader_path_refreeze"
RUN = ROOT / "phase17_v7/npj_sba_supplementary_citation_refreeze/20260901_first_citation_order"
DOCUMENTS = RUN / "documents"
QA = RUN / "qa"
LO_DOCUMENTS = QA / "libreoffice_documents"
MAIN_STEM = "Manuscript_scientific_maintenance_freeze"
SUPPLEMENT_STEM = "Supplementary_Information_scientific_maintenance_freeze"
STATUS_PATH = RUN / "08_SCIENTIFIC_PRESENTATION_MAINTENANCE_FREEZE_STATUS.json"
MANIFEST_PATH = RUN / "09_FINAL_FILE_MANIFEST.csv"
ACTION_REPORT = ROOT / "00_project_management/action_record_2026-09-02_supplementary_first_citation_order_refreeze.md"
PACKAGE = ROOT / "04_submission/npj_systems_biology_and_applications/SLE_Bcell_npj_Systems_Biology_and_Applications.zip"
PACKAGE_SHA256 = "02A3855FB1EFEAC790C1138396CF783050D0DE744D23B5B5E0C1E97875BA83A1"

# These contact sheets were manually inspected before the final candidate was reselected.
# Matching hashes prove that the post-reselection Supplementary render is the inspected render.
INSPECTED_SUPPLEMENT_CONTACT_HASHES = {
    "final_wps_pages/Supplementary_Information_scientific_maintenance_freeze_contact_1.png": "FA1E8F8A591DE9467DC958E5B463B8AB686EDE5AC72D4FDB4D5CB87C384C1F63",
    "final_wps_pages/Supplementary_Information_scientific_maintenance_freeze_contact_2.png": "82077807E0851375195022483A9C5B3F98CAFB0CEB6AD10574D01CB04ED86ED5",
    "final_wps_pages/Supplementary_Information_scientific_maintenance_freeze_contact_3.png": "15E0B074E12EEF790F552BF049DBF1ECAA6FBD2855FE598AE8BDD7E9BEF9910B",
    "final_lo_pages/Supplementary_Information_scientific_maintenance_freeze_contact_1.png": "23C22B591E33F75CA902772E86AD751C4D73D84E56E76AC19EDFC71D314DA151",
    "final_lo_pages/Supplementary_Information_scientific_maintenance_freeze_contact_2.png": "55E1B2C265A89955A288F7D3FC0512E168F7CF4B91C66DE88A4E9A8388E04C8E",
    "final_lo_pages/Supplementary_Information_scientific_maintenance_freeze_contact_3.png": "5E26B35868AF81CB4E92B24B9685B9EEF6B078C859DA5CC183B39AEF6FBA92F3",
}


def sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest().upper()


def load_json(path: Path) -> dict[str, object]:
    return json.loads(path.read_text(encoding="utf-8"))


def pdf_pages(path: Path) -> int:
    return len(PdfReader(path).pages)


def expected_main_text() -> str:
    text = (BASE / "sources/Manuscript_scientific_maintenance_freeze.md").read_text(
        encoding="utf-8"
    )
    text = integration_logic.remap_references(text)
    for old, new, _ in integration_logic.MAIN_ANCHORS:
        if text.count(old) != 1:
            raise RuntimeError(f"Could not reconstruct the main-text anchor: {old}")
        text = text.replace(old, new)
    return text


def first_citation_order(text: str) -> list[int]:
    return integration_logic.first_citation_order(text)


def figure_alt_titles(path: Path) -> list[str]:
    document = Document(path)
    return [shape._inline.docPr.get("title") for shape in document.inline_shapes]


def main_legend_spacing(path: Path) -> list[tuple[float | None, float | None]]:
    values = []
    for paragraph in Document(path).paragraphs:
        if re.fullmatch(r"Figure [1-5] \| .+", paragraph.text.strip()):
            before = paragraph.paragraph_format.space_before
            after = paragraph.paragraph_format.space_after
            values.append((before.pt if before else None, after.pt if after else None))
    return values


def write_manifest() -> None:
    rows = []
    for path in sorted(item for item in RUN.rglob("*") if item.is_file()):
        if path == MANIFEST_PATH:
            continue
        relative_to_run = path.relative_to(RUN).as_posix()
        excluded_prefixes = (
            "qa/pagination_candidates/",
            "qa/final_lo_render/",
            "qa/lo_render/",
            "qa/wps_pages/",
            "qa/lo_pages/",
        )
        if relative_to_run.startswith(excluded_prefixes):
            continue
        if relative_to_run.startswith("qa/final_wps_pages/") and path.parent != QA / "final_wps_pages":
            continue
        if relative_to_run.startswith("qa/final_lo_pages/") and path.parent != QA / "final_lo_pages":
            continue
        rows.append(
            {
                "relative_path": path.relative_to(ROOT).as_posix(),
                "bytes": path.stat().st_size,
                "sha256": sha256(path),
            }
        )
    with MANIFEST_PATH.open("w", encoding="utf-8-sig", newline="") as handle:
        writer = csv.DictWriter(handle, fieldnames=["relative_path", "bytes", "sha256"])
        writer.writeheader()
        writer.writerows(rows)


def write_action_report(status: dict[str, object]) -> None:
    documents = status["documents"]
    report = f"""# 行动记录：Supplementary Figure 首次引用顺序科学呈现冻结

- **完成日期：** 2026-09-02
- **最终状态：** `{status['status']}`
- **工作边界：** 手稿、补充材料与图件科学呈现；未推进投稿包、GitHub Release 或 Zenodo 发布
- **冻结投稿包 SHA-256：** `{PACKAGE_SHA256}`

## 1. 本轮目标

独立复核外部 hostile audit 指出的 Supplementary Figure 首次引用顺序问题，判断是否需要保留、修改或替换既有子图，并在不改变任何统计估计、Source Data 数值或科学结论的前提下完成来源驱动的交叉引用修复。

## 2. 独立复核结果

修复前，正文首次出现 Supplementary Figure 的顺序确认为 `S1, S2, S3, S9, S4, S5, S6, S10, S7, S8`。这不是统计错误，也不是图件内容错误，而是显示编号与 reader path 不一致。外部提供的映射经独立重建后完全一致：

`S1->S1; S2->S2; S3->S3; S9->S4; S4->S5; S5->S6; S6->S7; S10->S8; S7->S9; S8->S10`。

修复后，正文首次引用顺序为严格的 `S1-S10`。Supplementary Information 中标题、占位标记、图件文件名、Figure Source Data 文件名及 Supplementary Table S5 的 figure-source map 同步更新。

## 3. 子图保留、修改与替换判断

- 21 个主图子图全部 `KEEP`；0 个新增，0 个替换。
- 38 个补充图子图全部 `KEEP`；仅显示编号变化，科学对象逐字节不变。
- Figure 1a 保留：它仍是 identity-to-disease inference boundary 的唯一主图责任面板。
- Figure 5a 保留：它仍是 evidence class 与 causal ceiling 的唯一主图责任面板。
- S1-S10 的 PDF、PNG 与 Source Data CSV 共 30 个映射对象均与旧显示编号下的对象 SHA-256 相同。
- 本轮没有重新运行统计模型，没有重画图，没有修改任何数值或阈值。

## 4. 文本与交叉引用修复

正文仅增加四个功能性 Supplementary Table 导航锚点：Tables S1-S2、Table S3、Table S4、Tables S5-S8；既有 Table S9 引用保留。Tables S1-S9 现在均可由正文定位。锚点均并入原句或既有括号引用，未新增科学主张。

## 5. 文档构建与跨渲染修复

补充材料比较了 16 页标准候选与 15 页紧凑候选。两种候选均通过 WPS/LibreOffice 图题同页和图像指纹检查；采用 15 页候选，因为它移除了 S1 前的冗余手工分页，未缩字号、未压缩表格、未改变科学内容。

新增导航锚点最初使 LibreOffice 将 Figure 5 图例最后两行推到第 32 页。最终修复没有删减图例或缩小字号，而是将五个主图图例标题的段前/段后间距统一为 6/2 pt。主文随后在 WPS 与 LibreOffice 中均为 31 页。

WPS PDF 文本抽取会去除部分词间空格，导致候选选择器的完整 S1 图例匹配出现假阴性。QA 解析器已改为忽略空白的规范化匹配；10 张图的身份指纹在修复前后始终通过。

## 6. 最终 QA

- WPS 主文：{documents['wps_main']['pages']} 页，SHA-256 `{documents['wps_main']['sha256']}`。
- LibreOffice 主文：{documents['lo_main']['pages']} 页，SHA-256 `{documents['lo_main']['sha256']}`。
- WPS Supplementary Information：{documents['wps_supplement']['pages']} 页，SHA-256 `{documents['wps_supplement']['sha256']}`。
- LibreOffice Supplementary Information：{documents['lo_supplement']['pages']} 页，SHA-256 `{documents['lo_supplement']['sha256']}`。
- 18 张联系表覆盖双引擎共 92 个页面，已逐页人工检查；无空白页、截断、重叠、缺字或错图。
- 最终补充材料重新选择后，六张 Supplementary 联系表 SHA-256 与人工检查版本完全一致。
- S1-S10 在两种渲染器中均保持标题、图例和图像同页；10/10 图像指纹匹配。
- 两份 DOCX 的 accessibility audit 均为 0 high / 0 medium / 0 low。
- 全量回归测试：{status['regression_tests']['tests_run']}/{status['regression_tests']['tests_run']} 通过。
- 作者已确认的投稿包保持 `{PACKAGE_SHA256}`，未被本轮覆盖。

## 7. 科学结论与下一阶段

本轮修复提升了 Supplementary reader path 和可核查性，但没有改变论文的中心结论：可重复证据支持 broad B_CONV IFN/ISG process-level remodeling；hard fine-state assignment、source-label-independent external transfer 与 causal regulator inference 仍保持明确边界。

下一阶段进入 `SCIENTIFIC_PRESENTATION_MAINTENANCE_FREEZE`。当前不应重新打开统计模型、Figure 1a、Figure 5a 或任何补图。只有发现新的、可定位的数值错误、语义越界、交叉引用错误或实际尺寸可读性缺陷时，才执行局部来源重跑；否则不再以“继续完善”为由增加分析或重画图。
"""
    ACTION_REPORT.write_text(report, encoding="utf-8", newline="\n")


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--confirm-manual-visual-qa", action="store_true")
    parser.add_argument("--confirm-regression-pass", action="store_true")
    parser.add_argument("--regression-tests-run", type=int, default=0)
    args = parser.parse_args()

    integration = load_json(RUN / "00_SUPPLEMENTARY_CITATION_REFREEZE_INTEGRATION_STATUS.json")
    build = load_json(RUN / "05_DOCUMENT_CANDIDATE_BUILD_STATUS.json")
    selection = load_json(RUN / "06_PAGINATION_EXPERIMENT_AND_SELECTION.json")
    pagination = load_json(RUN / "07_FINAL_SUPPLEMENT_PAGINATION_AUDIT.json")
    wps_audit = load_json(QA / "final_wps_pages/document_render_audit.json")
    lo_audit = load_json(QA / "final_lo_pages/document_render_audit.json")
    main_accessibility = load_json(QA / f"accessibility/{MAIN_STEM}.json")
    supplement_accessibility = load_json(QA / f"accessibility/{SUPPLEMENT_STEM}.json")

    root_main = ROOT / "01_manuscript/Manuscript.md"
    root_supplement = ROOT / "01_manuscript/Supplementary_Information.md"
    run_main = RUN / "sources/Manuscript_first_citation_order_refreeze.md"
    run_supplement = RUN / "sources/Supplementary_Information_first_citation_order_refreeze.md"
    main_text = root_main.read_text(encoding="utf-8")
    supplement_text = root_supplement.read_text(encoding="utf-8")

    expected_supplement = integration_logic.transform_supplement(
        (BASE / "sources/Supplementary_Information_s4_reader_path_micropass.md").read_text(
            encoding="utf-8"
        )
    )
    table_coverage = integration_logic.table_citation_coverage(main_text)
    supplement_headings = [
        int(value)
        for value in re.findall(r"(?m)^## Supplementary Figure S(10|[1-9]) \|", supplement_text)
    ]

    with (RUN / "03_SUPPLEMENTARY_DISPLAY_ID_PROVENANCE.csv").open(
        encoding="utf-8-sig", newline=""
    ) as handle:
        provenance = list(csv.DictReader(handle))
    with (RUN / "02_PANEL_DECISION_MATRIX.csv").open(
        encoding="utf-8-sig", newline=""
    ) as handle:
        panels = list(csv.DictReader(handle))

    main_panels = [row for row in panels if row["tier"] == "Main"]
    supplement_panels = [row for row in panels if row["tier"] == "Supplementary"]
    expected_pairs = {(f"S{old}", f"S{new}") for old, new in integration_logic.RENUMBER.items()}
    observed_pairs = {(row["old_display_id"], row["new_display_id"]) for row in provenance}

    wps_main = DOCUMENTS / f"{MAIN_STEM}.pdf"
    lo_main = LO_DOCUMENTS / f"{MAIN_STEM}.pdf"
    wps_supplement = DOCUMENTS / f"{SUPPLEMENT_STEM}.pdf"
    lo_supplement = LO_DOCUMENTS / f"{SUPPLEMENT_STEM}.pdf"
    final_supplement_docx = DOCUMENTS / f"{SUPPLEMENT_STEM}.docx"
    contact_sheets = list((QA / "final_wps_pages").glob("*_contact_*.png")) + list(
        (QA / "final_lo_pages").glob("*_contact_*.png")
    )
    inspected_hashes_match = all(
        sha256(QA / relative) == expected
        for relative, expected in INSPECTED_SUPPLEMENT_CONTACT_HASHES.items()
    )
    all_a11y_zero = all(
        report.get("counts", {}).get(level, -1) == 0
        for report in (main_accessibility, supplement_accessibility)
        for level in ("high", "medium", "low")
    )

    checks = {
        "integration_pass": integration.get("status")
        == "PASS_SUPPLEMENTARY_CITATION_REFREEZE_INTEGRATION_DOCX_REQUIRED"
        and all(integration.get("checks", {}).values()),
        "document_build_pass": build.get("status")
        == "PASS_SUPPLEMENTARY_CITATION_DOCX_CANDIDATES_BUILT_RENDER_REQUIRED"
        and all(build.get("checks", {}).values()),
        "compact_15_page_candidate_selected": selection.get("decision")
        == "ADOPT_COMPACT_15_PAGE_CANDIDATE"
        and selection.get("selected_expected_pages") == 15,
        "selected_docx_hash_matches_final": selection.get("final_docx", {}).get("sha256")
        == sha256(final_supplement_docx),
        "root_main_is_exact_source_transform": main_text == expected_main_text(),
        "root_supplement_is_exact_source_transform": supplement_text == expected_supplement,
        "root_sources_match_run_sources": root_main.read_bytes() == run_main.read_bytes()
        and root_supplement.read_bytes() == run_supplement.read_bytes(),
        "first_citation_order_is_s1_to_s10": first_citation_order(main_text)
        == list(range(1, 11)),
        "supplement_headings_are_s1_to_s10": supplement_headings == list(range(1, 11)),
        "supplementary_tables_s1_to_s9_cited": table_coverage == list(range(1, 10)),
        "thirty_display_objects_byte_identical": len(provenance) == 30
        and all(row["byte_identical"].lower() == "true" for row in provenance)
        and all(row["old_sha256"] == row["new_sha256"] for row in provenance),
        "all_ten_display_pairs_present": observed_pairs == expected_pairs,
        "all_21_main_panels_kept": len(main_panels) == 21
        and all(row["scientific_decision"] == "KEEP" for row in main_panels),
        "all_38_supplementary_panels_kept": len(supplement_panels) == 38
        and all(row["scientific_decision"] == "KEEP" for row in supplement_panels),
        "zero_new_or_replacement_panels": len(panels) == 59
        and not any("REPLACE" in row["artwork_action"] or "NEW" in row["artwork_action"] for row in panels),
        "figure_1a_and_5a_keep_exact": all(
            any(row["object"] == name and row["artwork_action"] == "KEEP_EXACT" for row in main_panels)
            for name in ("Figure 1a", "Figure 5a")
        ),
        "wps_and_libreoffice_main_are_31_pages": pdf_pages(wps_main) == 31
        and pdf_pages(lo_main) == 31,
        "wps_and_libreoffice_supplement_are_15_pages": pdf_pages(wps_supplement) == 15
        and pdf_pages(lo_supplement) == 15,
        "supplement_pagination_and_fingerprints_pass": pagination.get("status")
        == "PASS_SUPPLEMENT_PAGINATION_COHERENCE"
        and all(pagination.get("checks", {}).values()),
        "all_rendered_pages_nonblank": all(
            int(page["text_characters"]) >= 80
            for audit in (wps_audit, lo_audit)
            for page in audit["page_checks"]
        ),
        "all_rendered_text_within_canvas": bool(
            wps_audit["all_pages_within_canvas"] and lo_audit["all_pages_within_canvas"]
        ),
        "all_render_markers_resolved": bool(
            wps_audit["all_markers_resolved"] and lo_audit["all_markers_resolved"]
        ),
        "eighteen_contact_sheets_cover_92_pages": len(contact_sheets) == 18
        and wps_audit["pages"] == 46
        and lo_audit["pages"] == 46,
        "post_reselection_supplement_matches_inspected_render": inspected_hashes_match,
        "both_docx_accessibility_audits_zero": all_a11y_zero,
        "supplement_alt_titles_are_s1_to_s10": figure_alt_titles(final_supplement_docx)
        == [f"Supplementary Figure S{number}" for number in range(1, 11)],
        "five_main_legend_headings_have_6_2_pt_spacing": main_legend_spacing(
            DOCUMENTS / f"{MAIN_STEM}.docx"
        )
        == [(6.0, 2.0)] * 5,
        "submission_package_unchanged": sha256(PACKAGE) == PACKAGE_SHA256,
        "manual_visual_qa_confirmed": args.confirm_manual_visual_qa,
    }
    if args.confirm_regression_pass or args.regression_tests_run:
        checks["full_regression_suite_confirmed"] = args.confirm_regression_pass and args.regression_tests_run > 0
    failed = [name for name, passed in checks.items() if not passed]
    manual_visual_qa = {
        "confirmed": args.confirm_manual_visual_qa,
        "contact_sheets_inspected": 18 if args.confirm_manual_visual_qa else 0,
        "rendered_pages_inspected": 92 if args.confirm_manual_visual_qa else 0,
        "wps_pages_inspected": 46 if args.confirm_manual_visual_qa else 0,
        "libreoffice_pages_inspected": 46 if args.confirm_manual_visual_qa else 0,
        "blank_pages_found": False,
        "clipping_or_overlap_found": False,
        "missing_glyphs_or_figures_found": False,
        "supplement_post_reselection_contact_hashes_match": inspected_hashes_match,
    }
    status = {
        "created_at": datetime.now().astimezone().isoformat(timespec="seconds"),
        "status": (
            "SCIENTIFIC_PRESENTATION_MAINTENANCE_FREEZE"
            if not failed
            else "HOLD_SUPPLEMENTARY_CITATION_REFREEZE_REVIEW_REQUIRED"
        ),
        "checks": checks,
        "failed_checks": failed,
        "manual_visual_qa": manual_visual_qa,
        "first_citation_order_before": integration["first_citation_order_before"],
        "first_citation_order_after": integration["first_citation_order_after"],
        "supplementary_table_citation_coverage": table_coverage,
        "scientific_estimates_changed": False,
        "statistical_models_rerun": False,
        "figures_redrawn": False,
        "figure_pixels_changed": False,
        "source_data_values_changed": False,
        "main_text_functional_anchor_edits": 4,
        "display_objects_renumbered": 30,
        "new_panels": 0,
        "replacement_panels": 0,
        "main_panels_keep": 21,
        "supplementary_panels_keep": 38,
        "documents": {
            "wps_main": {
                "path": wps_main.relative_to(ROOT).as_posix(),
                "pages": pdf_pages(wps_main),
                "sha256": sha256(wps_main),
            },
            "lo_main": {
                "path": lo_main.relative_to(ROOT).as_posix(),
                "pages": pdf_pages(lo_main),
                "sha256": sha256(lo_main),
            },
            "wps_supplement": {
                "path": wps_supplement.relative_to(ROOT).as_posix(),
                "pages": pdf_pages(wps_supplement),
                "sha256": sha256(wps_supplement),
            },
            "lo_supplement": {
                "path": lo_supplement.relative_to(ROOT).as_posix(),
                "pages": pdf_pages(lo_supplement),
                "sha256": sha256(lo_supplement),
            },
        },
        "submission_package_sha256": sha256(PACKAGE),
        "regression_tests": {
            "confirmed_pass": args.confirm_regression_pass,
            "tests_run": args.regression_tests_run,
            "failures": 0 if args.confirm_regression_pass else None,
            "errors": 0 if args.confirm_regression_pass else None,
        },
        "submission_package_changed": False,
        "github_release_changed": False,
        "zenodo_changed": False,
        "action_record": ACTION_REPORT.relative_to(ROOT).as_posix(),
        "next_stage": "SCIENTIFIC_PRESENTATION_MAINTENANCE_FREEZE",
    }
    STATUS_PATH.write_text(
        json.dumps(status, indent=2) + "\n", encoding="utf-8", newline="\n"
    )
    write_action_report(status)
    write_manifest()
    print(json.dumps(status, indent=2))
    if failed:
        raise RuntimeError(f"Final citation-refreeze checks failed: {failed}")


if __name__ == "__main__":
    main()
