"""Apply reviewed local prose changes to a separate manuscript, never the release."""

import argparse
from collections import defaultdict
from copy import deepcopy
from datetime import datetime
import json
from pathlib import Path
import re

from docx import Document
from docx.oxml import OxmlElement

from phase17_postc9_14_audit_refined_manuscript import (
    ROOT, citation_order, extract_docx, file_record, normalized, numeric_tokens, section,
)


def apply_changes(document, changes):
    grouped = defaultdict(list)
    for change in changes:
        grouped[change["paragraph"]].append(change)
    before = [paragraph.text for paragraph in document.paragraphs]
    record = []
    for index, edits in sorted(grouped.items()):
        paragraph = document.paragraphs[index]
        revised = paragraph.text
        for change in edits:
            if revised.count(change["old"]) != 1:
                raise ValueError("Replacement must have exactly one match: " + change["id"])
            revised = revised.replace(change["old"], change["new"], 1)
        if numeric_tokens(before[index]) != numeric_tokens(revised):
            raise ValueError("A numerical token changed in paragraph " + str(index))
        source_run = max(paragraph.runs, key=lambda run: len(run.text))
        properties = deepcopy(source_run._r.rPr)
        paragraph.clear()
        label = re.match(r"^(Background|Methods|Results|Conclusions): ", revised)
        pieces = [(revised[:label.end() - 1], True), (revised[label.end() - 1:], False)] if label else [(revised, None)]
        for text, bold in pieces:
            run = paragraph.add_run(text)
            if properties is not None:
                run._r.insert(0, deepcopy(properties))
            if bold is not None:
                run.bold = bold
        record.append({"paragraph": index, "change_ids": [c["id"] for c in edits],
                       "before": before[index], "after": paragraph.text})
    for index, text in enumerate(before):
        if index not in grouped and document.paragraphs[index].text != text:
            raise ValueError("Unrequested paragraph modification")
    return record


def suppress_header_footer_numbers(document):
    parts = {}
    for item in document.sections:
        for container in (item.header, item.footer):
            parts[str(container.part.partname)] = container
    count = 0
    for container in parts.values():
        for paragraph in container.paragraphs:
            props = paragraph._p.get_or_add_pPr()
            if not props.findall("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}suppressLineNumbers"):
                props.append(OxmlElement("w:suppressLineNumbers"))
            count += 1
    return count


def main():
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--input-docx", required=True, type=Path)
    parser.add_argument("--changes", required=True, type=Path)
    parser.add_argument("--output-dir", required=True, type=Path)
    args = parser.parse_args()
    output = args.output_dir.resolve()
    if not output.is_relative_to(ROOT / "00_project_management") or output.name != "review_candidate":
        raise ValueError("Write only to a dedicated project-management review_candidate")
    changes = json.loads(args.changes.read_text(encoding="utf-8"))
    source = file_record(args.input_docx)
    if source["sha256"] != changes["input_sha256"]:
        raise ValueError("Received DOCX does not match the reviewed input")
    before, _, refs, _ = extract_docx(args.input_docx)
    document = Document(args.input_docx)
    edits = apply_changes(document, changes["changes"])
    suppress_count = suppress_header_footer_numbers(document)
    output.mkdir(parents=True, exist_ok=True)
    destination = output / "Manuscript.docx"
    document.save(destination)
    after, _, new_refs, _ = extract_docx(destination)
    if new_refs != refs:
        raise ValueError("Bibliography changed")
    if normalized(section(before, "Figure legends")) != normalized(section(after, "Figure legends")):
        raise ValueError("Figure legends changed")
    body = after.split("## Background\n", 1)[1].split("## References\n", 1)[0]
    if citation_order(body) != list(range(1, 33)):
        raise ValueError("Reference order changed")
    (output / "Manuscript.md").write_text(after, encoding="utf-8")
    receipt = {
        "created_at": datetime.now().astimezone().isoformat(timespec="seconds"),
        "input": source, "change_specification": file_record(args.changes),
        "outputs": [file_record(destination), file_record(output / "Manuscript.md")],
        "paragraphs_changed": len(edits), "replacements": len(changes["changes"]),
        "paragraph_changes": edits, "header_footer_line_number_suppression": suppress_count,
        "reference_count": len(refs), "reference_order_preserved": True,
        "figure_legends_preserved": True, "numeric_tokens_preserved_in_each_changed_paragraph": True,
        "canonical_sources_modified": False, "scientific_reanalysis": False,
        "author_approval": "PENDING_EXACT_REFINED_FILES", "submission_authorized": False,
        "scope": "Standalone manuscript review candidate; not a submission package or a replacement of corrected_candidate.zip",
    }
    (output / "build_receipt.json").write_text(json.dumps(receipt, indent=2) + "\n", encoding="utf-8")
    print(json.dumps({key: receipt[key] for key in ("paragraphs_changed", "replacements", "header_footer_line_number_suppression",
          "reference_count", "figure_legends_preserved", "author_approval")}, indent=2))


if __name__ == "__main__":
    main()
