#!/usr/bin/env python3
"""Build the methodologically revised RP as a styled Word document."""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "01_manuscript" / "research_proposal_v14_methodologically_revised_2026-08-10.md"
OUTPUT = ROOT / "01_manuscript" / "research_proposal_v14_methodologically_revised_2026-08-10.docx"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
BLACK = RGBColor(0, 0, 0)
MUTED = RGBColor(89, 89, 89)
LIGHT_FILL = "F4F6F9"


def set_run_font(run, size=None, color=BLACK, bold=None, italic=None, name="Calibri"):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa):
    total = sum(widths_dxa)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for cell, width in zip(row.cells, widths_dxa):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def mark_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = tr_pr.find(qn("w:tblHeader"))
    if tbl_header is None:
        tbl_header = OxmlElement("w:tblHeader")
        tr_pr.append(tbl_header)
    tbl_header.set(qn("w:val"), "true")


def add_page_field(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    set_run_font(run, size=9, color=MUTED)
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char1, instr, fld_char2])


def add_inline(paragraph, text):
    token = re.compile(r"(\*\*.+?\*\*|\*.+?\*|`.+?`)")
    cursor = 0
    for match in token.finditer(text):
        if match.start() > cursor:
            set_run_font(paragraph.add_run(text[cursor:match.start()]), size=11)
        value = match.group(0)
        if value.startswith("**"):
            run = paragraph.add_run(value[2:-2])
            set_run_font(run, size=11, bold=True)
        elif value.startswith("*"):
            run = paragraph.add_run(value[1:-1])
            set_run_font(run, size=11, italic=True)
        else:
            run = paragraph.add_run(value[1:-1])
            set_run_font(run, size=10.5, color=DARK_BLUE, name="Consolas")
        cursor = match.end()
    if cursor < len(text):
        set_run_font(paragraph.add_run(text[cursor:]), size=11)


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    heading_tokens = {
        "Heading 1": (16, BLUE, 16, 8),
        "Heading 2": (13, BLUE, 12, 6),
        "Heading 3": (12, DARK_BLUE, 8, 4),
    }
    for name, (size, color, before, after) in heading_tokens.items():
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        style.paragraph_format.left_indent = Inches(0.5)
        style.paragraph_format.first_line_indent = Inches(-0.319)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.208


def add_cover(doc):
    section = doc.sections[0]
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(
        header.add_run("MBI6013 | Research Proposal | Revised analysis plan"),
        size=9,
        color=MUTED,
    )
    add_page_field(section.footer.paragraphs[0])

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(38)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(10)
    set_run_font(p.add_run("MBI6013 RESEARCH PROPOSAL"), size=12, color=MUTED, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    set_run_font(
        p.add_run("Donor- and cohort-resolved single-cell analysis\nof compositional and transcriptional B-cell\nremodeling in systemic lupus erythematosus"),
        size=18.5,
        color=BLACK,
        bold=True,
    )

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(28)
    set_run_font(
        p.add_run("Methodologically revised version 14"),
        size=13,
        color=DARK_BLUE,
        italic=True,
    )

    table = doc.add_table(rows=3, cols=2)
    table.style = "Table Grid"
    values = [
        ("Document status", "Active research plan"),
        ("Study design", "Secondary analysis of public human single-cell transcriptomic data"),
        ("Revision date", "10 August 2026"),
    ]
    for row, (label, value) in zip(table.rows, values):
        set_cell_shading(row.cells[0], LIGHT_FILL)
        p0 = row.cells[0].paragraphs[0]
        p0.paragraph_format.space_after = Pt(0)
        set_run_font(p0.add_run(label), size=10.5, bold=True)
        p1 = row.cells[1].paragraphs[0]
        p1.paragraph_format.space_after = Pt(0)
        set_run_font(p1.add_run(value), size=10.5)
    set_table_geometry(table, [2400, 6960])

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(26)
    p.paragraph_format.space_after = Pt(0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(
        p.add_run("Outcome-locked, sample-level and reproducible by design"),
        size=11,
        color=MUTED,
        italic=True,
    )
    doc.add_page_break()


def parse_table(lines, start):
    rows = []
    index = start
    while index < len(lines) and lines[index].strip().startswith("|"):
        cells = [cell.strip() for cell in lines[index].strip().strip("|").split("|")]
        rows.append(cells)
        index += 1
    if len(rows) >= 2 and all(set(cell) <= {"-", ":"} for cell in rows[1]):
        rows.pop(1)
    return rows, index


def add_markdown_table(doc, rows):
    if not rows:
        return
    n_cols = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=n_cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    mark_table_header(table.rows[0])
    narrative_weights = [max(10, max(len(row[i]) if i < len(row) else 0 for row in rows)) for i in range(n_cols)]
    total_weight = sum(narrative_weights)
    widths = [int(9360 * weight / total_weight) for weight in narrative_weights]
    widths[-1] += 9360 - sum(widths)
    for row_index, row in enumerate(rows):
        for col_index, value in enumerate(row):
            cell = table.cell(row_index, col_index)
            if row_index == 0:
                set_cell_shading(cell, LIGHT_FILL)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.1
            add_inline(p, value)
            for run in p.runs:
                run.font.size = Pt(9.2)
                if row_index == 0:
                    run.bold = True
    set_table_geometry(table, widths)
    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(2)


def build():
    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.45)
    section.footer_distance = Inches(0.45)
    configure_styles(doc)
    add_cover(doc)

    # Skip the Markdown cover block through the first blank after Study type.
    body_start = next(i for i, line in enumerate(lines) if line.strip() == "## Project overview")
    index = body_start
    paragraph_buffer = []

    def flush_paragraph():
        if not paragraph_buffer:
            return
        text = " ".join(part.strip() for part in paragraph_buffer)
        p = doc.add_paragraph()
        add_inline(p, text)
        paragraph_buffer.clear()

    while index < len(lines):
        raw = lines[index]
        line = raw.strip()
        if not line:
            flush_paragraph()
            index += 1
            continue
        if line.startswith("|"):
            flush_paragraph()
            rows, index = parse_table(lines, index)
            add_markdown_table(doc, rows)
            continue
        if line.startswith("### "):
            flush_paragraph()
            doc.add_paragraph(line[4:], style="Heading 2")
        elif line.startswith("## "):
            flush_paragraph()
            doc.add_paragraph(line[3:], style="Heading 1")
        elif line.startswith("- "):
            flush_paragraph()
            p = doc.add_paragraph(style="List Bullet")
            add_inline(p, line[2:])
        elif re.match(r"^\d+\.\s", line):
            flush_paragraph()
            p = doc.add_paragraph(style="List Number")
            add_inline(p, re.sub(r"^\d+\.\s+", "", line))
            if doc.paragraphs[-2].text == "References" or any(
                previous.text == "References" for previous in doc.paragraphs[-3:-1]
            ):
                p.paragraph_format.left_indent = Inches(0.35)
                p.paragraph_format.first_line_indent = Inches(-0.25)
        elif line.startswith("> "):
            flush_paragraph()
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.3)
            p.paragraph_format.right_indent = Inches(0.3)
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(8)
            add_inline(p, line[2:])
            for run in p.runs:
                run.italic = True
                run.font.color.rgb = DARK_BLUE
        else:
            paragraph_buffer.append(line)
        index += 1
    flush_paragraph()

    for paragraph in doc.paragraphs:
        if paragraph.style.name == "Normal":
            paragraph.paragraph_format.widow_control = True

    core = doc.core_properties
    core.title = "Donor- and cohort-resolved single-cell analysis of B-cell remodeling in SLE"
    core.subject = "MBI6013 methodologically revised research proposal"
    core.author = "Research team"
    core.keywords = "SLE; B cells; single-cell RNA-seq; pseudobulk; composition"
    core.comments = "Generated from the version-controlled v14 Markdown source."
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
