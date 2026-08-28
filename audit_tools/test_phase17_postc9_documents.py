"""Regression checks for review-document formatting without editing outputs."""

import unittest

from docx import Document
from docx.oxml.ns import qn

from phase17_c8s_04_build_documents import add_inline, add_markdown_table, table_widths


class ReviewDocumentTests(unittest.TestCase):
    def test_inline_emphasis_preserves_text(self):
        paragraph = Document().add_paragraph()
        add_inline(paragraph, "A *Science* paper with **bold** and `code`.")
        self.assertEqual(paragraph.text, "A Science paper with bold and code.")
        self.assertTrue(next(run for run in paragraph.runs if run.text == "Science").italic)
        self.assertTrue(next(run for run in paragraph.runs if run.text == "bold").bold)
        self.assertEqual(next(run for run in paragraph.runs if run.text == "code").font.name, "Courier New")

    def test_table_header_stays_with_data_and_rows_do_not_split(self):
        document = Document()
        add_markdown_table(document, "| Metric | Value |\n|---|---|\n| Precision | 0.885210 |", 11)
        table = document.tables[0]
        self.assertEqual(len(table.rows), 2)
        for cell in table.rows[0].cells:
            self.assertTrue(cell.paragraphs[0].paragraph_format.keep_with_next)
        for row in table.rows:
            self.assertIsNotNone(row._tr.get_or_add_trPr().find(qn("w:cantSplit")))
        self.assertEqual(table.cell(1, 1).text, "0.885210")

    def test_dataset_table_has_explicit_widths(self):
        widths = table_widths([["Resource", "Role", "Biological unit", "Active scope"]])
        self.assertEqual(widths, [1600, 1700, 2100, 3960])
        self.assertEqual(sum(widths), 9360)
        depletion = table_widths([["Depletion", "Method", "Positive direction", "Dedicated q<0.05", "Main qualification"]])
        self.assertEqual(depletion[1], 1300)
        self.assertEqual(sum(depletion), 9360)


if __name__ == "__main__":
    unittest.main()
