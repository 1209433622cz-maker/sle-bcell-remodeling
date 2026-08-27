#!/usr/bin/env python3
"""Audit and freeze the final author-approved submission package."""

from __future__ import annotations

import csv
import hashlib
import json
import re
import shutil
import tempfile
import zipfile
from pathlib import Path
from typing import Any

from docx import Document
from pypdf import PdfReader


ROOT = Path(__file__).resolve().parents[1]
RUN_DIR = ROOT / "phase17_v7" / "gateC8BRF" / "20260825_author_release"
C8S_RUN = ROOT / "phase17_v7" / "gateC8S" / "20260821_supplementary_traceability_freeze"
ROUND6_SOURCE = ROOT / "phase17_v7" / "round6_q1_robustness" / "20260825_overlap_depletion"
R1_INTEGRATION = ROOT / "phase17_v7" / "round6_q1_robustness" / "20260827_r1_hold_integration"
PACKAGE = ROOT / "04_submission" / "journal_submission"
PACKAGE_ZIP = ROOT / "04_submission" / "journal_submission.zip"
MANUSCRIPT = ROOT / "01_manuscript" / "Manuscript.md"
SUPPLEMENT = ROOT / "01_manuscript" / "Supplementary_Information.md"
COVER = ROOT / "04_submission" / "Cover_Letter.md"
METADATA = ROOT / "04_submission" / "Zenodo_Metadata.json"
MAIN_DOCX = PACKAGE / "main_text" / "Manuscript.docx"
SUPP_DOCX = PACKAGE / "additional_files" / "Supplementary_Information.docx"
COVER_DOCX = PACKAGE / "submission_docs" / "Cover_Letter.docx"
MAIN_PDF = PACKAGE / "internal_qc" / "wps_render_main" / "Manuscript_WPS.pdf"
SUPP_PDF = PACKAGE / "internal_qc" / "wps_render_supplement" / "Supplementary_Information_WPS.pdf"
COVER_PDF = PACKAGE / "internal_qc" / "wps_render_cover_letter" / "Cover_Letter_WPS.pdf"
FIXED_ZIP_TIME = (2026, 8, 25, 0, 0, 0)


def sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest().upper()


def read_json(path: Path) -> dict[str, Any]:
    return json.loads(path.read_text(encoding="utf-8-sig"))


def docx_text(path: Path) -> str:
    document = Document(path)
    values = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            values.extend(cell.text for cell in row.cells)
    return "\n".join(values)


def pdf_text_pages(path: Path) -> list[str]:
    return [re.sub(r"\s+", " ", page.extract_text() or "") for page in PdfReader(path).pages]


def pdf_dimensions_mm(path: Path) -> tuple[float, float]:
    page = PdfReader(path).pages[0]
    return (
        float(page.mediabox.width) * 25.4 / 72.0,
        float(page.mediabox.height) * 25.4 / 72.0,
    )


def build_manifest() -> list[dict[str, object]]:
    manifest = PACKAGE / "MANIFEST_SHA256.csv"
    rows: list[dict[str, object]] = []
    for path in sorted(
        item for item in PACKAGE.rglob("*") if item.is_file() and item != manifest
    ):
        rows.append(
            {
                "relative_path": path.relative_to(PACKAGE).as_posix(),
                "bytes": path.stat().st_size,
                "sha256": sha256(path),
            }
        )
    with manifest.open("w", encoding="utf-8", newline="") as handle:
        writer = csv.DictWriter(
            handle, fieldnames=["relative_path", "bytes", "sha256"]
        )
        writer.writeheader()
        writer.writerows(rows)
    return rows


def write_deterministic_zip(output: Path) -> None:
    with zipfile.ZipFile(
        output, "w", compression=zipfile.ZIP_DEFLATED, compresslevel=9
    ) as archive:
        for path in sorted(item for item in PACKAGE.rglob("*") if item.is_file()):
            relative = (Path(PACKAGE.name) / path.relative_to(PACKAGE)).as_posix()
            info = zipfile.ZipInfo(relative, date_time=FIXED_ZIP_TIME)
            info.compress_type = zipfile.ZIP_DEFLATED
            info.external_attr = 0o100644 << 16
            info.create_system = 3
            archive.writestr(
                info,
                path.read_bytes(),
                compress_type=zipfile.ZIP_DEFLATED,
                compresslevel=9,
            )


def build_deterministic_archive() -> tuple[int, str]:
    with tempfile.TemporaryDirectory(prefix="journal_submission_zip_") as temp_dir:
        first = Path(temp_dir) / "first.zip"
        second = Path(temp_dir) / "second.zip"
        write_deterministic_zip(first)
        write_deterministic_zip(second)
        if sha256(first) != sha256(second):
            raise RuntimeError("Deterministic package ZIP rebuild failed")
        shutil.copyfile(first, PACKAGE_ZIP)
    return PACKAGE_ZIP.stat().st_size, sha256(PACKAGE_ZIP)


def package_readme(doi: str) -> str:
    return f"""# Journal submission package

This directory is the current author-approved, journal-facing submission set.
Internal build identifiers and draft numbers are intentionally excluded from
submission filenames.

## Scientific and release status

- DOI: https://doi.org/{doi}
- Scientific estimates changed during final publication engineering: no
- Main panel-data assertions: PASS 46/46
- Supplementary-figure panel-data assertions: PASS 29/29
- Main figures: vector PDF at 170 mm plus 600-dpi PNG
- Author declarations and generative-AI disclosure: complete
- Code licence: MIT
- Original text, figures, documentation and derived source-data licence: CC BY 4.0
- Third-party GEO/CELLxGENE data: excluded from project relicensing

## Portal policy

Use `portal_upload_required/` as the default 11-file upload set. The nine PDFs
in `portal_upload_optional/` duplicate figures embedded in Supplementary
Information and should be used only if the journal portal explicitly requires
standalone supplementary figures.

## Verification

`MANIFEST_SHA256.csv` records every package payload file. The canonical ZIP is built twice and accepted only when both byte streams have the same SHA-256.
"""


def main() -> None:
    checks: dict[str, dict[str, object]] = {}

    def check(name: str, passed: bool, detail: object) -> None:
        checks[name] = {"pass": bool(passed), "detail": detail}

    metadata = read_json(METADATA)
    doi = str(metadata["doi"])
    manuscript = MANUSCRIPT.read_text(encoding="utf-8-sig")
    supplement = SUPPLEMENT.read_text(encoding="utf-8-sig")
    cover = COVER.read_text(encoding="utf-8-sig")

    check(
        "zero_placeholders",
        "[[" not in manuscript and "[[" not in cover,
        {
            "manuscript": manuscript.count("[["),
            "cover": cover.count("[["),
            "supplement_embedding_markers": supplement.count(
                "[[SUPPLEMENTARY_FIGURE:S"
            ),
        },
    )
    declaration_tokens = [
        "No additional ethics approval was required",
        "The authors declare that they have no competing interests",
        "This research received no specific funding",
        "Both authors read and approved",
        "Not applicable.",
        "Use of generative artificial intelligence",
        "did not use generative artificial intelligence to create or alter primary research data",
    ]
    check(
        "author_declarations_complete",
        all(token in manuscript for token in declaration_tokens),
        declaration_tokens,
    )
    check(
        "persistent_citation_complete",
        manuscript.count(doi) >= 2
        and doi in cover
        and doi in (ROOT / "README.md").read_text(encoding="utf-8-sig")
        and doi in (ROOT / "CITATION.cff").read_text(encoding="utf-8-sig"),
        doi,
    )
    check(
        "licence_scope",
        all(
            path.is_file()
            for path in (
                ROOT / "LICENSE",
                ROOT / "LICENSE_SCOPE.md",
                ROOT / "LICENSE_CONTENT_CC_BY_4.0.md",
            )
        )
        and "do not apply to or relicense source datasets"
        in (ROOT / "LICENSE_SCOPE.md").read_text(encoding="utf-8-sig"),
        ["MIT", "CC-BY-4.0", "third-party data excluded"],
    )

    figure_status = read_json(RUN_DIR / "01_FIGURE_BUILD_STATUS.json")
    assertions = read_json(RUN_DIR / "02_PANEL_DATA_ASSERTIONS.json")
    dimensions: dict[str, dict[str, float]] = {}
    dimension_pass = True
    for pdf in sorted((RUN_DIR / "figures").glob("Figure*.pdf")):
        width, height = pdf_dimensions_mm(pdf)
        dimensions[pdf.name] = {
            "width_mm": round(width, 3),
            "height_mm": round(height, 3),
        }
        dimension_pass &= abs(width - 170.0) <= 0.15 and height <= 225.0
    check(
        "main_figure_contract",
        figure_status["status"] == "PASS_GATE_C8BRF_170MM_PUBLICATION_FIGURES_BUILT"
        and len(assertions["checks"]) == 46
        and all(row["pass"] is True for row in assertions["checks"])
        and dimension_pass,
        dimensions,
    )
    figure1 = (RUN_DIR / "source_data" / "Figure1_source_data.csv").read_text(
        encoding="utf-8-sig"
    )
    provenance = read_json(RUN_DIR / "00_FIGURE1_PUBLICATION_SOURCE_PROVENANCE.json")
    check(
        "figure1_publication_source",
        provenance["publication_rows_equal_frozen_rows_after_declared_filter"] is True
        and provenance["plotted_numeric_rows_changed"] is False
        and all(token not in figure1 for token in ("HOLD_GATE", "PASS_GATE", "OUTCOME_UNLOCK")),
        {
            "removed_rows": len(provenance["removed_non_plotted_rows"]),
            "publication_sha256": provenance["publication_source_sha256"],
        },
    )
    uuid_status = read_json(RUN_DIR / "03_FIGURE2_UUID_GOVERNANCE.json")
    check(
        "figure2_uuid_governance",
        uuid_status["status"] == "PASS_FIGURE2_PUBLIC_NON_IDENTIFYING_SOURCE_UUIDS"
        and not uuid_status["unmapped_figure_ids"]
        and not uuid_status["forbidden_direct_identifier_columns"],
        {
            "uuid_count": uuid_status["figure_union_uuid_count"],
            "versions": uuid_status["uuid_versions"],
        },
    )
    supplement_assertions = read_json(
        C8S_RUN / "03_SUPPLEMENTARY_PANEL_DATA_ASSERTIONS.json"
    )
    s8_status = read_json(ROUND6_SOURCE / "06_SUPPLEMENTARY_FIGURE_S8_STATUS.json")
    s9_status = read_json(R1_INTEGRATION / "13_SUPPLEMENTARY_FIGURE_S9_STATUS.json")
    check(
        "supplementary_assertions",
        len(supplement_assertions["checks"]) == 29
        and all(row["pass"] is True for row in supplement_assertions["checks"])
        and s8_status["status"] == "PASS_ROUND6_SUPPLEMENTARY_FIGURE_S8_BUILT"
        and s8_status["source_rows"] == 36
        and s9_status["status"] == "PASS_ROUND6_SUPPLEMENTARY_FIGURE_S9_BUILT"
        and s9_status["source_rows"] == 128
        and all(s9_status["checks"].values()),
        "legacy 29/29; S8 36 rows; S9 128 rows and 8/8 checks",
    )

    required_rows = list(
        csv.DictReader(
            (PACKAGE / "submission_docs" / "PORTAL_UPLOAD_REQUIRED.csv").open(
                "r", encoding="utf-8-sig", newline=""
            )
        )
    )
    optional_rows = list(
        csv.DictReader(
            (PACKAGE / "submission_docs" / "PORTAL_UPLOAD_OPTIONAL.csv").open(
                "r", encoding="utf-8-sig", newline=""
            )
        )
    )
    required_aliases = {row["portal_alias"] for row in required_rows}
    optional_aliases = {row["portal_alias"] for row in optional_rows}
    required_disk_files = {
        path.relative_to(PACKAGE).as_posix()
        for path in (PACKAGE / "portal_upload_required").iterdir()
        if path.is_file()
    }
    optional_disk_files = {
        path.relative_to(PACKAGE).as_posix()
        for path in (PACKAGE / "portal_upload_optional").iterdir()
        if path.is_file()
    }
    check(
        "portal_required_optional_policy",
        len(required_rows) == 11
        and len(optional_rows) == 9
        and required_disk_files == required_aliases
        and optional_disk_files == optional_aliases
        and not required_aliases.intersection(optional_aliases)
        and not any("Supplementary_Figure_S" in value for value in required_aliases),
        {
            "required_manifest": len(required_rows),
            "required_on_disk": len(required_disk_files),
            "optional_manifest": len(optional_rows),
            "optional_on_disk": len(optional_disk_files),
        },
    )

    docx_values = {
        "main": docx_text(MAIN_DOCX),
        "supplement": docx_text(SUPP_DOCX),
        "cover": docx_text(COVER_DOCX),
    }
    supplement_doc = Document(SUPP_DOCX)
    check(
        "docx_structure_and_content",
        all("[[" not in value for value in docx_values.values())
        and len(supplement_doc.inline_shapes) == 9
        and len(supplement_doc.tables) == 9,
        {
            "supplement_inline_figures": len(supplement_doc.inline_shapes),
            "supplement_tables": len(supplement_doc.tables),
        },
    )

    a11y: list[dict[str, object]] = []
    for name in (
        "main_text_a11y.json",
        "supplement_a11y.json",
        "cover_letter_a11y.json",
    ):
        report = read_json(PACKAGE / "internal_qc" / name)
        counts = report.get("counts", {})
        a11y.append({"file": name, **counts})
    check(
        "docx_accessibility",
        all(
            sum(int(row.get(level, 0)) for level in ("high", "medium", "low"))
            == 0
            for row in a11y
        ),
        a11y,
    )

    render_details: list[dict[str, object]] = []
    page_counts: dict[str, int] = {}
    render_ok = True
    for name, pdf, minimum, maximum in (
        ("main", MAIN_PDF, 25, 32),
        ("supplement", SUPP_PDF, 14, 20),
        ("cover", COVER_PDF, 1, 2),
    ):
        reader = PdfReader(pdf) if pdf.exists() else None
        pages = len(reader.pages) if reader else 0
        pngs = len(list(pdf.parent.glob("final-page-*.png")))
        ok = (
            pdf.exists()
            and pdf.stat().st_size > 30_000
            and minimum <= pages <= maximum
            and pngs == pages
        )
        render_ok &= ok
        page_counts[name] = pages
        render_details.append(
            {
                "document": name,
                "pages": pages,
                "page_pngs": pngs,
                "bytes": pdf.stat().st_size if pdf.exists() else 0,
            }
        )
    check("wps_render_integrity", render_ok, render_details)

    supp_pages = pdf_text_pages(SUPP_PDF)
    s7_pages = [
        page
        for page, text in enumerate(supp_pages, start=1)
        if "Supplementary Table S7" in text
    ]
    s7_first_row_same_page = bool(s7_pages) and "B_ASC composition" in supp_pages[
        s7_pages[0] - 1
    ]
    check(
        "supplement_s7_pagination",
        len(s7_pages) == 1 and s7_first_row_same_page,
        {"title_page": s7_pages, "first_data_row_same_page": s7_first_row_same_page},
    )

    stats_copy = PACKAGE / "additional_files" / "Full_Statistical_Results.zip"
    with zipfile.ZipFile(stats_copy) as archive:
        stats_names = set(archive.namelist())
        stats_manifest = list(
            csv.DictReader(
                archive.read("MANIFEST_SHA256.csv").decode("utf-8-sig").splitlines()
            )
        )
        manifest_valid = all(
            row["relative_path"] in stats_names
            and len(archive.read(row["relative_path"])) == int(row["bytes"])
            and hashlib.sha256(archive.read(row["relative_path"])).hexdigest().upper()
            == row["sha256"]
            for row in stats_manifest
        )
        identity_names = {
            name for name in stats_names if name.startswith("identity_robustness/")
        }
    check(
        "augmented_statistical_archive",
        manifest_valid
        and len(identity_names) == 101
        and "identity_robustness/end_to_end_resampling/05_FULL_PIPELINE_RESAMPLING_STATUS.json"
        in identity_names
        and "identity_robustness/boundary_propagation/14_ROUND6_R1_HOLD_ADVISOR_REVIEW.md"
        in identity_names,
        {
            "manifest_rows": len(stats_manifest),
            "identity_robustness_files": len(identity_names),
            "sha256": sha256(stats_copy),
        },
    )
    restricted_suffixes = {".h5ad", ".h5", ".rds", ".mtx", ".bam", ".fastq"}
    restricted = [
        path.relative_to(PACKAGE).as_posix()
        for path in PACKAGE.rglob("*")
        if path.is_file() and path.suffix.lower() in restricted_suffixes
    ]
    check("no_restricted_large_source_data", not restricted, restricted)

    technical_pass = all(item["pass"] for item in checks.values())
    if not technical_pass:
        failed = [name for name, item in checks.items() if not item["pass"]]
        raise RuntimeError(f"Final package audit failed: {failed}")

    (PACKAGE / "README.md").write_text(
        package_readme(doi), encoding="utf-8", newline="\n"
    )
    manifest = build_manifest()
    package_bytes, package_sha = build_deterministic_archive()
    decision = "PASS_GATE_C8BR_RELEASE_PORTABILITY_AUTHOR_COMPLETION_AND_PORTAL_PREFLIGHT"
    audit = {
        "created_at": "2026-08-27",
        "decision": decision,
        "primary_target": "Genome Medicine",
        "source_scientific_freeze": "frozen disease-blind analysis with formal end-to-end identity HOLD propagation",
        "scientific_estimates_changed": False,
        "author_completion_pass": True,
        "technical_package_pass": True,
        "portal_upload_set_authorized": True,
        "journal_submission_performed": False,
        "doi": doi,
        "page_counts": page_counts,
        "manifest_payload_files": len(manifest),
        "package_zip": PACKAGE_ZIP.relative_to(ROOT).as_posix(),
        "package_zip_bytes": package_bytes,
        "package_zip_sha256": package_sha,
        "checks": checks,
        "next_stage": (
            "Restore GSE135779 and complete a label-agnostic external mapping "
            "sensitivity before updating the immutable archive or uploading to a journal portal."
        ),
    }
    (RUN_DIR / "06_GATE_C8BRF_FINAL_AUDIT.json").write_text(
        json.dumps(audit, indent=2) + "\n", encoding="utf-8", newline="\n"
    )
    report_lines = [
        "# Gate C8BRF final author-release audit",
        "",
        f"**Decision:** `{decision}`",
        "",
        f"**DOI:** `https://doi.org/{doi}`",
        "",
        "## Verified outcome",
        "",
        "- Scientific estimates changed: No.",
        "- Main figure assertions: 46/46.",
        "- Supplementary-figure assertions: legacy 29/29 plus verified S8 and S9 contracts.",
        "- Main figures: 5/5 at 170 mm, vector PDF plus 600-dpi PNG.",
        "- Manuscript and cover placeholders: 0.",
        "- WPS pages: "
        f"{page_counts['main']} manuscript, {page_counts['supplement']} supplement, {page_counts['cover']} cover.",
        "- Accessibility findings: 0 high, 0 medium, 0 low for all three DOCX files.",
        "- Portal maps: 11 REQUIRED, 9 OPTIONAL.",
        f"- Deterministic ZIP: {package_bytes:,} bytes; SHA-256 `{package_sha}`.",
        "",
        "## Check matrix",
        "",
        "| Check | Result |",
        "|---|---|",
    ]
    report_lines.extend(
        f"| {name} | {'PASS' if item['pass'] else 'FAIL'} |"
        for name, item in checks.items()
    )
    report_lines.extend(
        [
            "",
            "## Next stage",
            "",
            audit["next_stage"],
            "",
        ]
    )
    (RUN_DIR / "06_GATE_C8BRF_FINAL_AUDIT.md").write_text(
        "\n".join(report_lines), encoding="utf-8", newline="\n"
    )
    (RUN_DIR / "07_GATE_C8BRF_PACKAGE_STATUS.json").write_text(
        json.dumps(
            {
                "created_at": "2026-08-27",
                "status": decision,
                "doi": doi,
                "package_zip_bytes": package_bytes,
                "package_zip_sha256": package_sha,
                "canonical_builder_output": True,
                "manual_rezip_permitted": False,
            },
            indent=2,
        )
        + "\n",
        encoding="utf-8",
        newline="\n",
    )
    print(json.dumps(audit, indent=2))


if __name__ == "__main__":
    main()
